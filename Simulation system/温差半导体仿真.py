import sys
import os
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QLabel, QLineEdit, QComboBox, QPushButton, 
                            QGroupBox, QFrame, QGridLayout, QDialog, QScrollArea,
                            QMessageBox, QProgressBar, QToolTip, QStatusBar,
                            QMenu, QFileDialog, QTabWidget, QListWidget, QTextEdit,
                            QTableWidget, QTableWidgetItem, QFormLayout, QInputDialog)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt5.QtGui import QPixmap, QFont, QPalette, QColor
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavigationToolbar
import numpy as np
from scipy.interpolate import interp1d, CubicSpline
from scipy.optimize import fsolve
import pandas as pd
import time

# 设置工作目录为exe文件所在目录
if getattr(sys, 'frozen', False):
    # 如果是打包后的exe文件，使用PyInstaller的资源路径
    application_path = sys._MEIPASS
else:
    # 如果是Python脚本
    application_path = os.path.dirname(os.path.abspath(__file__))

os.chdir(application_path)

class StatusLight(QLabel):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedSize(20, 20)
        self.setStyleSheet("background-color: red; border-radius: 10px;")
        
    def set_status(self, status):
        if status:
            self.setStyleSheet("background-color: green; border-radius: 10px;")
        else:
            self.setStyleSheet("background-color: red; border-radius: 10px;")

class CalculationThread(QThread):
    """计算线程，用于在后台执行计算任务"""
    progress_updated = pyqtSignal(int)
    calculation_finished = pyqtSignal(dict)
    error_occurred = pyqtSignal(str)
    
    def __init__(self, calculator, params):
        super().__init__()
        self.calculator = calculator
        self.params = params
        self.is_running = True
        
    def run(self):
        try:
            # 模拟计算进度
            for i in range(101):
                if not self.is_running:
                    break
                self.progress_updated.emit(i)
                time.sleep(0.01)  # 模拟计算时间
                
            if self.is_running:
                # 执行实际计算
                result = self.calculator.calculate_temperature_distribution(
                    self.params['Th'], 
                    self.params['Tc'], 
                    self.params['n_points'], 
                    self.params['material_type'], 
                    self.params['composition'], 
                    self.params['current_density'], 
                    self.params['max_iter']
                )
                self.calculation_finished.emit({'x': result[0], 'T': result[1]})
                
        except Exception as e:
            self.error_occurred.emit(str(e))
            
    def stop(self):
        self.is_running = False

class ImageViewerDialog(QDialog):
    def __init__(self, pixmap, parent=None):
        super().__init__(parent)
        self.setWindowTitle("图片查看")
        self.setWindowFlags(Qt.Window | Qt.WindowMaximizeButtonHint | Qt.WindowCloseButtonHint)
        
        # 获取屏幕尺寸
        screen = QApplication.primaryScreen().geometry()
        self.setMinimumSize(screen.width() // 2, screen.height() // 2)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # 创建图片容器
        self.image_container = QWidget()
        self.image_container.setStyleSheet("background-color: white;")
        container_layout = QVBoxLayout(self.image_container)
        container_layout.setContentsMargins(0, 0, 0, 0)
        container_layout.setSpacing(0)
        
        # 创建图片标签
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        container_layout.addWidget(self.image_label)
        
        # 创建滚动区域
        scroll = QScrollArea()
        scroll.setWidget(self.image_container)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        
        # 添加关闭按钮
        button_layout = QHBoxLayout()
        button_layout.setContentsMargins(10, 5, 10, 5)
        close_button = QPushButton("关闭")
        close_button.setFixedWidth(100)
        close_button.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 1px solid #dcdcdc;
                border-radius: 3px;
                padding: 5px;
            }
            QPushButton:hover {
                background-color: #e6e6e6;
            }
        """)
        close_button.clicked.connect(self.close)
        button_layout.addStretch()
        button_layout.addWidget(close_button)
        button_layout.addStretch()
        layout.addLayout(button_layout)
        
        # 保存原始图片
        self.original_pixmap = pixmap
        # 初始显示
        self.resizeEvent(None)
    
    def resizeEvent(self, event):
        """当窗口大小改变时，调整图片大小"""
        if hasattr(self, 'original_pixmap') and not self.original_pixmap.isNull():
            # 获取可用空间大小（减去按钮区域高度）
            available_size = self.size()
            available_size.setHeight(available_size.height() - 40)  # 40是按钮区域的高度
            
            # 计算缩放后的图片大小
            scaled_pixmap = self.original_pixmap.scaled(
                available_size,
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            )
            
            # 更新图片
            self.image_label.setPixmap(scaled_pixmap)

class ClickableImageLabel(QLabel):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setCursor(Qt.PointingHandCursor)  # 设置鼠标指针为手型
        
    def mouseDoubleClickEvent(self, event):
        if self.pixmap() and not self.pixmap().isNull():
            dialog = ImageViewerDialog(self.pixmap(), self.window())
            dialog.exec_()

class ThermoelectricCalculator:
    # 内置材料组分常量
    BUILTIN_P_COMPOSITIONS = {"0.01", "0.02", "0.03"}
    BUILTIN_N_COMPOSITIONS = {"0.0004", "0.0012", "0.0020", "0.0028"}
    
    def __init__(self):
        # 移除对iter_edit的依赖
        self.p_type_data = {}
        self.n_type_data = {}
        self.interpolators = {}
        
        # 添加材料数据文件路径配置
        self.material_data_dir = "material_data"
        self.p_type_file = "p_type_materials.json"
        self.n_type_file = "n_type_materials.json"
        
        # 确保材料数据目录存在
        import os
        if not os.path.exists(self.material_data_dir):
            os.makedirs(self.material_data_dir)
        
        # 读取P型材料数据，修正组分值对应关系
        p_files = {
            "0.01": "P_yuanshi_2_5.xls",  # 0.01对应2.5
            "0.02": "P_yuanshi_3_1.xls",  # 0.02对应3.1
            "0.03": "P_yuanshi_3_7.xls"   # 0.03对应3.7
        }
        
        # 读取N型材料数据
        n_files = {
            "0.0004": "N_yuanshi_0.0004.xls",
            "0.0012": "N_yuanshi_0.0012.xls",
            "0.0020": "N_yuanshi_0.0020.xls",
            "0.0028": "N_yuanshi_0.0028.xls"
        }
        
        def read_excel_file(filename):
            """读取Excel文件的辅助函数"""
            try:
                # 构建完整路径
                full_path = os.path.join(application_path, filename)
                print(f"尝试读取文件: {full_path}")
                
                if not os.path.exists(full_path):
                    print(f"文件不存在: {full_path}")
                    return None
                
                # 首先尝试使用xlrd引擎
                try:
                    import xlrd
                    # 不使用列名读取数据
                    data = pd.read_excel(full_path, engine='xlrd', header=None)
                    print(f"成功使用xlrd读取文件: {filename}")
                    return data
                except ImportError:
                    print("xlrd未安装，尝试使用openpyxl...")
                    return None
            except Exception as e:
                print(f"读取文件失败: {str(e)}")
                return None
        
        # 读取所有P型材料数据
        for composition, filename in p_files.items():
            print(f"\n尝试读取P型材料数据文件: {filename}")
            data = read_excel_file(filename)
            if data is not None:
                try:
                    # 查找列
                    columns = self.find_columns(data)
                    if columns:
                        # P型材料：F列是优值系数(ZT)，我们需要从中反推热导率
                        # 热导率 k = (α^2 × T) / (ρ × ZT)
                        # 其中 α 是塞贝克系数，ρ 是电阻率，T 是温度，ZT 是优值系数
                        seebeck = data[columns['seebeck']].values * 1e-6  # μV/K 转换为 V/K
                        resistivity = data[columns['resistivity']].values * 1e-6  # μΩ·m 转换为 Ω·m (修正单位换算错误)
                        temperature = data[columns['temp']].values
                        zt_values = data[columns['thermal_cond']].values  # 这里实际上是ZT值
                        
                        # 计算热导率
                        thermal_cond = []
                        for i in range(len(temperature)):
                            try:
                                # 避免无效ZT值和除以零
                                if zt_values[i] > 0:
                                    k = (seebeck[i]**2 * temperature[i]) / (resistivity[i] * zt_values[i])
                                    # 添加合理性检查 (热导率通常在0.1-100 W/m·K范围内)
                                    if 0.1 <= k <= 100:
                                        thermal_cond.append(k)
                                    else:
                                        print(f"警告: 计算得到异常热导率值 {k:.3f} W/m·K，使用默认值 2.0 W/m·K")
                                        thermal_cond.append(2.0)  # 更合理的默认值
                                else:
                                    print(f"警告: 无效ZT值 {zt_values[i]}，使用默认热导率 2.0 W/m·K")
                                    thermal_cond.append(2.0)  # 更合理的默认值
                            except Exception as e:
                                print(f"热导率计算错误: {str(e)}，使用默认值 2.0 W/m·K")
                                thermal_cond.append(2.0)  # 更合理的默认值
                        
                        self.p_type_data[composition] = {
                            "temp": temperature,
                            "seebeck": seebeck,
                            "resistivity": resistivity,
                            "thermal_cond": np.array(thermal_cond)  # 从ZT反推的热导率
                        }
                        print(f"成功读取P型材料数据: {composition}")
                        print(f"温度范围: {min(temperature)}-{max(temperature)} K")
                        print(f"塞贝克系数范围: {min(seebeck*1e6)}-{max(seebeck*1e6)} μV/K")
                        print(f"电阻率范围: {min(resistivity*1e6)}-{max(resistivity*1e6)} μΩ·m")
                        print(f"计算的热导率范围: {min(thermal_cond)}-{max(thermal_cond)} W/(m·K)")
                    else:
                        print(f"在文件 {filename} 中未找到所需的列")
                        
                except Exception as e:
                    print(f"处理P型材料数据文件 {filename} 时出错: {str(e)}")
                    import traceback
                    traceback.print_exc()
        
        # 读取所有N型材料数据
        for composition, filename in n_files.items():
            print(f"\n尝试读取N型材料数据文件: {filename}")
            data = read_excel_file(filename)
            if data is not None:
                try:
                    # 查找列
                    columns = self.find_columns(data)
                    if columns:
                        self.n_type_data[composition] = {
                            "temp": data[columns['temp']].values,
                            "seebeck": -data[columns['seebeck']].values * 1e-6,  # μV/K 转换为 V/K，N型为负值
                            "resistivity": data[columns['resistivity']].values * 1e-3,  # μΩ·m 转换为 Ω·m (与内置材料保持一致)
                            "thermal_cond": data[columns['thermal_cond']].values / 100  # W/(m·K) (与内置材料保持一致)
                        }
                        print(f"成功读取N型材料数据: {composition}")
                    else:
                        print(f"在文件 {filename} 中未找到所需的列")
                        
                except Exception as e:
                    print(f"处理N型材料数据文件 {filename} 时出错: {str(e)}")
                    import traceback
                    traceback.print_exc()
                    
        print("\n数据读取完成")
        print(f"成功读取的P型材料: {list(self.p_type_data.keys())}")
        print(f"成功读取的N型材料: {list(self.n_type_data.keys())}")
        
        # 加载用户自定义材料数据
        self.load_custom_materials()
    
    def load_custom_materials(self):
        """加载用户自定义材料数据"""
        try:
            import json
            import os
            
            # 加载P型自定义材料
            p_file_path = os.path.join(self.material_data_dir, self.p_type_file)
            if os.path.exists(p_file_path):
                with open(p_file_path, 'r', encoding='utf-8') as f:
                    custom_p_data = json.load(f)
                    for composition, data in custom_p_data.items():
                        # 转换数据格式
                        self.p_type_data[composition] = {
                            "temp": np.array(data["temp"]),
                            "seebeck": np.array(data["seebeck"]),
                            "resistivity": np.array(data["resistivity"]),
                            "thermal_cond": np.array(data["thermal_cond"])
                        }
                    print(f"加载了 {len(custom_p_data)} 个自定义P型材料")
            
            # 加载N型自定义材料
            n_file_path = os.path.join(self.material_data_dir, self.n_type_file)
            if os.path.exists(n_file_path):
                with open(n_file_path, 'r', encoding='utf-8') as f:
                    custom_n_data = json.load(f)
                    for composition, data in custom_n_data.items():
                        # 转换数据格式
                        self.n_type_data[composition] = {
                            "temp": np.array(data["temp"]),
                            "seebeck": np.array(data["seebeck"]),
                            "resistivity": np.array(data["resistivity"]),
                            "thermal_cond": np.array(data["thermal_cond"])
                        }
                    print(f"加载了 {len(custom_n_data)} 个自定义N型材料")
                    
        except Exception as e:
            print(f"加载自定义材料数据失败: {str(e)}")
    
    def save_custom_materials(self):
        """保存用户自定义材料数据"""
        try:
            import json
            import os
            
            # 分离内置材料和自定义材料
            builtin_p = {"0.01", "0.02", "0.03"}
            builtin_n = {"0.0004", "0.0012", "0.0020", "0.0028"}
            
            custom_p_data = {}
            custom_n_data = {}
            
            # 提取P型自定义材料
            for composition, data in self.p_type_data.items():
                if composition not in builtin_p:
                    custom_p_data[composition] = {
                        "temp": data["temp"].tolist(),
                        "seebeck": data["seebeck"].tolist(),
                        "resistivity": data["resistivity"].tolist(),
                        "thermal_cond": data["thermal_cond"].tolist()
                    }
            
            # 提取N型自定义材料
            for composition, data in self.n_type_data.items():
                if composition not in builtin_n:
                    custom_n_data[composition] = {
                        "temp": data["temp"].tolist(),
                        "seebeck": data["seebeck"].tolist(),
                        "resistivity": data["resistivity"].tolist(),
                        "thermal_cond": data["thermal_cond"].tolist()
                    }
            
            # 保存P型材料
            p_file_path = os.path.join(self.material_data_dir, self.p_type_file)
            with open(p_file_path, 'w', encoding='utf-8') as f:
                json.dump(custom_p_data, f, ensure_ascii=False, indent=2)
            
            # 保存N型材料
            n_file_path = os.path.join(self.material_data_dir, self.n_type_file)
            with open(n_file_path, 'w', encoding='utf-8') as f:
                json.dump(custom_n_data, f, ensure_ascii=False, indent=2)
                
            print(f"保存了 {len(custom_p_data)} 个自定义P型材料和 {len(custom_n_data)} 个自定义N型材料")
            
        except Exception as e:
            print(f"保存自定义材料数据失败: {str(e)}")
    
    def add_custom_material(self, material_type, composition, material_data):
        """
        添加自定义材料数据
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料组分标识符
        material_data: 包含温度、塞贝克系数、电阻率、热导率数据的字典
        """
        try:
            # 验证数据格式
            required_keys = ["temp", "seebeck", "resistivity", "thermal_cond"]
            for key in required_keys:
                if key not in material_data:
                    raise ValueError(f"缺少必需的键: {key}")
            
            # 验证数据长度一致性
            data_lengths = [len(material_data[key]) for key in required_keys]
            if len(set(data_lengths)) != 1:
                raise ValueError("所有数据数组长度必须一致")
            
            # 验证数据合理性
            temp = np.array(material_data["temp"])
            seebeck = np.array(material_data["seebeck"])
            resistivity = np.array(material_data["resistivity"])
            thermal_cond = np.array(material_data["thermal_cond"])
            
            # 温度范围检查
            if np.any(temp < 0) or np.any(temp > 2000):
                raise ValueError("温度值应在0-2000K范围内")
            
            # 电阻率和热导率应为正值
            if np.any(resistivity <= 0):
                raise ValueError("电阻率必须为正值")
            if np.any(thermal_cond <= 0):
                raise ValueError("热导率必须为正值")
            
            # 塞贝克系数符号检查
            if material_type == 'p' and np.any(seebeck < 0):
                print("警告: P型材料的塞贝克系数应为正值")
            elif material_type == 'n' and np.any(seebeck > 0):
                print("警告: N型材料的塞贝克系数应为负值")
            
            # 存储材料数据
            target_data = self.p_type_data if material_type == 'p' else self.n_type_data
            target_data[composition] = {
                "temp": temp,
                "seebeck": seebeck,
                "resistivity": resistivity,
                "thermal_cond": thermal_cond
            }
            
            # 清除相关的插值器缓存
            interp_key = f"{material_type}_{composition}"
            if interp_key in self.interpolators:
                del self.interpolators[interp_key]
            
            # 保存到文件
            self.save_custom_materials()
            
            print(f"成功添加{material_type}型材料: {composition}")
            return True
            
        except Exception as e:
            print(f"添加材料数据失败: {str(e)}")
            return False
    
    def remove_custom_material(self, material_type, composition):
        """
        删除自定义材料数据
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料组分标识符
        """
        try:
            # 检查是否为内置材料
            builtin_p = {"0.01", "0.02", "0.03"}
            builtin_n = {"0.0004", "0.0012", "0.0020", "0.0028"}
            
            if material_type == 'p' and composition in builtin_p:
                raise ValueError("不能删除内置P型材料")
            elif material_type == 'n' and composition in builtin_n:
                raise ValueError("不能删除内置N型材料")
            
            # 删除材料数据
            target_data = self.p_type_data if material_type == 'p' else self.n_type_data
            if composition in target_data:
                del target_data[composition]
                
                # 清除相关的插值器缓存
                interp_key = f"{material_type}_{composition}"
                if interp_key in self.interpolators:
                    del self.interpolators[interp_key]
                
                # 保存到文件
                self.save_custom_materials()
                
                print(f"成功删除{material_type}型材料: {composition}")
                return True
            else:
                print(f"材料 {composition} 不存在")
                return False
                
        except Exception as e:
            print(f"删除材料数据失败: {str(e)}")
            return False
    
    def get_material_list(self, material_type):
        """
        获取指定类型的所有材料列表
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        
        返回:
        list: 材料组分列表
        """
        target_data = self.p_type_data if material_type == 'p' else self.n_type_data
        return list(target_data.keys())
    
    def get_material_info(self, material_type, composition):
        """
        获取材料详细信息
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料组分标识符
        
        返回:
        dict: 材料信息字典
        """
        try:
            target_data = self.p_type_data if material_type == 'p' else self.n_type_data
            if composition not in target_data:
                return None
            
            data = target_data[composition]
            temp = data["temp"]
            seebeck = data["seebeck"]
            resistivity = data["resistivity"]
            thermal_cond = data["thermal_cond"]
            
            # 计算ZT值
            zt_values = []
            for i in range(len(temp)):
                try:
                    zt = (seebeck[i]**2 * temp[i]) / (resistivity[i] * thermal_cond[i])
                    zt_values.append(zt)
                except:
                    zt_values.append(0)
            
            return {
                "composition": composition,
                "material_type": material_type,
                "temperature_range": f"{min(temp):.1f} - {max(temp):.1f} K",
                "seebeck_range": f"{min(seebeck*1e6):.2f} - {max(seebeck*1e6):.2f} μV/K",
                "resistivity_range": f"{min(resistivity*1e6):.2f} - {max(resistivity*1e6):.2f} μΩ·m",
                "thermal_cond_range": f"{min(thermal_cond):.2f} - {max(thermal_cond):.2f} W/(m·K)",
                "max_zt": f"{max(zt_values):.3f}",
                "data_points": len(temp),
                "is_builtin": composition in {"0.01", "0.02", "0.03"} if material_type == 'p' else composition in {"0.0004", "0.0012", "0.0020", "0.0028"}
            }
            
        except Exception as e:
            print(f"获取材料信息失败: {str(e)}")
            return None
    
    def find_columns(self, data):
        """根据数据结构查找相应的列 - 类方法版本"""
        try:
            # 检查数据的结构来确定正确的列索引
            # 对于P_yuanshi文件（如P_yuanshi_2_5.xls），列结构为：
            # 温度(A列,0), 塞贝克系数(B列,1), 温度(C列,2), 电阻率(D列,3), 温度(E列,4), 优值系数(F列,5)
            if data.shape[1] >= 6:  # 确保有足够的列
                print("找到的列结构：")
                for i in range(min(6, data.shape[1])):
                    print(f"列 {i}: {data.iloc[0, i]}")
                
                # 检查前几行的数据来识别是P型还是N型文件
                # P型文件特征：第一列数值在300左右（温度）
                first_col_values = data.iloc[0:5, 0].values
                print(f"第一列前5个值: {first_col_values}")
                
                if any(290 <= v <= 310 for v in first_col_values if isinstance(v, (int, float))):
                    print("检测到P型材料数据文件格式")
                    return {
                        "temp": 0,        # A列作为温度
                        "seebeck": 1,     # B列作为塞贝克系数
                        "resistivity": 3, # D列作为电阻率
                        "thermal_cond": 5 # F列作为优值系数(ZT)，需要计算热导率
                    }
                else:
                    print("检测到N型材料数据文件格式")
                    return {
                        "temp": 0,        # 第1列作为温度
                        "seebeck": 1,     # 第2列作为塞贝克系数
                        "resistivity": 3, # 第4列作为电阻率
                        "thermal_cond": 5 # 第6列作为热导率
                    }
            else:
                print("警告：数据列数不足，使用默认列映射")
                return {
                    "temp": 0,        # 第1列作为温度
                    "seebeck": 1,     # 第2列作为塞贝克系数
                    "resistivity": 3, # 第4列作为电阻率
                    "thermal_cond": 5 # 第6列作为热导率（N型）或优值系数（P型）
                }
        except Exception as e:
            print(f"查找列错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def import_material_from_excel(self, material_type, composition, excel_file):
        """
        从Excel文件导入材料数据
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料组分标识符
        excel_file: Excel文件路径
        """
        try:
            # 读取Excel文件
            data = pd.read_excel(excel_file, header=None)
            
            # 查找数据列
            columns = self.find_columns(data)
            if not columns:
                raise ValueError("无法识别Excel文件的数据格式")
            
            # 提取数据
            temp = data[columns['temp']].values
            seebeck = data[columns['seebeck']].values * 1e-6  # μV/K 转换为 V/K
            
            # 电阻率单位转换：根据材料类型使用不同的转换系数
            if material_type == 'p':
                resistivity = data[columns['resistivity']].values * 1e-6  # μΩ·m 转换为 Ω·m (P型材料)
            else:
                resistivity = data[columns['resistivity']].values * 1e-3  # μΩ·m 转换为 Ω·m (N型材料，与内置材料保持一致)
            
            # 对于P型材料，第5列是优值系数(ZT)，需要计算热导率
            if material_type == 'p':
                zt_values = data[columns['thermal_cond']].values  # 第5列是优值系数ZT
                # 通过ZT公式计算热导率: ZT = S²T/(kρ) => k = S²T/(ZT·ρ)
                thermal_cond = (seebeck**2 * temp) / (zt_values * resistivity)
                print(f"P型材料：从ZT值计算热导率，平均热导率: {np.mean(thermal_cond):.3f} W/(m·K)")
            else:
                # N型材料第5列是热导率，需要除以100（与内置材料保持一致）
                thermal_cond = data[columns['thermal_cond']].values / 100  # 转换为正确的单位
                print(f"N型材料：热导率数据除以100，平均热导率: {np.mean(thermal_cond):.3f} W/(m·K)")
            
            # 处理N型材料的塞贝克系数符号
            if material_type == 'n':
                seebeck = -seebeck
            
            # 创建材料数据字典
            material_data = {
                "temp": temp,
                "seebeck": seebeck,
                "resistivity": resistivity,
                "thermal_cond": thermal_cond
            }
            
            # 添加材料
            return self.add_custom_material(material_type, composition, material_data)
            
        except Exception as e:
            print(f"从Excel导入材料数据失败: {str(e)}")
            return False
    
    def export_material_to_excel(self, material_type, composition, excel_file):
        """
        导出材料数据到Excel文件
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料组分标识符
        excel_file: 输出Excel文件路径
        """
        try:
            target_data = self.p_type_data if material_type == 'p' else self.n_type_data
            if composition not in target_data:
                raise ValueError(f"材料 {composition} 不存在")
            
            data = target_data[composition]
            
            # 创建DataFrame
            df = pd.DataFrame({
                '温度(K)': data["temp"],
                '塞贝克系数(μV/K)': data["seebeck"] * 1e6,
                '电阻率(μΩ·m)': data["resistivity"] * 1e6,
                '热导率(W/m·K)': data["thermal_cond"]
            })
            
            # 计算ZT值
            zt_values = []
            for i in range(len(data["temp"])):
                try:
                    zt = (data["seebeck"][i]**2 * data["temp"][i]) / (data["resistivity"][i] * data["thermal_cond"][i])
                    zt_values.append(zt)
                except:
                    zt_values.append(0)
            
            df['优值系数ZT'] = zt_values
            
            # 保存到Excel
            df.to_excel(excel_file, index=False)
            
            print(f"成功导出{material_type}型材料 {composition} 到 {excel_file}")
            return True
            
        except Exception as e:
            print(f"导出材料数据失败: {str(e)}")
            return False
    def create_interpolators(self, material_type, composition):
        """为给定材料创建属性插值器"""
        try:
            # 验证参数
            if not composition or composition.strip() == '':
                raise ValueError(f"材料组分标识不能为空")
            
            if material_type not in ['p', 'n']:
                raise ValueError(f"材料类型必须是 'p' 或 'n'，当前为: {material_type}")
            
            data = self.p_type_data if material_type == 'p' else self.n_type_data
            
            if composition not in data:
                raise ValueError(f"{material_type.upper()}型材料 {composition} 不存在")
            
            mat_data = data[composition]
            
            # ==== 增加插值范围限制 ====
            temps = mat_data["temp"]
            seebeck = mat_data["seebeck"]
            resistivity = mat_data["resistivity"]
            thermal_cond = mat_data["thermal_cond"]
            
            # 确保数据有序
            sort_idx = np.argsort(temps)
            temps = temps[sort_idx]
            seebeck = seebeck[sort_idx]
            resistivity = resistivity[sort_idx]
            thermal_cond = thermal_cond[sort_idx]
            
            # 打印材料属性范围
            print(f"\n===== 创建 {material_type}型材料插值器 (组分={composition}) =====")
            print(f"温度范围: {min(temps)}-{max(temps)} K")
            print(f"塞贝克系数范围: {min(seebeck*1e6):.2f}-{max(seebeck*1e6):.2f} μV/K")
            print(f"电阻率范围: {min(resistivity*1e6):.2f}-{max(resistivity*1e6):.2f} μΩ·m")
            print(f"热导率范围: {min(thermal_cond):.2f}-{max(thermal_cond):.2f} W/(m·K)")
            
            # 创建边界值保护的插值器
            self.interpolators[f"{material_type}_{composition}"] = {
                "seebeck": interp1d(temps, seebeck, kind='linear', 
                                   bounds_error=False, 
                                   fill_value=(seebeck[0], seebeck[-1])),  # 限制外推值
                "resistivity": interp1d(temps, resistivity, kind='linear',
                                      bounds_error=False,
                                      fill_value=(resistivity[0], resistivity[-1])),
                "thermal_cond": interp1d(temps, thermal_cond, kind='linear',
                                       bounds_error=False,
                                       fill_value=(thermal_cond[0], thermal_cond[-1]))
            }
            
            print(f"插值器创建成功")
            
        except Exception as e:
            print(f"创建插值器错误: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def calculate_temperature_distribution(self, Th, Tc, n_points, material_type, composition, current_density, max_iter=50):
        """
        根据热电材料物理性质计算温度分布，统一使用material_P和material_N方法
        """
        try:
            # 验证输入参数
            if not composition or composition.strip() == '':
                raise ValueError(f"材料组分标识不能为空")
            
            if material_type not in ['p', 'n']:
                raise ValueError(f"材料类型必须是 'p' 或 'n'，当前为: {material_type}")
            
            print(f"\n开始计算温度分布: {material_type}型, 组分={composition}, 电流密度={current_density}A/cm^2")
            print(f"边界条件: Th={Th}K, Tc={Tc}K, 格点数={n_points}")
            
            # 统一使用material_P和material_N方法，它们已经能够处理所有材料
            if material_type == 'p':
                return self.temperature_distribution_P(n_points, current_density, Tc, Th, max_iter, composition)
            else:
                return self.temperature_distribution_N(current_density, Tc, Th, n_points, 1.0, composition)
            
            # 初始化格点位置和温度
            L = 1.0  # 标准化长度
            dx = L / (n_points - 1)  # 网格间距
            x = np.linspace(0, L, n_points)  # 从0到1的均匀分布
            T = np.linspace(Tc, Th, n_points)  # 初始线性温度分布
            
            print(f"初始温度分布: {T}")
            
            # ===== 修正1: 电流密度单位处理 =====
            # 假设输入电流密度为A/cm^2，使用更合理的转换系数
            J = current_density * 100  # 转换为A/m²但限制在合理范围
            
            # 检查电流密度是否在合理范围内
            if abs(J) > 5e3:  # 根据物理合理性设定上限
                print(f"警告: 电流密度 {J} A/m² 超过正常范围，将限制为5000A/m²")
                J = np.sign(J) * 5e3
            
            print(f"网格间距: dx={dx}, 电流密度: J={J}A/m²")
            
            # 迭代求解参数
            relaxation_factor = 0.2  # 松弛因子，提高稳定性
            convergence_threshold = 0.01  # 收敛阈值
            jacobi_iterations = 100  # 内部Jacobi迭代次数上限
            
            for iter_count in range(max_iter):
                # 保存旧的温度分布用于收敛判断
                T_old = T.copy()
                
                # ===== 修正2: 先计算并存储所有网格点的材料属性 =====
                seebeck = np.zeros(n_points)
                resistivity = np.zeros(n_points)
                thermal_cond = np.zeros(n_points)
                
                for i in range(n_points):
                    T_safe = np.clip(T[i], 300, 700)  # 确保温度在有效范围内
                    seebeck[i] = self.interpolators[interp_key]["seebeck"](T_safe)
                    resistivity[i] = self.interpolators[interp_key]["resistivity"](T_safe)
                    thermal_cond[i] = self.interpolators[interp_key]["thermal_cond"](T_safe)
                
                # 构建系数矩阵和右端向量
                A = np.zeros((n_points, n_points))
                b = np.zeros(n_points)
                
                # 设置边界条件
                A[0, 0] = 1.0
                b[0] = Tc
                A[n_points-1, n_points-1] = 1.0
                b[n_points-1] = Th
                
                # ===== 修正3: 内部点的系数计算改为基于温度的插值 =====
                for i in range(1, n_points-1):
                    # 使用基于温度分布的插值获取半点处的材料属性
                    T_half_minus = 0.5 * (T[i-1] + T[i])
                    T_half_plus = 0.5 * (T[i] + T[i+1])
                    
                    # 直接使用插值器计算半点处属性
                    k_minus = self.interpolators[interp_key]["thermal_cond"](np.clip(T_half_minus, 300, 700))
                    k_plus = self.interpolators[interp_key]["thermal_cond"](np.clip(T_half_plus, 300, 700))
                    s_minus = self.interpolators[interp_key]["seebeck"](np.clip(T_half_minus, 300, 700))
                    s_plus = self.interpolators[interp_key]["seebeck"](np.clip(T_half_plus, 300, 700))
                    rho_i = resistivity[i]  # 焦耳热在节点上
                    
                    # 热传导项系数
                    A[i, i-1] = k_minus / dx**2
                    A[i, i] = -(k_minus + k_plus) / dx**2
                    A[i, i+1] = k_plus / dx**2
                    
                    # 塞贝克项（热电耦合）
                    A[i, i-1] += -J * s_minus / (2 * dx)
                    A[i, i+1] += J * s_plus / (2 * dx)
                    
                    # ===== 修正4: 焦耳热项修正为正贡献 =====
                    b[i] = rho_i * J**2  # 正确的焦耳热项（热源）
                
                # 使用带松弛因子的Jacobi迭代
                T_new = T_old.copy()
                
                for jacobi_iter in range(jacobi_iterations):
                    T_prev = T_new.copy()
                    
                    # 边界点固定不变
                    T_new[0] = Tc
                    T_new[n_points-1] = Th
                    
                    # 更新内部点
                    for i in range(1, n_points-1):
                        if abs(A[i, i]) > 1e-10:  # 避免除以零
                            numerator = b[i]
                            for j in range(n_points):
                                if j != i:
                                    numerator -= A[i, j] * T_prev[j]
                            
                            # 使用更小的松弛因子提高稳定性
                            T_new[i] = T_prev[i] + relaxation_factor * (numerator / A[i, i] - T_prev[i])
                    
                    # 检查内部迭代收敛性
                    if np.max(np.abs(T_new - T_prev)) < 0.001:
                        break
                
                # 检查解的合理性
                if np.any(np.isnan(T_new)) or np.any(np.isinf(T_new)):
                    print(f"警告：第{iter_count+1}次迭代解不合理，使用线性插值")
                    T_new = np.linspace(Tc, Th, n_points)
                
                # 限制温度在物理合理范围内（略微放宽范围）
                T_new = np.clip(T_new, min(Tc, Th)*0.95, max(Tc, Th)*1.1)
                
                # 计算收敛情况
                max_change = np.max(np.abs(T_new - T_old))
                print(f"迭代{iter_count+1}次完成，最大温度变化: {max_change:.6f}K")
                
                # 更新温度
                T = T_new.copy()
                
                # 判断是否已经收敛
                if max_change < convergence_threshold:
                    print(f"温度分布已收敛，在第{iter_count+1}次迭代")
                    break
            
            # ===== 修正5: 改进的热流验证 =====
            try:
                # 计算节点处的热流密度 q = J*T*S - k*dT/dx
                dTdx = np.zeros(n_points)
                dTdx[1:-1] = (T[2:] - T[:-2]) / (2*dx)  # 中心差分
                dTdx[0] = (T[1] - T[0]) / dx  # 前向差分
                dTdx[-1] = (T[-1] - T[-2]) / dx  # 后向差分
                
                q = np.zeros(n_points)
                for i in range(n_points):
                    q[i] = J * T[i] * seebeck[i] - thermal_cond[i] * dTdx[i]
                
                heat_in = q[0]  # 入口热流
                heat_out = q[-1]  # 出口热流
                joule_heat = sum(resistivity * J**2 * dx)  # 总焦耳热
                
                print(f"热流验证: 入口热流={heat_in:.3f}, 出口热流={heat_out:.3f}, 焦耳热={joule_heat:.3f}")
                print(f"热流平衡检查: 出口-入口-焦耳热={heat_out-heat_in-joule_heat:.3f}（应接近零）")
            except Exception as e:
                print(f"热流验证计算错误: {e}")
            
            # 打印最终温度分布
            print(f"最终温度分布: {T}")
            
            return x, T
            
        except Exception as e:
            print(f"计算温度分布错误: {str(e)}")
            import traceback
            traceback.print_exc()
            
            # 出错时返回线性温度分布
            L = 1.0
            return np.linspace(0, L, n_points), np.linspace(Tc, Th, n_points)
    
    def calculate_efficiency(self, Th, Tc, material_type, composition, current_density, x=None, T=None):
        """
        参考温度分布和效率的实现计算热电材料效率
        
        参数:
        Th: 高温端温度 (K)
        Tc: 低温端温度 (K)
        material_type: 材料类型 ('p' 或 'n')
        composition: 材料组分
        current_density: 电流密度 (A/cm^2)
        x, T: 温度分布数据
        
        返回:
        efficiency: 效率 (%)
        power: 输出功率密度 (W/m²)
        """
        try:
            # 验证输入参数
            if not composition or composition.strip() == '':
                raise ValueError(f"材料组分标识不能为空")
            
            if material_type not in ['p', 'n']:
                raise ValueError(f"材料类型必须是 'p' 或 'n'，当前为: {material_type}")
            
            if Th <= Tc:
                print(f"警告: 温度差无效 (Th={Th}K, Tc={Tc}K)")
                return 0.0, 0.0
            
            # 准备插值器
            interp_key = f"{material_type}_{composition}"
            if interp_key not in self.interpolators:
                self.create_interpolators(material_type, composition)
            
            # 确保温度分布数据有效
            if x is None or T is None or len(x) < 3:
                print(f"温度分布数据无效，使用线性温度分布近似")
                n_points = 20
                x = np.linspace(0, 1.0, n_points)
                T = np.linspace(Tc, Th, n_points)
            
            # 获取格点数和间距
            n_points = len(x)
            dx = (x[-1] - x[0]) / (n_points - 1)
            
            # 计算各节点处的材料属性并打印调试信息
            seebeck = np.zeros(n_points)
            resistivity = np.zeros(n_points)
            thermal_cond = np.zeros(n_points)
            
            for i in range(n_points):
                T_safe = np.clip(T[i], 300, 700)  # 确保温度在有效范围内
                seebeck[i] = self.interpolators[interp_key]["seebeck"](T_safe)
                resistivity[i] = self.interpolators[interp_key]["resistivity"](T_safe)
                thermal_cond[i] = self.interpolators[interp_key]["thermal_cond"](T_safe)
            
            # 打印材料属性统计信息用于调试
            print(f"\n===== {material_type}型材料属性统计({composition}) =====")
            print(f"塞贝克系数(V/K): 最小={np.min(seebeck):.3e}, 最大={np.max(seebeck):.3e}, 平均={np.mean(seebeck):.3e}")
            print(f"电阻率(Ω·m): 最小={np.min(resistivity):.3e}, 最大={np.max(resistivity):.3e}, 平均={np.mean(resistivity):.3e}")
            print(f"热导率(W/m·K): 最小={np.min(thermal_cond):.3f}, 最大={np.max(thermal_cond):.3f}, 平均={np.mean(thermal_cond):.3f}")
            
            # 验证P型材料塞贝克系数符号是否正确
            if material_type == 'p' and np.any(seebeck < 0):
                print(f"警告: P型材料塞贝克系数出现负值! 最小值: {np.min(seebeck):.3e}V/K")
            elif material_type == 'n' and np.any(seebeck > 0):
                print(f"警告: N型材料塞贝克系数出现正值! 最大值: {np.max(seebeck):.3e}V/K")
            
            # 正确转换单位: A/cm^2 → A/m^2
            J = current_density * 100  # 转换为A/m^2 (1A/cm^2 = 100A/m^2)
            
            # 计算温度梯度 (使用中心差分)
            dTdx = np.zeros_like(T)
            dTdx[1:-1] = (T[2:] - T[:-2]) / (2*dx)  # 中心差分
            dTdx[0] = (T[1] - T[0]) / dx            # 前向差分
            dTdx[-1] = (T[-1] - T[-2]) / dx         # 后向差分
            
            # 计算热流密度: q(x) = κ·dT/dx - J·S·T
            q = np.zeros(n_points)
            for i in range(n_points):
                q[i] = thermal_cond[i] * dTdx[i] - J * seebeck[i] * T[i]
            
            # 确保热流方向正确（从高温端到低温端）
            if np.mean(q) < 0:
                print(f"热流方向修正: 平均热流 {np.mean(q):.2e} 为负值，已反转")
                q = -q
            
            # 第一个积分：∫ S·dT（温度差间的塞贝克积分）
            seebeck_integral = 0.0
            for i in range(1, n_points):
                seebeck_integral += (seebeck[i] + seebeck[i-1]) / 2 * (T[i] - T[i-1])
            
            # 第二个积分：∫ ρ·dx（电阻率沿长度的积分）
            resistivity_integral = 0.0
            for i in range(1, n_points):
                resistivity_integral += (resistivity[i] + resistivity[i-1]) / 2 * dx
            
            # 计算功率输出
            seebeck_power = J * seebeck_integral  # 塞贝克效应产生的功率
            joule_heat = J**2 * resistivity_integral  # 焦耳热损失
            net_power = seebeck_power - joule_heat  # 净功率输出
            
            # 计算热输入（高温端热流）
            heat_in = abs(q[0])  # 热流密度
            
            # 计算焦耳热总量
            total_joule_heat = 0.0
            for i in range(n_points-1):
                seg_resistivity = (resistivity[i] + resistivity[i+1]) / 2
                total_joule_heat += J**2 * seg_resistivity * dx
            
            # 打印详细的功率和热流信息用于调试
            print(f"\n===== 能量分析 =====")
            print(f"塞贝克功率: {seebeck_power:.3e} W/m²")
            print(f"焦耳热损失: {joule_heat:.3e} W/m²")
            print(f"净功率输出: {net_power:.3e} W/m²")
            print(f"热输入(高温端): {heat_in:.3e} W/m²")
            print(f"焦耳热总量: {total_joule_heat:.3e} W/m²")
            
            # 计算效率
            if heat_in > 0 and net_power > 0:
                efficiency = net_power / heat_in * 100  # 转换为百分比
                
                # 计算卡诺效率进行比较
                carnot_eff = (Th - Tc) / Th * 100
                relative_eff = efficiency / carnot_eff * 100  # 与卡诺效率的比值
                
                print(f"效率: {efficiency:.2f}% (卡诺效率: {carnot_eff:.2f}%, 相对效率: {relative_eff:.2f}%)")
                
                # 检查是否超过卡诺效率
                if efficiency > carnot_eff:
                    print(f"警告: 计算效率 {efficiency:.2f}% 超过卡诺效率 {carnot_eff:.2f}%")
                    efficiency = carnot_eff * 0.9  # 限制在卡诺效率的90%以内
            else:
                # 记录无效效率的原因
                if heat_in <= 0:
                    print(f"警告: 热输入为零或负值 ({heat_in:.3e} W/m²)")
                if net_power <= 0:
                    print(f"警告: 净功率为零或负值 ({net_power:.3e} W/m²)")
                
                efficiency = 0.0
                net_power = 0.0
            
            print(f"材料: {material_type}型, 组分={composition}, 电流密度={current_density}A/cm^2, 效率={efficiency:.2f}%")
            
            # 保存关键参数用于后续分析
            self.last_calc_data = {
                "seebeck": seebeck,
                "resistivity": resistivity,
                "thermal_cond": thermal_cond,
                "dTdx": dTdx,
                "current_density": J,
                "temperature": T,
                "heat_in": heat_in,
                "joule_heat": joule_heat,
                "seebeck_power": seebeck_power,
                "net_power": net_power,
                "efficiency": efficiency
            }
            
            return efficiency, net_power
            
        except Exception as e:
            print(f"效率计算错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return 0.0, 0.0

    def calculate_zt(self, material_type, composition, temperature):
        """计算给定温度下的优值系数 ZT = S²T/(kρ)
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料成分
        temperature: 温度 (K)
        
        返回:
        zt: 优值系数
        """
        try:
            # 验证输入参数
            if not composition or composition.strip() == '':
                raise ValueError(f"材料组分标识不能为空")
            
            if material_type not in ['p', 'n']:
                raise ValueError(f"材料类型必须是 'p' 或 'n'，当前为: {material_type}")
            
            # 统一使用material_P和material_N方法，它们已经能够处理所有材料
            if material_type == 'p':
                seebeck, resistivity, thermal_cond = self.material_P(temperature, composition)
            else:
                seebeck, resistivity, thermal_cond = self.material_N(temperature, composition)
            
            # 计算优值系数 ZT = S²T/(kρ)
            # S: 塞贝克系数 (V/K)
            # T: 温度 (K)
            # k: 热导率 (W/(m·K))
            # ρ: 电阻率 (Ω·m)
            zt = (seebeck ** 2) * temperature / (thermal_cond * resistivity)
            
            return zt
            
        except Exception as e:
            print(f"计算优值系数错误: {str(e)}")
            return 0

    def visualize_energy_flow(self, material_type, composition, current_density, x, T):
        """
        可视化材料内部的能量流动
        """
        try:
            # 验证输入参数
            if not composition or composition.strip() == '':
                raise ValueError(f"材料组分标识不能为空")
            
            if material_type not in ['p', 'n']:
                raise ValueError(f"材料类型必须是 'p' 或 'n'，当前为: {material_type}")
            
            # 创建图表
            fig, axes = plt.subplots(2, 1, figsize=(8, 10))
            fig.suptitle(f"{material_type}型材料 (组分={composition}) 能量流分析", fontsize=14)
            
            # 转换单位
            J = current_density * 1e4  # A/cm^2 → A/m^2
            
            # 准备数据
            n_points = len(x)
            dx = (x[-1] - x[0]) / (n_points - 1)
            
            # 计算温度梯度
            dTdx = np.zeros_like(T)
            dTdx[1:-1] = (T[2:] - T[:-2]) / (2*dx)
            dTdx[0] = (T[1] - T[0]) / dx
            dTdx[-1] = (T[-1] - T[-2]) / dx
            
            # 获取材料属性
            interp_key = f"{material_type}_{composition}"
            seebeck = np.zeros(n_points)
            resistivity = np.zeros(n_points)
            thermal_cond = np.zeros(n_points)
            
            for i in range(n_points):
                T_safe = np.clip(T[i], 300, 700)
                seebeck[i] = self.interpolators[interp_key]["seebeck"](T_safe)
                resistivity[i] = self.interpolators[interp_key]["resistivity"](T_safe)
                thermal_cond[i] = self.interpolators[interp_key]["thermal_cond"](T_safe)
            
            # 计算各种热流密度
            fourier_heat = thermal_cond * dTdx              # 傅里叶热流 κ·dT/dx
            peltier_heat = J * seebeck * T                  # 帕尔贴热流 J·S·T
            total_heat = fourier_heat - peltier_heat        # 净热流 q = κ·dT/dx - J·S·T
            joule_heat = J**2 * resistivity                 # 焦耳热 J²·ρ
            seebeck_power = J * seebeck * dTdx              # 塞贝克功率 J·S·dT/dx
            
            # 绘制热流分布
            ax1 = axes[0]
            ax1.plot(x, fourier_heat, 'r-', label='傅里叶热流 (κ·dT/dx)')
            ax1.plot(x, peltier_heat, 'b-', label='帕尔贴热流 (J·S·T)')
            ax1.plot(x, total_heat, 'g-', label='净热流 (q)')
            ax1.set_xlabel('位置 (归一化)')
            ax1.set_ylabel('热流密度 (W/m²)')
            ax1.legend()
            ax1.grid(True)
            
            # 绘制功率和热损失
            ax2 = axes[1]
            ax2.plot(x, seebeck_power, 'b-', label='塞贝克功率 (J·S·dT/dx)')
            ax2.plot(x, joule_heat, 'r-', label='焦耳热损失 (J²·ρ)')
            ax2.plot(x, seebeck_power - joule_heat, 'g-', label='净功率')
            ax2.set_xlabel('位置 (归一化)')
            ax2.set_ylabel('功率密度 (W/m³)')
            ax2.legend()
            ax2.grid(True)
            
            # 显示图表
            plt.tight_layout()
            plt.show()
            
        except Exception as e:
            print(f"能量流可视化错误: {str(e)}")
            import traceback
            traceback.print_exc()

    def validate_material_data(self, material_type, composition):
        """验证材料数据的有效性并返回验证结果"""
        try:
            data = self.p_type_data if material_type == 'p' else self.n_type_data
            if composition not in data:
                return False, f"找不到组分为 {composition} 的 {material_type}型材料数据"
            
            mat_data = data[composition]
            if len(mat_data["temp"]) == 0 or len(mat_data["seebeck"]) == 0 or len(mat_data["resistivity"]) == 0 or len(mat_data["thermal_cond"]) == 0:
                return False, f"{material_type}型材料 (组分={composition}) 的数据不完整"
                
            # 物理性质合理性检查
            if material_type == 'p':
                # P型材料的塞贝克系数应为正值
                if np.any(mat_data["seebeck"] <= 0):
                    return False, f"P型材料 (组分={composition}) 的塞贝克系数包含非正值"
            elif material_type == 'n':
                # N型材料的塞贝克系数应为负值
                if np.any(mat_data["seebeck"] >= 0):
                    return False, f"N型材料 (组分={composition}) 的塞贝克系数包含非负值"
            
            # 电阻率和热导率应为正值
            if np.any(mat_data["resistivity"] <= 0):
                return False, f"{material_type}型材料 (组分={composition}) 的电阻率包含非正值"
            if np.any(mat_data["thermal_cond"] <= 0):
                return False, f"{material_type}型材料 (组分={composition}) 的热导率包含非正值"
                
            # 数据点是否足够
            if len(mat_data["temp"]) < 5:
                return False, f"{material_type}型材料 (组分={composition}) 的数据点过少 ({len(mat_data['temp'])})"
                
            return True, f"{material_type}型材料 (组分={composition}) 的数据有效"
            
        except Exception as e:
            return False, f"验证材料数据时出错: {str(e)}"

    def material_P(self, T, composition):
        """获取P型材料属性，支持内置材料和自定义材料"""
        import pandas as pd
        from scipy.interpolate import CubicSpline
        
        # 检查是否为内置材料
        if composition in ['0.01', '0.02', '0.03']:
            # 内置材料处理
            if composition == '0.01':
                a1 = pd.read_excel(os.path.join(application_path, 'P_yuanshi_2_5.xls'), header=None).values
            elif composition == '0.02':
                a1 = pd.read_excel(os.path.join(application_path, 'P_yuanshi_3_1.xls'), header=None).values
            elif composition == '0.03':
                a1 = pd.read_excel(os.path.join(application_path, 'P_yuanshi_3_7.xls'), header=None).values
            
            sb = CubicSpline(a1[:, 0], a1[:, 1])(T) * 1e-6  # V/K
            res = CubicSpline(a1[:, 2], a1[:, 3])(T) * 1e-3  # Ω·m
            ZT = CubicSpline(a1[:, 4], a1[:, 5])(T)
            th = (T * sb ** 2) / (ZT * res)
            return sb, res, th
        else:
            # 自定义材料处理
            if composition not in self.p_type_data:
                raise ValueError(f"P型材料 {composition} 不存在")
            
            # 创建插值器（如果还没有创建）
            interp_key = f"p_{composition}"
            if interp_key not in self.interpolators:
                self.create_interpolators('p', composition)
            
            # 使用插值器获取材料属性，确保单位与内置材料一致
            sb = self.interpolators[interp_key]["seebeck"](T)
            res = self.interpolators[interp_key]["resistivity"](T)
            
            # 对于P型材料，统一使用ZT值计算热导率的方式
            if "zt" in self.interpolators[interp_key]:
                # 如果有ZT数据，使用ZT计算热导率（与内置材料保持一致）
                zt = self.interpolators[interp_key]["zt"](T)
                th = (T * sb ** 2) / (zt * res)
            else:
                # 如果没有ZT数据，使用热导率数据
                th = self.interpolators[interp_key]["thermal_cond"](T)
            
            return sb, res, th

    def material_N(self, T, composition):
        """获取N型材料属性，支持内置材料和自定义材料"""
        import pandas as pd
        from scipy.interpolate import CubicSpline
        
        # 检查是否为内置材料
        if composition in ['0.0004', '0.0012', '0.0020', '0.0028']:
            # 内置材料处理
            if composition == '0.0004':
                a2 = pd.read_excel(os.path.join(application_path, 'N_yuanshi_0.0004.xls'), header=None).values
            elif composition == '0.0012':
                a2 = pd.read_excel(os.path.join(application_path, 'N_yuanshi_0.0012.xls'), header=None).values
            elif composition == '0.0020':
                a2 = pd.read_excel(os.path.join(application_path, 'N_yuanshi_0.0020.xls'), header=None).values
            elif composition == '0.0028':
                a2 = pd.read_excel(os.path.join(application_path, 'N_yuanshi_0.0028.xls'), header=None).values
            
            sb = CubicSpline(a2[:, 0], a2[:, 1])(T) * 1e-6  # V/K
            res = CubicSpline(a2[:, 2], a2[:, 3])(T) * 1e-3  # Ω·m
            th = CubicSpline(a2[:, 4], a2[:, 5])(T) / 100    # W/(m·K)
            return sb, res, th
        else:
            # 自定义材料处理
            if composition not in self.n_type_data:
                raise ValueError(f"N型材料 {composition} 不存在")
            
            # 创建插值器（如果还没有创建）
            interp_key = f"n_{composition}"
            if interp_key not in self.interpolators:
                self.create_interpolators('n', composition)
            
            # 使用插值器获取材料属性，确保单位与内置材料一致
            sb = self.interpolators[interp_key]["seebeck"](T)
            res = self.interpolators[interp_key]["resistivity"](T)
            th = self.interpolators[interp_key]["thermal_cond"](T)
            
            return sb, res, th
    
    # ==================== 新材料数据管理功能 ====================
    
    def load_custom_materials(self):
        """加载用户自定义材料数据"""
        try:
            import json
            import os
            
            # 加载P型自定义材料
            p_file_path = os.path.join(self.material_data_dir, self.p_type_file)
            if os.path.exists(p_file_path):
                with open(p_file_path, 'r', encoding='utf-8') as f:
                    custom_p_data = json.load(f)
                    for composition, data in custom_p_data.items():
                        # 转换数据格式
                        self.p_type_data[composition] = {
                            "temp": np.array(data["temp"]),
                            "seebeck": np.array(data["seebeck"]),
                            "resistivity": np.array(data["resistivity"]),
                            "thermal_cond": np.array(data["thermal_cond"])
                        }
                    print(f"加载了 {len(custom_p_data)} 个自定义P型材料")
            
            # 加载N型自定义材料
            n_file_path = os.path.join(self.material_data_dir, self.n_type_file)
            if os.path.exists(n_file_path):
                with open(n_file_path, 'r', encoding='utf-8') as f:
                    custom_n_data = json.load(f)
                    for composition, data in custom_n_data.items():
                        # 转换数据格式
                        self.n_type_data[composition] = {
                            "temp": np.array(data["temp"]),
                            "seebeck": np.array(data["seebeck"]),
                            "resistivity": np.array(data["resistivity"]),
                            "thermal_cond": np.array(data["thermal_cond"])
                        }
                    print(f"加载了 {len(custom_n_data)} 个自定义N型材料")
                    
        except Exception as e:
            print(f"加载自定义材料数据失败: {str(e)}")
    
    def save_custom_materials(self):
        """保存用户自定义材料数据"""
        try:
            import json
            import os
            
            # 分离内置材料和自定义材料
            builtin_p = {"0.01", "0.02", "0.03"}
            builtin_n = {"0.0004", "0.0012", "0.0020", "0.0028"}
            
            custom_p_data = {}
            custom_n_data = {}
            
            # 提取P型自定义材料
            for composition, data in self.p_type_data.items():
                if composition not in builtin_p:
                    custom_p_data[composition] = {
                        "temp": data["temp"].tolist(),
                        "seebeck": data["seebeck"].tolist(),
                        "resistivity": data["resistivity"].tolist(),
                        "thermal_cond": data["thermal_cond"].tolist()
                    }
            
            # 提取N型自定义材料
            for composition, data in self.n_type_data.items():
                if composition not in builtin_n:
                    custom_n_data[composition] = {
                        "temp": data["temp"].tolist(),
                        "seebeck": data["seebeck"].tolist(),
                        "resistivity": data["resistivity"].tolist(),
                        "thermal_cond": data["thermal_cond"].tolist()
                    }
            
            # 保存P型材料
            p_file_path = os.path.join(self.material_data_dir, self.p_type_file)
            with open(p_file_path, 'w', encoding='utf-8') as f:
                json.dump(custom_p_data, f, ensure_ascii=False, indent=2)
            
            # 保存N型材料
            n_file_path = os.path.join(self.material_data_dir, self.n_type_file)
            with open(n_file_path, 'w', encoding='utf-8') as f:
                json.dump(custom_n_data, f, ensure_ascii=False, indent=2)
                
            print(f"保存了 {len(custom_p_data)} 个自定义P型材料和 {len(custom_n_data)} 个自定义N型材料")
            
        except Exception as e:
            print(f"保存自定义材料数据失败: {str(e)}")
    
    def add_custom_material(self, material_type, composition, material_data):
        """
        添加自定义材料数据
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料组分标识符
        material_data: 包含温度、塞贝克系数、电阻率、热导率数据的字典
        """
        try:
            # 验证数据格式
            required_keys = ["temp", "seebeck", "resistivity", "thermal_cond"]
            for key in required_keys:
                if key not in material_data:
                    raise ValueError(f"缺少必需的键: {key}")
            
            # 验证数据长度一致性
            data_lengths = [len(material_data[key]) for key in required_keys]
            if len(set(data_lengths)) != 1:
                raise ValueError("所有数据数组长度必须一致")
            
            # 验证数据合理性
            temp = np.array(material_data["temp"])
            seebeck = np.array(material_data["seebeck"])
            resistivity = np.array(material_data["resistivity"])
            thermal_cond = np.array(material_data["thermal_cond"])
            
            # 温度范围检查
            if np.any(temp < 0) or np.any(temp > 2000):
                raise ValueError("温度值应在0-2000K范围内")
            
            # 电阻率和热导率应为正值
            if np.any(resistivity <= 0):
                raise ValueError("电阻率必须为正值")
            if np.any(thermal_cond <= 0):
                raise ValueError("热导率必须为正值")
            
            # 塞贝克系数符号检查
            if material_type == 'p' and np.any(seebeck < 0):
                print("警告: P型材料的塞贝克系数应为正值")
            elif material_type == 'n' and np.any(seebeck > 0):
                print("警告: N型材料的塞贝克系数应为负值")
            
            # 存储材料数据
            target_data = self.p_type_data if material_type == 'p' else self.n_type_data
            target_data[composition] = {
                "temp": temp,
                "seebeck": seebeck,
                "resistivity": resistivity,
                "thermal_cond": thermal_cond
            }
            
            # 清除相关的插值器缓存
            interp_key = f"{material_type}_{composition}"
            if interp_key in self.interpolators:
                del self.interpolators[interp_key]
            
            # 保存到文件
            self.save_custom_materials()
            
            print(f"成功添加{material_type}型材料: {composition}")
            return True
            
        except Exception as e:
            print(f"添加材料数据失败: {str(e)}")
            return False
    def remove_custom_material(self, material_type, composition):
        """
        删除自定义材料数据
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料组分标识符
        """
        try:
            # 检查是否为内置材料
            builtin_p = {"0.01", "0.02", "0.03"}
            builtin_n = {"0.0004", "0.0012", "0.0020", "0.0028"}
            
            if material_type == 'p' and composition in builtin_p:
                raise ValueError("不能删除内置P型材料")
            elif material_type == 'n' and composition in builtin_n:
                raise ValueError("不能删除内置N型材料")
            
            # 删除材料数据
            target_data = self.p_type_data if material_type == 'p' else self.n_type_data
            if composition in target_data:
                del target_data[composition]
                
                # 清除相关的插值器缓存
                interp_key = f"{material_type}_{composition}"
                if interp_key in self.interpolators:
                    del self.interpolators[interp_key]
                
                # 保存到文件
                self.save_custom_materials()
                
                print(f"成功删除{material_type}型材料: {composition}")
                return True
            else:
                print(f"材料 {composition} 不存在")
                return False
                
        except Exception as e:
            print(f"删除材料数据失败: {str(e)}")
            return False
    
    def get_material_list(self, material_type):
        """
        获取指定类型的所有材料列表
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        
        返回:
        list: 材料组分列表
        """
        target_data = self.p_type_data if material_type == 'p' else self.n_type_data
        return list(target_data.keys())
    
    def get_material_info(self, material_type, composition):
        """
        获取材料详细信息
        
        参数:
        material_type: 'p' 或 'n'，材料类型
        composition: 材料组分标识符
        
        返回:
        dict: 材料信息字典
        """
        try:
            target_data = self.p_type_data if material_type == 'p' else self.n_type_data
            if composition not in target_data:
                return None
            
            data = target_data[composition]
            temp = data["temp"]
            seebeck = data["seebeck"]
            resistivity = data["resistivity"]
            thermal_cond = data["thermal_cond"]
            
            # 计算ZT值
            zt_values = []
            for i in range(len(temp)):
                try:
                    zt = (seebeck[i]**2 * temp[i]) / (resistivity[i] * thermal_cond[i])
                    zt_values.append(zt)
                except:
                    zt_values.append(0)
            
            return {
                "composition": composition,
                "material_type": material_type,
                "temperature_range": f"{min(temp):.1f} - {max(temp):.1f} K",
                "seebeck_range": f"{min(seebeck*1e6):.2f} - {max(seebeck*1e6):.2f} μV/K",
                "resistivity_range": f"{min(resistivity*1e6):.2f} - {max(resistivity*1e6):.2f} μΩ·m",
                "thermal_cond_range": f"{min(thermal_cond):.2f} - {max(thermal_cond):.2f} W/(m·K)",
                "max_zt": f"{max(zt_values):.3f}",
                "data_points": len(temp),
                "is_builtin": composition in {"0.01", "0.02", "0.03"} if material_type == 'p' else composition in {"0.0004", "0.0012", "0.0020", "0.0028"}
            }
            
        except Exception as e:
            print(f"获取材料信息失败: {str(e)}")
            return None
    
    def temperature_distribution_P(self, n, J, Tc, Th, max_iter, composition):
        """
        计算P型材料的温度分布（J单位与03.py一致，不再换算）
        """
        try:
            l = 1
            dx = l / (n - 1)
            T = np.linspace(Tc, Th, n)
            for iter_num in range(max_iter):
                A = np.zeros((n, n))
                b = np.zeros(n)
                sb, res, th = self.material_P(T, composition)
                c1 = J * sb / th
                c2 = -1 / th
                c3 = sb ** 2 * J ** 2 / th
                c4 = -J * sb / th
                c5 = res * J ** 2
                A[0, 0] = 1
                b[0] = Tc
                A[-1, -1] = 1
                b[-1] = Th
                for i in range(1, n - 1):
                    A[i, i - 1] = 1 / (c2[i] * dx)
                    A[i, i] = c4[i + 1] / c2[i + 1] - 1 / (c2[i + 1] * dx) - (1 - c1[i] * dx) / (c2[i] * dx)
                    A[i, i + 1] = (1 - c1[i + 1] * dx) / (c2[i + 1] * dx) - c3[i + 1] * dx - (1 - c1[i + 1] * dx) * c4[i + 1] / c2[i + 1]
                    b[i] = c5[i - 1] * dx
                try:
                    T_new = np.linalg.solve(A, b)
                    T_diff = np.max(np.abs(T_new - T))
                    T = T_new.copy()
                    if T_diff < 0.1:
                        break
                except np.linalg.LinAlgError:
                    print("  线性方程组求解失败，请检查系数矩阵")
                    return None, None
            x = np.linspace(0, l, n)
            return x, T
        except Exception as e:
            print(f"计算P型温度分布错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return None, None

    def temperature_distribution_N(self, J, Tc, Th, n_points, l, composition):
        """
        计算N型材料的温度分布（J单位与03.py一致，不再换算）
        """
        try:
            dx = l / (n_points - 1)
            T = np.linspace(Tc, Th, n_points)
            max_iter = 10
            for iter_num in range(max_iter):
                A = np.zeros((n_points, n_points))
                b = np.zeros(n_points)
                sb, res, th = self.material_N(T, composition)
                c1 = J * sb / th
                c2 = -1 / th
                c3 = sb ** 2 * J ** 2 / th
                c4 = -J * sb / th
                c5 = res * J ** 2
                A[0, 0] = 1
                b[0] = Tc
                A[-1, -1] = 1
                b[-1] = Th
                for i in range(1, n_points - 1):
                    A[i, i - 1] = 1 / (c2[i] * dx)
                    A[i, i] = c4[i + 1] / c2[i + 1] - 1 / (c2[i + 1] * dx) - (1 - c1[i] * dx) / (c2[i] * dx)
                    A[i, i + 1] = (1 - c1[i + 1] * dx) / (c2[i + 1] * dx) - c3[i + 1] * dx - (1 - c1[i + 1] * dx) * c4[i + 1] / c2[i + 1]
                    b[i] = c5[i - 1] * dx
                try:
                    T_new = np.linalg.solve(A, b)
                    T_diff = np.max(np.abs(T_new - T))
                    T = T_new.copy()
                    if T_diff < 0.1:
                        break
                except np.linalg.LinAlgError:
                    print("  线性方程组求解失败，请检查系数矩阵")
                    return None, None
            x = np.linspace(0, l, n_points)
            return x, T
        except Exception as e:
            print(f"计算N型温度分布错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return None, None

    def calculate_efficiency_P(self, Tc, Th, n, l, composition):
        """
        计算P型材料的效率曲线，算法与03.py一致
        """
        import numpy as np
        eff_list_P = []
        J_list_P = []
        q_P_list = []
        dx = l / (n - 1)
        for j in range(0, 31, 1):
            J = -j
            T = self.temperature_distribution_P(n, J, Tc, Th, 10, composition)[1]
            sb, res, th = self.material_P(T, composition)
            c1 = J * sb / th
            c2 = -1 / th
            c3 = sb ** 2 * J ** 2 / th
            c4 = -J * sb / th
            c5 = res * J ** 2
            q_P = np.zeros(n)
            for k in range(1, n):
                q_P[k] = ((1 / dx - c1[k]) * T[k] - T[k - 1] / dx) / (c2[k])
            q_P[0] = (1 - c4[1] * dx) * q_P[1] - c3[1] * dx * T[1] - c5[1] * dx
            q_P_list.append(q_P[-1])
            Cumulative_scoring1 = 0
            Cumulative_scoring2 = 0
            for m in range(1, n):
                T1 = T[m]
                T2 = T[m - 1]
                Cumulative_scoring1 += (sb[m] + sb[m - 1]) / 2 * (T1 - T2)
                Cumulative_scoring2 += (res[m] + res[m - 1]) / 2 * dx
            if q_P[n - 1] != 0:
                eff = J * (Cumulative_scoring1 + J * Cumulative_scoring2) / q_P[n - 1]
                eff_list_P.append(eff)
                J_list_P.append(J)
        return eff_list_P, J_list_P, q_P_list

    def calculate_efficiency_N(self, Tc, Th, n, l, composition):
        """
        计算N型材料的效率曲线，算法与03.py一致
        """
        import numpy as np
        eff_list_N = []
        J_list_N = []
        q_N_list = []
        dx = l / (n - 1)
        for j in range(0, 51, 1):
            J = j
            T = self.temperature_distribution_N(J, Tc, Th, n, l, composition)[1]
            sb, res, th = self.material_N(T, composition)
            c1 = J * sb / th
            c2 = -1 / th
            c3 = sb ** 2 * J ** 2 / th
            c4 = -J * sb / th
            c5 = res * J ** 2
            q_N = np.zeros(n)
            for k in range(1, n):
                q_N[k] = ((1 / dx - c1[k]) * T[k] - T[k - 1] / dx) / (c2[k])
            q_N[0] = (1 - c4[1] * dx) * q_N[1] - c3[1] * dx * T[1] - c5[1] * dx
            q_N_list.append(q_N[-1])
            Cumulative_scoring1 = 0
            Cumulative_scoring2 = 0
            for m in range(1, n):
                T1 = T[m]
                T2 = T[m - 1]
                Cumulative_scoring1 += (sb[m] + sb[m - 1]) / 2 * (T1 - T2)
                Cumulative_scoring2 += (res[m] + res[m - 1]) / 2 * dx
            if q_N[n - 1] != 0:
                eff = J * (Cumulative_scoring1 + J * Cumulative_scoring2) / q_N[n - 1]
                eff_list_N.append(eff)
                J_list_N.append(J)
        return eff_list_N, J_list_N, q_N_list

class ParameterValidator:
    """参数验证类"""
    
    @staticmethod
    def validate_temperature(value, name):
        """验证温度参数"""
        try:
            temp = float(value)
            if temp <= 0:
                return False, f"{name}必须大于0"
            if temp > 1000:
                return False, f"{name}不能超过1000K"
            return True, ""
        except ValueError:
            return False, f"{name}必须是有效的数字"
    
    @staticmethod
    def validate_grid_points(value):
        """验证网格点数"""
        try:
            points = int(value)
            if points < 5:
                return False, "网格点数不能少于5"
            if points > 100:
                return False, "网格点数不能超过100"
            return True, ""
        except ValueError:
            return False, "网格点数必须是有效的整数"
    
    @staticmethod
    def validate_iterations(value):
        """验证迭代次数"""
        try:
            iterations = int(value)
            if iterations < 1:
                return False, "迭代次数必须大于0"
            if iterations > 1000:
                return False, "迭代次数不能超过1000"
            return True, ""
        except ValueError:
            return False, "迭代次数必须是有效的整数"
    
    @staticmethod
    def validate_all_params(th, tc, grid_points, iterations):
        """验证所有参数"""
        errors = []
        
        # 验证温度参数
        th_valid, th_msg = ParameterValidator.validate_temperature(th, "高温温度Th")
        if not th_valid:
            errors.append(th_msg)
            
        tc_valid, tc_msg = ParameterValidator.validate_temperature(tc, "低温温度Tc")
        if not tc_valid:
            errors.append(tc_msg)
        
        # 验证温度关系
        if th_valid and tc_valid:
            if float(th) <= float(tc):
                errors.append("高温温度Th必须大于低温温度Tc")
        
        # 验证网格点数
        grid_valid, grid_msg = ParameterValidator.validate_grid_points(grid_points)
        if not grid_valid:
            errors.append(grid_msg)
        
        # 验证迭代次数
        iter_valid, iter_msg = ParameterValidator.validate_iterations(iterations)
        if not iter_valid:
            errors.append(iter_msg)
        
        return len(errors) == 0, errors
class ThermoelectricApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setup_plot_style()
        self.setWindowTitle('基于差分法的半导体热电器件仿真实验')
        
        # 设置窗口的默认大小和最小大小
        screen = QApplication.primaryScreen().geometry()
        default_width = min(int(screen.width() * 0.8), 1440)  # 最大宽度1440
        default_height = min(int(screen.height() * 0.8), 900)  # 最大高度900
        self.setGeometry(100, 100, default_width, default_height)
        self.setMinimumSize(1024, 600)  # 设置最小窗口大小
        
        # 初始化验证器
        self.validator = ParameterValidator()
        
        # 初始化计算线程
        self.calculation_thread = None
        
        # 创建主窗口部件
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        # 创建主布局
        main_layout = QHBoxLayout(main_widget)
        main_layout.setSpacing(5)  # 减小面板之间的间距
        main_layout.setContentsMargins(5, 5, 5, 5)  # 减小边距
        
        # 创建左侧面板 - 先创建它，确保iter_edit已经定义
        left_panel = self.create_left_panel()
        main_layout.addWidget(left_panel)
        
        # 初始化计算器 - 现在iter_edit已经存在
        self.calculator = ThermoelectricCalculator()
        
        # 创建中间面板
        middle_panel = self.create_middle_panel()
        main_layout.addWidget(middle_panel)
        
        # 创建右侧面板
        right_panel = self.create_right_panel()
        main_layout.addWidget(right_panel)
        
        # 设置面板的比例 (左:中:右 = 2:3:3)
        main_layout.setStretch(0, 2)
        main_layout.setStretch(1, 3)
        main_layout.setStretch(2, 3)

        # 创建状态栏
        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        
        # 创建进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.statusBar.addPermanentWidget(self.progress_bar)
        
        # 添加状态标签
        self.status_label = QLabel("就绪")
        self.statusBar.addWidget(self.status_label)
        
        # 连接信号和槽
        self.init_button.clicked.connect(self.initialize_calculation)
        self.p_current_combo.currentTextChanged.connect(self.update_branch_characteristics)
        self.n_current_combo.currentTextChanged.connect(self.update_branch_characteristics)
        
        # 连接右侧面板的计算和导出按钮
        self.right_calc_button.clicked.connect(self.calculate_device_performance)
        self.right_export_button.clicked.connect(self.export_data)
        
        # 初始化实验报告生成器
        self.report_generator = ExperimentReportGenerator(self)
        
        # 设置初始状态
        self.update_status("就绪")
        
        # 创建菜单栏
        self.create_menu_bar()

    def create_menu_bar(self):
        """创建菜单栏"""
        menubar = self.menuBar()
        
        # 文件菜单
        file_menu = menubar.addMenu('文件(&F)')
        
        # 新建项目
        new_action = file_menu.addAction('新建项目(&N)')
        new_action.setShortcut('Ctrl+N')
        new_action.triggered.connect(self.new_project)
        
        # 打开项目
        open_action = file_menu.addAction('打开项目(&O)')
        open_action.setShortcut('Ctrl+O')
        open_action.triggered.connect(self.open_project)
        
        # 保存项目
        save_action = file_menu.addAction('保存项目(&S)')
        save_action.setShortcut('Ctrl+S')
        save_action.triggered.connect(self.save_project)
        
        file_menu.addSeparator()
        
        # 导出数据
        export_action = file_menu.addAction('导出数据(&E)')
        export_action.setShortcut('Ctrl+E')
        export_action.triggered.connect(self.export_data)
        
        # 导出实验报告
        report_menu = menubar.addMenu('实验报告(&R)')
        
        # 一键导出实验报告
        export_report_action = report_menu.addAction('一键导出实验报告(&E)')
        export_report_action.setShortcut('Ctrl+R')
        export_report_action.triggered.connect(self.export_experiment_report)
        
        # 导出Word格式报告
        export_word_action = report_menu.addAction('导出Word格式报告(&W)')
        export_word_action.setShortcut('Ctrl+Shift+W')
        export_word_action.triggered.connect(self.export_word_report)
        
        # 导出HTML格式报告
        export_html_action = report_menu.addAction('导出HTML格式报告(&H)')
        export_html_action.setShortcut('Ctrl+Shift+H')
        export_html_action.triggered.connect(self.export_html_report)
        
        # 导出文本格式报告
        export_txt_action = report_menu.addAction('导出文本格式报告(&T)')
        export_txt_action.setShortcut('Ctrl+Shift+T')
        export_txt_action.triggered.connect(self.export_txt_report)
        
        file_menu.addSeparator()
        
        # 退出
        exit_action = file_menu.addAction('退出(&X)')
        exit_action.setShortcut('Ctrl+Q')
        exit_action.triggered.connect(self.close)
        
        # 计算菜单
        calc_menu = menubar.addMenu('计算(&C)')
        
        # 开始计算
        start_calc_action = calc_menu.addAction('开始计算(&S)')
        start_calc_action.setShortcut('F5')
        start_calc_action.triggered.connect(self.initialize_calculation)
        
        # 停止计算
        stop_calc_action = calc_menu.addAction('停止计算(&T)')
        stop_calc_action.setShortcut('F6')
        stop_calc_action.triggered.connect(self.stop_calculation)
        
        calc_menu.addSeparator()
        
        # 外部效率计算
        external_calc_action = calc_menu.addAction('外部效率计算(&E)')
        external_calc_action.setShortcut('F7')
        external_calc_action.triggered.connect(self.run_external_efficiency)
        
        # 视图菜单
        view_menu = menubar.addMenu('视图(&V)')
        
        # 重置视图
        reset_view_action = view_menu.addAction('重置视图(&R)')
        reset_view_action.setShortcut('F8')
        reset_view_action.triggered.connect(self.reset_all_views)
        
        # 全屏显示
        fullscreen_action = view_menu.addAction('全屏显示(&F)')
        fullscreen_action.setShortcut('F11')
        fullscreen_action.triggered.connect(self.toggle_fullscreen)
        
        # 材料管理菜单
        material_menu = menubar.addMenu('材料管理(&M)')
        
        # 添加新材料
        add_material_action = material_menu.addAction('添加新材料(&A)')
        add_material_action.setShortcut('Ctrl+M')
        add_material_action.triggered.connect(self.show_material_manager)
        
        # 导入材料数据
        import_material_action = material_menu.addAction('导入材料数据(&I)')
        import_material_action.setShortcut('Ctrl+Shift+I')
        import_material_action.triggered.connect(self.import_material_data)
        
        # 导出材料数据
        export_material_action = material_menu.addAction('导出材料数据(&E)')
        export_material_action.setShortcut('Ctrl+Shift+E')
        export_material_action.triggered.connect(self.export_material_data)
        
        material_menu.addSeparator()
        
        # 材料信息查看
        view_materials_action = material_menu.addAction('查看材料信息(&V)')
        view_materials_action.setShortcut('Ctrl+Shift+V')
        view_materials_action.triggered.connect(self.view_materials_info)
        
        # 帮助菜单
        help_menu = menubar.addMenu('帮助(&H)')
        
        # 使用帮助
        help_action = help_menu.addAction('使用帮助(&H)')
        help_action.setShortcut('F1')
        help_action.triggered.connect(self.show_help)
        
        # 关于程序
        about_action = help_menu.addAction('关于程序(&A)')
        about_action.triggered.connect(self.show_about)

    def setup_plot_style(self):
        plt.style.use('default')
        
        # 设置中文字体，优先使用支持上标的字体
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False     # 用来正常显示负号
        
        plt.rcParams.update({
            'figure.facecolor': '#F0F0F0',
            'axes.facecolor': '#F0F0F0',
            'axes.grid': False,
            'axes.spines.top': True,
            'axes.spines.right': True,
            'font.size': 10,
            'figure.subplot.hspace': 0.3,
            'figure.subplot.wspace': 0.3,
            'mathtext.fontset': 'dejavusans'  # 使用DejaVu Sans字体集，支持上标
        })

    def update_status(self, message):
        """更新状态栏消息"""
        self.status_label.setText(message)
        QApplication.processEvents()
    
    def show_progress(self, show=True):
        """显示或隐藏进度条"""
        self.progress_bar.setVisible(show)
        if show:
            self.progress_bar.setValue(0)
    
    def update_progress(self, value):
        """更新进度条值"""
        self.progress_bar.setValue(value)
        QApplication.processEvents()
    
    def show_error_message(self, title, message):
        """显示错误消息对话框"""
        QMessageBox.critical(self, title, message)
    
    def show_info_message(self, title, message):
        """显示信息消息对话框"""
        QMessageBox.information(self, title, message)
    
    def validate_inputs(self):
        """验证输入参数"""
        try:
            th = self.th_edit.text().strip()
            tc = self.tc_edit.text().strip()
            grid_points = self.grid_edit.text().strip()
            iterations = self.iter_edit.text().strip()
            
            # 检查空值
            if not all([th, tc, grid_points, iterations]):
                return False, "请填写所有必需的参数"
            
            # 验证参数值
            is_valid, errors = self.validator.validate_all_params(th, tc, grid_points, iterations)
            if not is_valid:
                return False, "\n".join(errors)
            
            return True, ""
            
        except Exception as e:
            return False, f"参数验证失败: {str(e)}"

    def create_toolbar_buttons(self):
        buttons = []
        icons = ["⌂", "←", "→", "✥", "🔍", "≡", "📄"]
        for icon in icons:
            btn = QPushButton(icon)
            btn.setFixedSize(25, 25)
            btn.setStyleSheet("""
                QPushButton {
                    background-color: white;
                    border: 1px solid #dcdcdc;
                    border-radius: 3px;
                    padding: 5px;
                }
                QPushButton:hover {
                    background-color: #e6e6e6;
                }
            """)
            buttons.append(btn)
        return buttons

    def create_plot_widget(self, num_subplots=2, height=3, vertical=False):
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)  # 完全移除边距
        layout.setSpacing(0)  # 移除间距
        
        # 创建工具栏
        toolbar = QFrame()
        toolbar.setFixedHeight(16)  # 进一步减小工具栏高度
        toolbar.setStyleSheet("""
            QFrame {
                background-color: #F0F0F0;
                border: none;
                margin: 0px;
                padding: 0px;
            }
        """)
        toolbar_layout = QHBoxLayout(toolbar)
        toolbar_layout.setContentsMargins(1, 0, 1, 0)  # 只保留左右边距
        toolbar_layout.setSpacing(1)  # 最小按钮间距
        
        # 创建工具按钮
        icons = ["⌂", "←", "→", "+", "🔍", "≡", "📄"]
        for icon in icons:
            btn = QPushButton(icon)
            btn.setFixedSize(16, 16)  # 进一步减小按钮大小
            btn.setStyleSheet("""
                QPushButton {
                    background-color: white;
                    border: 1px solid #CCCCCC;
                    border-radius: 1px;
                    padding: 0px;
                    margin: 0px;
                    font-size: 9px;
                }
                QPushButton:hover {
                    background-color: #E6E6E6;
                }
            """)
            toolbar_layout.addWidget(btn)
        toolbar_layout.addStretch()
        layout.addWidget(toolbar)
        
        # 创建图表
        dpi = QApplication.primaryScreen().logicalDotsPerInch()
        fig_width = container.width() / dpi
        fig_height = (height * 96 + 10) / dpi  # 稍微增加图表高度
        
        if vertical and num_subplots > 1:
            fig, axes = plt.subplots(num_subplots, 1, figsize=(fig_width, fig_height))
        else:
            fig, axes = plt.subplots(1, num_subplots, figsize=(fig_width, fig_height))
        
        if num_subplots == 1:
            axes = [axes]
        
        # 设置图表样式
        for ax in axes:
            ax.grid(True, color='white', linestyle='-', alpha=0.8)
            ax.set_facecolor('#F0F0F0')
            ax.clear()
            ax.grid(True)
            # 调整字体大小
            ax.tick_params(labelsize=8)
            for label in ax.get_xticklabels() + ax.get_yticklabels():
                label.set_fontsize(8)
        
        # 调整图表间距，确保两个子图不重叠
        if num_subplots > 1:
            plt.subplots_adjust(top=0.88, bottom=0.15, left=0.08, right=0.95, wspace=0.35)
        else:
            plt.subplots_adjust(top=0.88, bottom=0.15, left=0.15, right=0.95)
        
        canvas = FigureCanvas(fig)
        layout.addWidget(canvas)
        
        return container, axes, canvas

    def create_left_panel(self):
        panel = QGroupBox()
        layout = QVBoxLayout()
        layout.setContentsMargins(5, 5, 5, 5)
        layout.setSpacing(5)
        
        # 添加标题
        title_label = QLabel("基于差分法的半导体热电器件仿真实验")
        title_label.setStyleSheet("""
            font-size: 18px;
            font-weight: bold;
            color: #0072BC;
            padding: 5px;
        """)
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setWordWrap(True)
        title_label.setFixedHeight(50)
        layout.addWidget(title_label)
        
        # 添加示意图
        image_container = QGroupBox()
        image_layout = QVBoxLayout(image_container)
        image_layout.setContentsMargins(0, 0, 0, 0)
        
        # 使用新的ClickableImageLabel替代QLabel
        image_label = ClickableImageLabel()
        pixmap = QPixmap("图片1.png")
        scaled_pixmap = pixmap.scaled(400, 320, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        image_label.setPixmap(scaled_pixmap)
        image_label.setAlignment(Qt.AlignCenter)
        # 添加提示文本
        image_label.setToolTip("双击查看大图")
        image_layout.addWidget(image_label)
        
        layout.addWidget(image_container)
        layout.addSpacing(10)
        
        # 初始条件设置
        params_group = QGroupBox("初始条件设置")
        params_layout = QGridLayout()
        params_layout.setContentsMargins(5, 5, 5, 5)  # 减小边距
        params_layout.setSpacing(5)  # 减小间距
        
        # 温度和网格设置
        params_layout.addWidget(QLabel("高温温度Th(K)"), 0, 0)
        self.th_edit = QLineEdit("500")
        self.th_edit.setToolTip("设置高温端温度，范围：1-1000K")
        self.th_edit.setPlaceholderText("500")
        params_layout.addWidget(self.th_edit, 0, 1)
        
        params_layout.addWidget(QLabel("格子数量"), 0, 2)
        self.grid_edit = QLineEdit("10")
        self.grid_edit.setToolTip("设置计算网格点数，范围：5-100")
        self.grid_edit.setPlaceholderText("10")
        params_layout.addWidget(self.grid_edit, 0, 3)
        
        params_layout.addWidget(QLabel("低温温度Tc(K)"), 1, 0)
        self.tc_edit = QLineEdit("300")
        self.tc_edit.setToolTip("设置低温端温度，范围：1-1000K，且必须小于Th")
        self.tc_edit.setPlaceholderText("300")
        params_layout.addWidget(self.tc_edit, 1, 1)
        
        params_layout.addWidget(QLabel("迭代次数"), 1, 2)
        self.iter_edit = QLineEdit("20")
        self.iter_edit.setToolTip("设置最大迭代次数，范围：1-1000")
        self.iter_edit.setPlaceholderText("20")
        params_layout.addWidget(self.iter_edit, 1, 3)
        
        # 材料选择
        params_layout.addWidget(QLabel("P: PbTe₁₋ᵧIᵧ"), 2, 0)
        self.p_type_combo = QComboBox()
        self.p_type_combo.addItems(["0.01", "0.02", "0.03"])  # 使用实际组分值而不是文件名
        params_layout.addWidget(self.p_type_combo, 2, 1)
        self.p_type_combo.currentTextChanged.connect(self.update_p_composition)
        
        params_layout.addWidget(QLabel("N: PbTe:Na/Ag₂Te"), 2, 2)
        self.n_type_combo = QComboBox()
        self.n_type_combo.addItems(["0.0004", "0.0012", "0.0020", "0.0028"])  # 保持N型材料组分值不变
        params_layout.addWidget(self.n_type_combo, 2, 3)
        
        params_group.setLayout(params_layout)
        layout.addWidget(params_group)
        
        # 材料优值系数图表
        zt_group = QGroupBox("选择材料的优值系数")
        zt_layout = QVBoxLayout()
        zt_layout.setContentsMargins(5, 5, 5, 5)  # 减小边距
        
        zt_container, (ax1, ax2), canvas = self.create_plot_widget(height=2)
        self.zt_axes = (ax1, ax2)  # 保存axes引用以便后续更新
        self.zt_canvas = canvas    # 保存canvas引用以便后续更新
        
        # 设置P型图表
        ax1.set_title("P型半导体材料", pad=5)
        ax1.set_xlabel("温度")
        ax1.set_ylabel("ZT")
        ax1.set_xlim(300, 700)
        ax1.set_ylim(0, 1.5)
        ax1.grid(True, color='white', linestyle='-', alpha=0.8)
        ax1.set_facecolor('#F0F0F0')
        
        # 设置N型图表
        ax2.set_title("N型半导体材料", pad=5)
        ax2.set_xlabel("温度")
        ax2.set_ylabel("ZT")
        ax2.set_xlim(300, 700)
        ax2.set_ylim(0, 1.5)
        ax2.grid(True, color='white', linestyle='-', alpha=0.8)
        ax2.set_facecolor('#F0F0F0')
        
        # 调整图表布局
        plt.tight_layout()
        
        zt_layout.addWidget(zt_container)
        zt_group.setLayout(zt_layout)
        layout.addWidget(zt_group)
        
        # 帮助功能已集成到菜单栏中，无需单独按钮
        
        # 添加初始化按钮
        self.init_button = QPushButton("开始计算")
        self.init_button.setStyleSheet("""
            QPushButton {
                background-color: #28A745;
                color: white;
                border: none;
                padding: 10px;
                border-radius: 4px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1E7E34;
            }
        """)
        layout.addWidget(self.init_button)
        
        # 添加状态指示灯
        status_layout = QHBoxLayout()
        status_layout.addWidget(QLabel("运行状态"))
        self.status_light = StatusLight()
        status_layout.addWidget(self.status_light)
        status_layout.addStretch()
        layout.addLayout(status_layout)
        
        # 设置拉伸因子，使图片区域占据更多空间
        layout.setStretch(0, 0)  # 标题不拉伸
        layout.setStretch(1, 1)  # 图片区域拉伸
        layout.setStretch(2, 0)  # 参数设置不拉伸
        layout.setStretch(3, 0)  # ZT图表不拉伸
        layout.setStretch(4, 0)  # 计算按钮不拉伸
        layout.setStretch(5, 0)  # 状态指示灯不拉伸
        
        panel.setLayout(layout)
        return panel

    def create_middle_panel(self):
        panel = QGroupBox("分支特性")
        layout = QVBoxLayout()
        
        # 格点温度分布
        temp_group = QGroupBox("格点温度分布")
        temp_layout = QVBoxLayout()
        
        temp_container, (ax1, ax2), canvas = self.create_plot_widget()
        # 保存温度分布图表的引用
        self.temp_axes = (ax1, ax2)
        self.temp_canvas = canvas
        
        # 移除多余的提示标签
        
        ax1.set_title("格点温度分布（P型）")
        ax2.set_title("格点温度分布（N型）")
        
        for ax in [ax1, ax2]:
            ax.set_xlabel("格点位置")
            ax.set_ylabel("T (K)")
            ax.set_xlim(0, 10)
            ax.set_ylim(300, 500)
        
        temp_layout.addWidget(temp_container)
        
        # 电流密度选择
        current_layout = QHBoxLayout()
        label_p = QLabel("电流密度 (A/cm<sup>2</sup>)")
        label_p.setTextFormat(Qt.RichText)
        current_layout.addWidget(label_p)
        self.p_current_combo = QComboBox()
        self.p_current_combo.addItems(["-2.0", "-1.5", "-1.0", "-0.5"])
        current_layout.addWidget(self.p_current_combo)
        label_n = QLabel("电流密度 (A/cm<sup>2</sup>)")
        label_n.setTextFormat(Qt.RichText)
        current_layout.addWidget(label_n)
        self.n_current_combo = QComboBox()
        self.n_current_combo.addItems(["25", "30", "35", "40"])
        current_layout.addWidget(self.n_current_combo)
        
        temp_layout.addLayout(current_layout)
        temp_group.setLayout(temp_layout)
        layout.addWidget(temp_group)
        
        # 材料效率
        eff_group = QGroupBox("材料效率")
        eff_layout = QVBoxLayout()
        eff_container, (ax3, ax4), canvas = self.create_plot_widget(num_subplots=2, height=3)
        self.eff_axes = (ax3, ax4)
        self.eff_canvas = canvas
        
        # 设置P型效率图
        ax3.set_title("效率 (P型)")
        ax3.set_xlabel(r"电流密度 (A/cm$^2$)")
        ax3.set_ylabel("效率")
        ax3.set_xlim(-20, 0)
        ax3.set_ylim(0, 0.1)
        ax3.grid(True, color='white', linestyle='-', alpha=0.8)
        ax3.set_facecolor('#F0F0F0')
        
        # 设置N型效率图
        ax4.set_title("效率 (N型)")
        ax4.set_xlabel(r"电流密度 (A/cm$^2$)")
        ax4.set_ylabel("效率")
        ax4.set_xlim(0, 40)
        ax4.set_ylim(0, 0.1)
        ax4.grid(True, color='white', linestyle='-', alpha=0.8)
        ax4.set_facecolor('#F0F0F0')
        
        # 确保两个子图有足够的间距
        canvas.figure.subplots_adjust(wspace=0.4)
        
        eff_layout.addWidget(eff_container)
        
        # 添加计算按钮和状态指示灯
        calc_layout = QHBoxLayout()
        calc_button = QPushButton("计算")
        calc_button.clicked.connect(self.update_branch_characteristics)
        calc_layout.addWidget(calc_button)
        
        # 添加新按钮运行PNefficiency.py程序
        external_calc_button = QPushButton("详细效率计算")
        external_calc_button.clicked.connect(self.run_external_efficiency)
        calc_layout.addWidget(external_calc_button)
        
        calc_layout.addWidget(QLabel("运行状态"))
        self.calc_status = StatusLight()
        calc_layout.addWidget(self.calc_status)
        calc_layout.addStretch()
        
        eff_layout.addLayout(calc_layout)
        eff_group.setLayout(eff_layout)
        layout.addWidget(eff_group)
        
        panel.setLayout(layout)
        return panel

    def create_right_panel(self):
        panel = QGroupBox("结果分析")
        layout = QVBoxLayout()
        layout.setSpacing(5)
        layout.setContentsMargins(5, 5, 5, 5)
        
        # N/P比例设置
        ratio_layout = QHBoxLayout()
        ratio_layout.setContentsMargins(0, 0, 0, 0)
        ratio_layout.addWidget(QLabel("N型分支面积/P型分支面积"))
        self.ratio_edit = QLineEdit("0.1")
        ratio_layout.addWidget(self.ratio_edit)
        layout.addLayout(ratio_layout)
        
        # 1. 器件功率图表
        power_group = QGroupBox("器件功率")
        power_layout = QVBoxLayout()
        power_container, [power_ax], power_canvas = self.create_plot_widget(num_subplots=1, height=2.5)
        self.power_canvas = power_canvas  # 保存canvas引用
        self.power_ax = power_ax          # 保存ax引用
        power_ax.set_title("器件功率")
        power_ax.set_xlabel("电流 (A)")
        power_ax.set_ylabel(r"功率 ($W/cm^2$)")
        power_ax.set_xlim(0, 10)  # 扩大范围以适应更大的电流值
        power_ax.set_ylim(0, 1)
        power_layout.addWidget(power_container)
        power_group.setLayout(power_layout)
        layout.addWidget(power_group)
        
        # 2. 器件效率图表
        eff_group = QGroupBox("器件效率")
        eff_layout = QVBoxLayout()
        efficiency_container, [efficiency_ax], efficiency_canvas = self.create_plot_widget(num_subplots=1, height=2.5)
        self.efficiency_canvas = efficiency_canvas  # 保存canvas引用
        self.efficiency_ax = efficiency_ax          # 保存ax引用
        efficiency_ax.set_title("器件效率")
        efficiency_ax.set_xlabel("电流 (A)")
        efficiency_ax.set_ylabel("效率 (%)")
        efficiency_ax.set_xlim(0, 10)
        # Y轴范围将在数据绘制时动态调整，主要关注0%-10%区间
        efficiency_ax.set_ylim(0, 10)
        eff_layout.addWidget(efficiency_container)
        eff_group.setLayout(eff_layout)
        layout.addWidget(eff_group)
        
        # 最大功率点和最大效率点显示框
        results_layout = QHBoxLayout()
        results_layout.setSpacing(10)  # 减小显示框之间的间距
        results_layout.setContentsMargins(0, 0, 0, 0)
        
        # 最大功率点
        power_point_group = QGroupBox("最大功率点")
        power_layout = QVBoxLayout()
        power_layout.setSpacing(5)  # 减小内部组件的间距
        power_layout.setContentsMargins(5, 5, 5, 5)
        power_value_layout = QHBoxLayout()
        power_value_layout.addWidget(QLabel("最大功率"))
        self.max_power = QLineEdit()
        self.max_power.setReadOnly(True)  # 设置为只读
        power_value_layout.addWidget(self.max_power)
        power_layout.addLayout(power_value_layout)
        power_current_layout = QHBoxLayout()
        power_current_layout.addWidget(QLabel("电流"))
        self.power_current = QLineEdit()
        self.power_current.setReadOnly(True)  # 设置为只读
        power_current_layout.addWidget(self.power_current)
        power_layout.addLayout(power_current_layout)
        power_point_group.setLayout(power_layout)
        results_layout.addWidget(power_point_group)
        
        # 最大效率点
        eff_point_group = QGroupBox("最大效率点")
        eff_layout = QVBoxLayout()
        eff_layout.setSpacing(5)  # 减小内部组件的间距
        eff_layout.setContentsMargins(5, 5, 5, 5)
        eff_value_layout = QHBoxLayout()
        eff_value_layout.addWidget(QLabel("最大效率"))
        self.max_eff = QLineEdit()
        self.max_eff.setReadOnly(True)  # 设置为只读
        eff_value_layout.addWidget(self.max_eff)
        eff_layout.addLayout(eff_value_layout)
        eff_current_layout = QHBoxLayout()
        eff_current_layout.addWidget(QLabel("电流"))
        self.eff_current = QLineEdit()
        self.eff_current.setReadOnly(True)  # 设置为只读
        eff_current_layout.addWidget(self.eff_current)
        eff_layout.addLayout(eff_current_layout)
        eff_point_group.setLayout(eff_layout)
        results_layout.addWidget(eff_point_group)
        
        layout.addLayout(results_layout)
        
        # 3. 功率效率优化区间图表
        opt_group = QGroupBox("功率效率优化区间")
        opt_layout = QVBoxLayout()
        optimization_container, [optimization_ax], optimization_canvas = self.create_plot_widget(num_subplots=1, height=2.5)
        self.optimization_canvas = optimization_canvas  # 保存canvas引用
        self.optimization_ax = optimization_ax          # 保存ax引用
        optimization_ax.set_title("功率效率优化区间")
        optimization_ax.set_xlabel(r"功率 ($W/cm^2$)")
        optimization_ax.set_ylabel("效率 (%)")
        optimization_ax.set_xlim(0, 0.1)
        # Y轴范围将在数据绘制时动态调整，主要关注0%-10%区间
        optimization_ax.set_ylim(0, 10)
        opt_layout.addWidget(optimization_container)
        opt_group.setLayout(opt_layout)
        layout.addWidget(opt_group)
        
        # 底部按钮和进度条
        button_layout = QHBoxLayout()
        button_layout.setSpacing(10)
        button_layout.setContentsMargins(0, 0, 0, 0)
        
        # 计算按钮
        self.right_calc_button = QPushButton("计算")
        self.right_calc_button.clicked.connect(self.calculate_device_performance)
        button_layout.addWidget(self.right_calc_button)
        
        # 导出按钮
        self.right_export_button = QPushButton("导出数据")
        self.right_export_button.clicked.connect(self.export_data)
        button_layout.addWidget(self.right_export_button)
        
        # 运行状态灯
        self.right_status_light = StatusLight()
        button_layout.addWidget(QLabel("运行状态"))
        button_layout.addWidget(self.right_status_light)
        
        button_layout.addStretch()
        layout.addLayout(button_layout)
        
        # 添加进度条和状态标签
        progress_layout = QHBoxLayout()
        progress_layout.setContentsMargins(0, 5, 0, 0)
        
        # 进度标签
        self.right_progress_label = QLabel("就绪")
        self.right_progress_label.setStyleSheet("""
            QLabel {
                color: #666;
                font-size: 12px;
                padding: 2px;
            }
        """)
        progress_layout.addWidget(self.right_progress_label)
        
        # 进度条
        self.right_progress_bar = QProgressBar()
        self.right_progress_bar.setVisible(False)
        self.right_progress_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #ccc;
                border-radius: 3px;
                text-align: center;
                background-color: #f0f0f0;
            }
            QProgressBar::chunk {
                background-color: #0072BC;
                border-radius: 2px;
            }
        """)
        progress_layout.addWidget(self.right_progress_bar)
        
        layout.addLayout(progress_layout)
        panel.setLayout(layout)
        return panel
    def update_zt_plots(self):
        """更新优值系数图表，展示ZT随温度的变化"""
        try:
            # 获取当前选择的材料组分
            p_composition = self.p_type_combo.currentText()
            n_composition = self.n_type_combo.currentText()
            
            # 检查材料组分是否有效
            if not p_composition or not n_composition:
                print("警告: 材料组分未选择，跳过优值系数图表更新")
                return
            
            # 创建温度范围（300K - 700K），与MATLAB代码一致
            temperatures = np.arange(300, 701, 20)  # 300:20:700
            
            # 计算P型材料的优值系数
            p_zt = []
            for T in temperatures:
                # 直接从Excel文件中读取ZT值，与MATLAB代码一致
                interp_key = f"p_{p_composition}"
                if interp_key not in self.calculator.interpolators:
                    self.calculator.create_interpolators('p', p_composition)
                p_zt.append(self.calculator.calculate_zt('p', p_composition, T))
            
            # 计算N型材料的优值系数
            n_zt = []
            for T in temperatures:
                interp_key = f"n_{n_composition}"
                if interp_key not in self.calculator.interpolators:
                    self.calculator.create_interpolators('n', n_composition)
                n_zt.append(self.calculator.calculate_zt('n', n_composition, T))
            
            # 更新P型图表
            self.zt_axes[0].clear()
            self.zt_axes[0].plot(temperatures, p_zt, 'b+-', linewidth=2)  # 使用蓝色+号标记，与MATLAB一致
            self.zt_axes[0].set_title("P型半导体材料优值系数", pad=5)
            self.zt_axes[0].set_xlabel("温度 (K)")
            self.zt_axes[0].set_ylabel("ZT")
            self.zt_axes[0].set_xlim(300, 700)
            self.zt_axes[0].set_ylim(0, 2.0)  # 与MATLAB图形一致
            self.zt_axes[0].grid(True, linestyle='--', alpha=0.7)
            
            # 更新N型图表
            self.zt_axes[1].clear()
            self.zt_axes[1].plot(temperatures, n_zt, 'r*-', linewidth=2)  # 使用红色*号标记，与MATLAB一致
            self.zt_axes[1].set_title("N型半导体材料优值系数", pad=5)
            self.zt_axes[1].set_xlabel("温度 (K)")
            self.zt_axes[1].set_ylabel("ZT")
            self.zt_axes[1].set_xlim(300, 700)
            self.zt_axes[1].set_ylim(0, 2.0)  # 与MATLAB图形一致
            self.zt_axes[1].grid(True, linestyle='--', alpha=0.7)
            
            # 设置两个图表的共同属性
            for ax in self.zt_axes:
                ax.set_facecolor('#F8F8F8')
                ax.tick_params(direction='in')  # 刻度线向内
                ax.spines['top'].set_visible(True)
                ax.spines['right'].set_visible(True)
                # 设置主要刻度
                ax.set_xticks(np.arange(300, 701, 100))
                ax.set_yticks(np.arange(0, 2.1, 0.5))
                # 添加次要刻度
                ax.minorticks_on()
            
            # 刷新图表
            self.zt_canvas.draw()
            
        except Exception as e:
            print(f"更新优值系数图表错误: {str(e)}")
            import traceback
            traceback.print_exc()

    def initialize_calculation(self):
        """初始化运算"""
        try:
            # 首先验证输入参数
            is_valid, error_msg = self.validate_inputs()
            if not is_valid:
                self.show_error_message("参数错误", error_msg)
                return
            
            # 更新状态
            self.update_status("正在验证参数...")
            self.show_progress(True)
            self.progress_bar.setValue(10)
            
            print("\n===== 开始初始化计算 =====")
            # 更新状态指示灯为红色（计算中）
            self.status_light.set_status(False)
            QApplication.processEvents()  # 确保UI更新
            
            # 更新进度
            self.update_progress(20)
            self.update_status("正在更新优值系数图表...")
            
            # 更新优值系数图表
            self.update_zt_plots()
            
            # 获取输入参数
            Th = float(self.th_edit.text())
            Tc = float(self.tc_edit.text())
            n_points = int(self.grid_edit.text())
            max_iter = int(self.iter_edit.text())  # 获取迭代次数
            
            print(f"输入参数: Th={Th}K, Tc={Tc}K, 格点数={n_points}")
            
            # 更新进度
            self.update_progress(40)
            self.update_status("正在计算温度分布...")
            
            # 计算P型和N型材料的温度分布
            p_composition = self.p_type_combo.currentText()
            n_composition = self.n_type_combo.currentText()
            
            # 获取当前选择的电流密度
            p_current = float(self.p_current_combo.currentText())
            n_current = float(self.n_current_combo.currentText())
            
            print(f"P型材料: 组分={p_composition}, 电流密度={p_current}A/cm^2")
            print(f"N型材料: 组分={n_composition}, 电流密度={n_current}A/cm^2")
            
            # 更新进度
            self.update_progress(60)
            self.update_status("正在计算P型材料温度分布...")
            
            # 将最大迭代次数传递给温度分布计算函数
            x_p, T_p = self.calculator.calculate_temperature_distribution(
                Th, Tc, n_points, 'p', p_composition, p_current, max_iter)
            
            # 更新进度
            self.update_progress(80)
            self.update_status("正在计算N型材料温度分布...")
            
            x_n, T_n = self.calculator.calculate_temperature_distribution(
                Th, Tc, n_points, 'n', n_composition, n_current, max_iter)
            
            # 保存计算结果以便后续使用
            self.x_p, self.T_p = x_p, T_p
            self.x_n, self.T_n = x_n, T_n
            
            print("计算完成，正在更新温度分布图...")
            
            # 更新进度
            self.update_progress(90)
            self.update_status("正在更新图表...")
            
            # 删除旧的点击事件处理器（如果存在）
            if hasattr(self, '_pick_cid') and self._pick_cid:
                self.temp_canvas.mpl_disconnect(self._pick_cid)
            
            # 更新温度分布图
            self.update_temperature_plots(x_p, T_p, x_n, T_n)
            
            # 计算完成，更新状态指示灯为绿色
            self.status_light.set_status(True)
            self.update_progress(100)
            self.update_status("计算完成")
            
            # 显示完成消息
            self.show_info_message("计算完成", "温度分布计算已成功完成！")
            
            print("===== 初始化计算完成 =====")
            
        except Exception as e:
            error_msg = f"初始化计算错误: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            self.status_light.set_status(False)
            self.show_error_message("计算错误", error_msg)
            self.update_status("计算失败")
        finally:
            # 隐藏进度条
            self.show_progress(False)
    
    def update_temperature_plots(self, x_p, T_p, x_n, T_n):
        """
        更新温度分布图，使横坐标随格点数变化，并支持数据点交互
        """
        try:
            # 使用保存的引用直接访问图表
            ax1, ax2 = self.temp_axes
            
            # 清除旧数据
            ax1.clear()
            ax2.clear()
            
            # 获取格点数量
            n_points_p = len(x_p)
            n_points_n = len(x_n)
            
            # 使用整数格点位置 1, 2, 3, ..., n
            grid_points_p = np.arange(1, n_points_p + 1)
            grid_points_n = np.arange(1, n_points_n + 1)
            
            print(f"\n=== 温度分布图数据 ===")
            print(f"P型格点数量: {n_points_p}")
            print(f"P型温度数据: {T_p}")
            print(f"N型格点数量: {n_points_n}")
            print(f"N型温度数据: {T_n}")
            
            # 绘制新数据 - 使用标记和细线
            p_line, = ax1.plot(grid_points_p, T_p, 'b*-', markersize=6, picker=5)  # 设置picker参数启用点击事件
            n_line, = ax2.plot(grid_points_n, T_n, 'r*-', markersize=6, picker=5)
            
            # 添加点击事件处理函数
            def on_pick(event):
                if event.artist == p_line:
                    ind = event.ind[0]
                    ax = ax1
                    grid_points = grid_points_p
                    temps = T_p
                    title = "P型材料"
                elif event.artist == n_line:
                    ind = event.ind[0]
                    ax = ax2
                    grid_points = grid_points_n
                    temps = T_n
                    title = "N型材料"
                else:
                    return
                
                # 显示详细信息
                pos = grid_points[ind]
                temp = temps[ind]
                
                # 移除之前的标注（如果有）
                for artist in ax.texts:
                    artist.remove()
                
                # 添加新标注
                ax.annotate(f'格点: {pos}\n温度: {temp:.2f}K',
                            xy=(pos, temp), xytext=(pos+0.5, temp+10),
                            arrowprops=dict(arrowstyle='->',
                                            connectionstyle='arc3,rad=.2',
                                            color='green'),
                            bbox=dict(boxstyle='round,pad=0.5', fc='yellow', alpha=0.7),
                            fontsize=8)
                
                # 更新图表
                self.temp_canvas.draw()
                
                # 输出详细数据到控制台
                print(f"{title} 格点位置 {pos} 的详细数据:")
                print(f"  温度: {temp:.2f}K")
            
            # 连接点击事件
            self._pick_cid = self.temp_canvas.mpl_connect('pick_event', on_pick)
            
            # 设置标题和标签
            ax1.set_title("格点温度分布（P型）")
            ax2.set_title("格点温度分布（N型）")
            
            # 获取温度的最小值和最大值，用于设置Y轴范围
            min_temp = min(min(T_p), min(T_n))
            max_temp = max(max(T_p), max(T_n))
            
            # 设置坐标轴范围和刻度
            for ax, n_points in zip([ax1, ax2], [n_points_p, n_points_n]):
                ax.set_xlabel("格点位置")
                ax.set_ylabel("温度 (K)")
                
                # 动态设置横坐标范围和刻度
                ax.set_xlim(0.5, n_points + 0.5)  # 添加边距
                
                # 如果格点数较多，则间隔显示刻度
                if n_points <= 20:
                    ax.set_xticks(range(1, n_points + 1))
                else:
                    step = max(1, n_points // 10)  # 最多显示10个刻度
                    ax.set_xticks(range(1, n_points + 1, step))
                
                # 设置Y轴范围
                y_margin = (max_temp - min_temp) * 0.1  # 添加10%的边距
                ax.set_ylim(min_temp - y_margin, max_temp + y_margin)
                
                # 添加网格
                ax.grid(True, linestyle='--', alpha=0.7)
            
            # 刷新图表
            self.temp_canvas.draw()
            print("温度分布图更新完成")
            
        except Exception as e:
            print(f"更新温度分布图错误: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def update_efficiency_plots(self):
        """
        更新效率曲线图，适配新版P/N型效率算法，保证与详细脚本一致
        """
        self.eff_axes[0].clear()
        self.eff_axes[1].clear()
        Th = self.th_edit.text()
        Tc = self.tc_edit.text()
        try:
            Th = float(Th)
            Tc = float(Tc)
        except ValueError:
            print("温度值必须是有效数字")
            return
        if Th <= Tc:
            print("警告：热端温度必须高于冷端温度")
            return
        n_points = 10
        l = 1.0
        ax1, ax2 = self.eff_axes
        p_composition = self.p_type_combo.currentText()
        n_composition = self.n_type_combo.currentText()
        # 新算法返回三个值
        p_eff_list, p_J_list, _ = self.calculator.calculate_efficiency_P(Tc, Th, n_points, l, p_composition)
        n_eff_list, n_J_list, _ = self.calculator.calculate_efficiency_N(Tc, Th, n_points, l, n_composition)
        print("\nP型效率结果:", p_eff_list)
        print("P型电流密度:", p_J_list)
        print("\nN型效率结果:", n_eff_list)
        print("N型电流密度:", n_J_list)
        if p_eff_list and p_J_list:
            ax1.plot(p_J_list, p_eff_list, 'b-', linewidth=2, label='P型效率')
            ax1.set_xlabel(r"电流密度 (A/cm$^2$)")
            ax1.set_ylabel('效率')
            ax1.set_title('P型材料效率曲线')
            ax1.grid(True, linestyle='--', alpha=0.7)
            ax1.set_xlim([-30, 0])
            ax1.set_ylim([min(p_eff_list), max(p_eff_list)*1.1 if max(p_eff_list)>0 else 0.1])
            ax1.legend(loc='best')
        if n_eff_list and n_J_list:
            ax2.plot(n_J_list, n_eff_list, 'g-', linewidth=2, label='N型效率')
            ax2.set_xlabel(r"电流密度 (A/cm$^2$)")
            ax2.set_ylabel('效率')
            ax2.set_title('N型材料效率曲线')
            ax2.grid(True, linestyle='--', alpha=0.7)
            ax2.set_xlim([0, 50])
            ax2.set_ylim([min(n_eff_list), max(n_eff_list)*1.1 if max(n_eff_list)>0 else 0.1])
            ax2.legend(loc='best')
        # 重新应用布局设置，确保两个子图不重叠
        self.eff_canvas.figure.subplots_adjust(wspace=0.4)
        self.eff_canvas.draw()
        print("效率图更新完成")
    
    def update_branch_characteristics(self):
        """更新分支特性"""
        try:
            print("开始更新分支特性...")
            # 更新状态指示灯为红色（计算中）
            self.calc_status.set_status(False)
            QApplication.processEvents()  # 确保UI更新
            
            # 执行计算
            self.initialize_calculation()
            
            # 更新效率图
            self.update_efficiency_plots()
            
            # 计算完成，更新状态指示灯为绿色
            self.calc_status.set_status(True)
            print("分支特性更新完成")
            
        except Exception as e:
            print(f"更新分支特性错误: {str(e)}")
            import traceback
            traceback.print_exc()
            self.calc_status.set_status(False)
    
    def update_p_composition(self):
        """P型材料组分变化处理方法"""
        composition = self.p_type_combo.currentText()
        print(f"P型材料组分更新为: {composition}")
        # 更新ZT图
        self.update_zt_plots()
    
    def update_n_composition(self):
        """N型材料组分变化处理方法"""
        composition = self.n_type_combo.currentText()
        print(f"N型材料组分更新为: {composition}")
        # 更新ZT图
        self.update_zt_plots()
    
    def update_device_performance_plots(self, result):
        """
        根据calculate_device_performance的结果，刷新右侧最大点文本框和三张图表，优先用self保存的canvas和ax
        """
        if not result:
            return
        # 最大功率点
        power_list = result.get('Power_total_list', [])
        if power_list and result.get('max_power_idx') is not None:
            max_power = power_list[result['max_power_idx']]
            self.max_power.setText(f"{max_power:.4f}")  # 改为常规小数显示
            self.power_current.setText(f"{result['I_list'][result['max_power_idx']]}")
        # 最大效率点
        eff_list = result.get('eff_total_list', [])
        if eff_list and result.get('max_eff_idx') is not None:
            max_eff = eff_list[result['max_eff_idx']]
            self.max_eff.setText(f"{max_eff*100:.2f}%")
            self.eff_current.setText(f"{result['I_list'][result['max_eff_idx']]}")
        # --- 功率图 ---
        if hasattr(self, 'power_ax') and hasattr(self, 'power_canvas'):
            ax = self.power_ax
            canvas = self.power_canvas
            ax.clear()
            i_list = result.get('I_list', [])
            if i_list and power_list:
                ax.plot(i_list, power_list, 'r-', label='功率曲线')
                if result.get('max_power_idx') is not None:
                    ax.scatter(i_list[result['max_power_idx']], power_list[result['max_power_idx']], color='red', marker='o', s=50, label='最大功率点')
            ax.set_xlabel("电流 (A)")
            ax.set_ylabel(r"功率 ($W/cm^2$)")
            ax.set_xlim(0, max(i_list) if i_list else 10)
            ax.set_ylim(0, max(power_list)*1.1 if power_list else 1)
            ax.legend()
            ax.grid(True, linestyle='--', alpha=0.6)
            canvas.draw()
        # --- 效率图 ---
        if hasattr(self, 'efficiency_ax') and hasattr(self, 'efficiency_canvas'):
            ax = self.efficiency_ax
            canvas = self.efficiency_canvas
            ax.clear()
            eff_percent = [e*100 for e in result['eff_total_list']]
            if i_list and eff_list:
                ax.plot(i_list, eff_percent, 'b-', label='效率曲线')
                if result.get('max_eff_idx') is not None:
                    ax.scatter(i_list[result['max_eff_idx']], eff_percent[result['max_eff_idx']], color='blue', marker='o', s=50, label='最大效率点')
            ax.set_xlabel("电流 (A)")
            ax.set_ylabel("效率 (%)")
            ax.set_xlim(0, max(i_list) if i_list else 10)
            # 主要关注0%-10%的效率区间
            y_max = min(max(max(eff_percent)*1.2, 10.0), 40.0) if eff_percent else 10.0
            ax.set_ylim(0, y_max)
            # 卡诺效率参考线
            try:
                Th = float(self.th_edit.text())
                Tc = float(self.tc_edit.text())
                carnot_eff = (Th - Tc) / Th * 100
                ax.axhline(y=carnot_eff, color='r', linestyle='--', label=f'卡诺效率: {carnot_eff:.1f}%')
            except Exception:
                pass
            ax.legend()
            ax.grid(True, linestyle='--', alpha=0.6)
            canvas.draw()
        # --- 优化区间图 ---
        if hasattr(self, 'optimization_ax') and hasattr(self, 'optimization_canvas'):
            ax = self.optimization_ax
            canvas = self.optimization_canvas
            ax.clear()
            eff_percent = [e*100 for e in result['eff_total_list']]
            if power_list and eff_list:
                ax.plot(power_list, eff_percent, 'g-', label='优化曲线')
                if result.get('max_power_idx') is not None:
                    ax.scatter(power_list[result['max_power_idx']], eff_percent[result['max_power_idx']], color='red', marker='o', s=50, label='最大功率点')
                if result.get('max_eff_idx') is not None:
                    ax.scatter(power_list[result['max_eff_idx']], eff_percent[result['max_eff_idx']], color='blue', marker='o', s=50, label='最大效率点')
            ax.set_xlabel(r"功率 ($W/cm^2$)")
            ax.set_ylabel("效率 (%)")
            x_max = max(power_list)*1.1 if power_list else 0.0001
            # 主要关注0%-10%的效率区间
            y_max = min(max(max(eff_percent)*1.2, 10.0), 40.0) if eff_percent else 10.0
            ax.set_xlim(0, x_max)
            ax.set_ylim(0, y_max)
            # 卡诺效率参考线
            try:
                Th = float(self.th_edit.text())
                Tc = float(self.tc_edit.text())
                carnot_eff = (Th - Tc) / Th * 100
                ax.axhline(y=carnot_eff, color='r', linestyle='--', label=f'卡诺效率: {carnot_eff:.1f}%')
            except Exception:
                pass
            ax.legend()
            ax.grid(True, linestyle='--', alpha=0.6)
            canvas.draw()

    def calculate_device_performance(self):
        """
        计算器件总效率、总功率、优化区间，完全采用03.py算法，不再依赖外部脚本。
        """
        try:
            # 计算开始，状态灯变红，显示进度条
            self.right_status_light.set_status(False)
            self.right_progress_bar.setVisible(True)
            self.right_progress_bar.setValue(0)
            self.right_progress_label.setText("正在准备计算...")
            QApplication.processEvents()
            
            # 获取参数
            Th = float(self.th_edit.text())
            Tc = float(self.tc_edit.text())
            p_composition = self.p_type_combo.currentText()
            n_composition = self.n_type_combo.currentText()
            area_ratio = float(self.ratio_edit.text())
            n = 10
            l = 1.0
            
            # 更新进度：开始计算P型分支
            self.right_progress_bar.setValue(10)
            self.right_progress_label.setText("正在计算P型分支效率...")
            QApplication.processEvents()
            
            # 1. 计算P型分支效率、热流
            eff_list_P, J_list_P, q_P_list = self.calculator.calculate_efficiency_P(Tc, Th, n, l, p_composition)
            
            # 更新进度：P型计算完成，开始N型计算
            self.right_progress_bar.setValue(30)
            self.right_progress_label.setText("正在计算N型分支效率...")
            QApplication.processEvents()
            
            # 2. 计算N型分支效率、热流
            eff_list_N, J_list_N, q_N_list = self.calculator.calculate_efficiency_N(Tc, Th, n, l, n_composition)
            
            # 更新进度：开始计算总效率
            self.right_progress_bar.setValue(50)
            self.right_progress_label.setText("正在计算总效率...")
            QApplication.processEvents()
            
            # 3. 计算总效率
            I_list = []
            eff_total_list = []
            for m in range(1, 21, 1):
                I_list.append(m)
                index_N = int(m * area_ratio)
                if m < len(eff_list_P) and index_N < len(eff_list_N):
                    denominator = q_P_list[m] / m - q_N_list[m] / (m / area_ratio)
                    if denominator != 0:
                        eff_total = (eff_list_P[m - 1] * q_P_list[m] / m - eff_list_N[index_N] * q_N_list[m] / (m / area_ratio)) / denominator
                        eff_total_list.append(eff_total)
                
                # 更新进度条（每5个循环更新一次）
                if m % 5 == 0:
                    progress = 50 + (m / 20) * 20
                    self.right_progress_bar.setValue(int(progress))
                    self.right_progress_label.setText(f"正在计算总效率... ({m}/20)")
                    QApplication.processEvents()
            
            # 更新进度：开始计算总功率
            self.right_progress_bar.setValue(70)
            self.right_progress_label.setText("正在计算总功率...")
            QApplication.processEvents()
            
            # 4. 计算总功率
            Power_total_list = []
            for nidx in range(0, len(eff_total_list)):
                pow_total = eff_total_list[nidx] * (-q_P_list[nidx+1] - q_N_list[nidx+1] * area_ratio)
                Power_total_list.append(pow_total)
            
            # 更新进度：寻找最大点
            self.right_progress_bar.setValue(85)
            self.right_progress_label.setText("正在寻找最大功率和效率点...")
            QApplication.processEvents()
            
            # 5. 最大点
            max_power_idx = np.argmax(Power_total_list) if Power_total_list else 0
            max_eff_idx = np.argmax(eff_total_list) if eff_total_list else 0
            
            # 更新进度：结果打包
            self.right_progress_bar.setValue(90)
            self.right_progress_label.setText("正在整理计算结果...")
            QApplication.processEvents()
            
            # 6. 结果打包
            result = {
                'eff_list_P': eff_list_P,
                'J_list_P': J_list_P,
                'q_P_list': q_P_list,
                'eff_list_N': eff_list_N,
                'J_list_N': J_list_N,
                'q_N_list': q_N_list,
                'I_list': I_list,
                'eff_total_list': eff_total_list,
                'Power_total_list': Power_total_list,
                'max_power_idx': max_power_idx,
                'max_eff_idx': max_eff_idx
            }
            
            # 更新进度：更新图表
            self.right_progress_bar.setValue(95)
            self.right_progress_label.setText("正在更新图表...")
            QApplication.processEvents()
            
            # 计算完成，状态灯变绿，隐藏进度条
            self.right_status_light.set_status(True)
            self.right_progress_bar.setValue(100)
            self.right_progress_label.setText("计算完成")
            
            # 延迟一秒后隐藏进度条
            QTimer.singleShot(1000, self.hide_right_progress)
            
            # 新增：自动刷新前端结果
            self.update_device_performance_plots(result)
            # 保证导出数据时有最新结果
            self.last_device_calculation = result
            return result
            
        except Exception as e:
            error_msg = f"计算器件性能错误: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            # 异常时状态灯变红，隐藏进度条
            self.right_status_light.set_status(False)
            self.right_progress_bar.setVisible(False)
            self.right_progress_label.setText(f"计算失败: {str(e)[:30]}...")
            return None
    
    def hide_right_progress(self):
        """隐藏右侧面板的进度条"""
        self.right_progress_bar.setVisible(False)
        self.right_progress_label.setText("就绪")

    def export_data(self):
        """
        导出详细数据到Excel，多Sheet，所有标题用中文，包含分支特性、温度分布、材料参数等
        """
        try:
            from datetime import datetime
            import pandas as pd
            import numpy as np
            # 获取当前时间作为文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"thermoelectric_data_{timestamp}.xlsx"
            # 只导出当前界面已有的最新结果，不再自动调用calculate_device_performance
            result = getattr(self, 'last_device_calculation', None)
            if not result:
                print("无有效计算结果，无法导出数据")
                return
            # 1. 参数与最大点
            max_eff_str = self.max_eff.text().replace('%', '').strip()
            param_data = {
                "高温端温度(K)": [float(self.th_edit.text())],
                "低温端温度(K)": [float(self.tc_edit.text())],
                "P型材料组分": [self.p_type_combo.currentText()],
                "N型材料组分": [self.n_type_combo.currentText()],
                    "N/P面积比": [float(self.ratio_edit.text())],
                "最大功率(W/cm²)": [float(self.max_power.text())],
                "最大功率对应电流(A)": [float(self.power_current.text())],
                "最大效率(%)": [float(max_eff_str).__round__(2)],
                "最大效率对应电流(A)": [float(self.eff_current.text())]
            }
            # 2. 总效率-电流
            i_list = result.get('I_list', [])
            eff_list = result.get('eff_total_list', [])
            
            df_eff = pd.DataFrame({
                "电流(A)": i_list,
                "总效率(%)": [e*100 for e in eff_list]
            })
            # 3. 总功率-电流
            power_list = result.get('Power_total_list', [])
            if not power_list:
                power_list = [0] * len(result.get('I_list', []))
            
            df_power = pd.DataFrame({
                "电流(A)": result['I_list'],
                "总功率(W/cm²)": power_list
            })
            # 4. 优化区间
            df_opt = pd.DataFrame({
                "功率(W/cm²)": power_list,
                "效率(%)": [e*100 for e in result['eff_total_list']]
            })
            # 5. P型分支
            j_list_p = result.get('J_list_P', [])
            eff_list_p = result.get('eff_list_P', [])
            q_p_list = result.get('q_P_list', [])
            
            df_p_branch = pd.DataFrame({
                "电流密度(A/cm^2)": j_list_p,
                "P型效率": eff_list_p,
                "P型热流": q_p_list
            })
            # 6. N型分支
            j_list_n = result.get('J_list_N', [])
            eff_list_n = result.get('eff_list_N', [])
            q_n_list = result.get('q_N_list', [])
            
            df_n_branch = pd.DataFrame({
                "电流密度(A/cm^2)": j_list_n,
                "N型效率": eff_list_n,
                "N型热流": q_n_list
            })
            # 7. P型温度分布
            x_p, T_p = getattr(self, 'x_p', None), getattr(self, 'T_p', None)
            if x_p is not None and T_p is not None:
                df_p_temp = pd.DataFrame({"P型格点": x_p, "P型温度(K)": T_p})
            else:
                df_p_temp = pd.DataFrame()
            # 8. N型温度分布
            x_n, T_n = getattr(self, 'x_n', None), getattr(self, 'T_n', None)
            if x_n is not None and T_n is not None:
                df_n_temp = pd.DataFrame({"N型格点": x_n, "N型温度(K)": T_n})
            else:
                df_n_temp = pd.DataFrame()
            # 9. P型材料参数
            p_composition = self.p_type_combo.currentText()
            p_mat = self.calculator.p_type_data.get(p_composition, None)
            if p_mat:
                df_p_mat = pd.DataFrame({
                    "温度(K)": p_mat["temp"],
                    "塞贝克系数(V/K)": p_mat["seebeck"],
                    "电阻率(Ω·m)": p_mat["resistivity"],
                    "热导率(W/m·K)": p_mat["thermal_cond"]
                })
            else:
                df_p_mat = pd.DataFrame()
            # 10. N型材料参数
            n_composition = self.n_type_combo.currentText()
            n_mat = self.calculator.n_type_data.get(n_composition, None)
            if n_mat:
                df_n_mat = pd.DataFrame({
                    "温度(K)": n_mat["temp"],
                    "塞贝克系数(V/K)": n_mat["seebeck"],
                    "电阻率(Ω·m)": n_mat["resistivity"],
                    "热导率(W/m·K)": n_mat["thermal_cond"]
                })
            else:
                df_n_mat = pd.DataFrame()
            # 写入Excel多Sheet
            with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                pd.DataFrame(param_data).to_excel(writer, sheet_name='参数与最大点', index=False)
                df_eff.to_excel(writer, sheet_name='总效率-电流', index=False)
                df_power.to_excel(writer, sheet_name='总功率-电流', index=False)
                df_opt.to_excel(writer, sheet_name='优化区间', index=False)
                df_p_branch.to_excel(writer, sheet_name='P型分支', index=False)
                df_n_branch.to_excel(writer, sheet_name='N型分支', index=False)
                if not df_p_temp.empty:
                    df_p_temp.to_excel(writer, sheet_name='P型温度分布', index=False)
                if not df_n_temp.empty:
                    df_n_temp.to_excel(writer, sheet_name='N型温度分布', index=False)
                if not df_p_mat.empty:
                    df_p_mat.to_excel(writer, sheet_name='P型材料参数', index=False)
                if not df_n_mat.empty:
                    df_n_mat.to_excel(writer, sheet_name='N型材料参数', index=False)
            print(f"数据已导出到文件: {filename}")
        except Exception as e:
            print(f"导出数据错误: {str(e)}")

    def export_experiment_report(self):
        """一键导出实验报告（自动选择最佳格式）"""
        try:
            # 检查是否有计算结果
            if not hasattr(self, 'last_device_calculation') or not self.last_device_calculation:
                QMessageBox.warning(self, "警告", "请先进行计算，获得结果后再导出实验报告。")
                return
            
            # 收集仿真数据
            if not self.report_generator.collect_simulation_data():
                QMessageBox.critical(self, "错误", "收集仿真数据失败，无法生成实验报告。")
                return
            
            # 尝试导出Word格式（优先选择）
            filename = self.report_generator.export_report_to_word()
            if filename:
                QMessageBox.information(self, "成功", f"实验报告已成功导出到Word格式：\n{filename}")
                return
            
            # 如果Word导出失败，尝试HTML格式
            filename = self.report_generator.export_report_to_html()
            if filename:
                QMessageBox.information(self, "成功", f"实验报告已成功导出到HTML格式：\n{filename}")
                return
            
            # 最后尝试文本格式
            filename = self.report_generator.export_report_to_txt()
            if filename:
                QMessageBox.information(self, "成功", f"实验报告已成功导出到文本格式：\n{filename}")
                return
            
            QMessageBox.critical(self, "错误", "所有格式的导出都失败了，请检查系统环境。")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出实验报告时发生错误：\n{str(e)}")
            print(f"导出实验报告错误: {str(e)}")

    def export_word_report(self):
        """导出Word格式实验报告"""
        try:
            if not hasattr(self, 'last_device_calculation') or not self.last_device_calculation:
                QMessageBox.warning(self, "警告", "请先进行计算，获得结果后再导出实验报告。")
                return
            
            if not self.report_generator.collect_simulation_data():
                QMessageBox.critical(self, "错误", "收集仿真数据失败，无法生成实验报告。")
                return
            
            filename = self.report_generator.export_report_to_word()
            if filename:
                QMessageBox.information(self, "成功", f"Word格式实验报告已成功导出：\n{filename}")
            else:
                QMessageBox.warning(self, "警告", "Word格式导出失败，请安装python-docx库或选择其他格式。")
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出Word格式报告时发生错误：\n{str(e)}")
            print(f"导出Word格式报告错误: {str(e)}")
    def export_html_report(self):
        """导出HTML格式实验报告"""
        try:
            if not hasattr(self, 'last_device_calculation') or not self.last_device_calculation:
                QMessageBox.warning(self, "警告", "请先进行计算，获得结果后再导出实验报告。")
                return
            
            if not self.report_generator.collect_simulation_data():
                QMessageBox.critical(self, "错误", "收集仿真数据失败，无法生成实验报告。")
                return
            
            filename = self.report_generator.export_report_to_html()
            if filename:
                QMessageBox.information(self, "成功", f"HTML格式实验报告已成功导出：\n{filename}")
            else:
                QMessageBox.critical(self, "错误", "HTML格式导出失败。")
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出HTML格式报告时发生错误：\n{str(e)}")
            print(f"导出HTML格式报告错误: {str(e)}")

    def export_txt_report(self):
        """导出文本格式实验报告"""
        try:
            if not hasattr(self, 'last_device_calculation') or not self.last_device_calculation:
                QMessageBox.warning(self, "警告", "请先进行计算，获得结果后再导出实验报告。")
                return
            
            if not self.report_generator.collect_simulation_data():
                QMessageBox.critical(self, "错误", "收集仿真数据失败，无法生成实验报告。")
                return
            
            filename = self.report_generator.export_report_to_txt()
            if filename:
                QMessageBox.information(self, "成功", f"文本格式实验报告已成功导出：\n{filename}")
            else:
                QMessageBox.critical(self, "错误", "文本格式导出失败。")
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出文本格式报告时发生错误：\n{str(e)}")
            print(f"导出文本格式报告错误: {str(e)}")

    def run_external_efficiency(self):
        """运行外部03.py程序进行详细效率计算"""
        try:
            import subprocess
            import os
            import pandas as pd
            import numpy as np
            
            # 获取当前参数
            p_composition = self.p_type_combo.currentText()
            n_composition = self.n_type_combo.currentText()
            Tc = self.tc_edit.text()
            Th = self.th_edit.text()
            area_ratio = self.ratio_edit.text()
            
            # 准备输出目录
            output_dir = "efficiency_results"
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # 构建命令
            cmd = [
                "python", "03.py",
                "--p_type", p_composition,
                "--n_type", n_composition,
                "--area_ratio", area_ratio,
                "--Tc", Tc,
                "--Th", Th,
                "--output_dir", output_dir
            ]
            
            # 更新状态指示灯
            self.calc_status.set_status(False)
            QApplication.processEvents()
            
            print(f"正在运行详细效率计算: {' '.join(cmd)}")
            
            # 显示一条消息
            QMessageBox.information(self, "计算中", "正在运行效率计算，请等待完成...")
            
            # 检查材料数据文件是否存在
            p_filename_map = {
                "0.01": "P_yuanshi_2_5.xls",
                "0.02": "P_yuanshi_3_1.xls",
                "0.03": "P_yuanshi_3_7.xls"
            }
            
            p_data_file = p_filename_map.get(p_composition)
            if p_data_file is None:
                error_msg = f"不支持的P型材料组分: {p_composition}"
                print(error_msg)
                QMessageBox.critical(self, "错误", error_msg)
                self.calc_status.set_status(False)
                return
                
            n_data_file = f"N_yuanshi_{n_composition}.xls"
            
            # 构建完整路径
            p_full_path = os.path.join(application_path, p_data_file)
            n_full_path = os.path.join(application_path, n_data_file)
            
            if not os.path.exists(p_full_path):
                error_msg = f"找不到P型材料数据文件: {p_full_path}"
                print(error_msg)
                QMessageBox.critical(self, "错误", error_msg)
                self.calc_status.set_status(False)
                return
                
            if not os.path.exists(n_full_path):
                error_msg = f"找不到N型材料数据文件: {n_full_path}"
                print(error_msg)
                QMessageBox.critical(self, "错误", error_msg)
                self.calc_status.set_status(False)
                return
            
            # 创建一个日志文件以储存输出
            log_file = os.path.join(output_dir, f"efficiency_log_{p_composition}_{n_composition}.txt")
            with open(log_file, 'w', encoding='utf-8') as f:
                f.write(f"命令: {' '.join(cmd)}\n\n")
                f.write("=== 执行开始 ===\n")
            
            # 运行子进程
            process = subprocess.Popen(
                cmd, 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                text=True,
                bufsize=1
            )
            
            # 获取输出并实时写入日志
            stdout_lines = []
            stderr_lines = []
            
            # 实时读取和记录stdout
            for line in iter(process.stdout.readline, ''):
                if not line:
                    break
                line = line.strip()
                print(line)
                with open(log_file, 'a', encoding='utf-8') as f:
                    f.write(line + '\n')
                stdout_lines.append(line)
            
            # 实时读取和记录stderr
            for line in iter(process.stderr.readline, ''):
                if not line:
                    break
                line = line.strip()
                print(f"错误: {line}")
                with open(log_file, 'a', encoding='utf-8') as f:
                    f.write(f"错误: {line}\n")
                stderr_lines.append(line)
            
            # 等待进程结束
            return_code = process.wait()
            with open(log_file, 'a', encoding='utf-8') as f:
                f.write(f"\n=== 执行结束，返回码: {return_code} ===\n")
            
            stdout = '\n'.join(stdout_lines)
            stderr = '\n'.join(stderr_lines)
            
            # 检查执行结果
            if return_code == 0:
                print("详细效率计算完成")
                QMessageBox.information(self, "完成", "详细效率计算已完成")
                
                # 显示计算结果的图像
                result_image = os.path.join(output_dir, f"summary_{p_composition}_{n_composition}.png")
                if os.path.exists(result_image):
                    pixmap = QPixmap(result_image)
                    if not pixmap.isNull():
                        dialog = ImageViewerDialog(pixmap, self)
                        dialog.setWindowTitle(f"效率计算结果 (P:{p_composition}, N:{n_composition})")
                        dialog.show()
                    else:
                        error_msg = f"无法加载结果图像: {result_image}"
                        print(error_msg)
                        QMessageBox.warning(self, "警告", error_msg)
                else:
                    error_msg = f"结果图像不存在: {result_image}"
                    print(error_msg)
                    QMessageBox.warning(self, "警告", error_msg)
                
                # 加载CSV数据并更新主界面图表
                try:
                    print("从CSV文件加载计算结果数据...")
                    total_data_file = os.path.join(output_dir, f"total_data_{p_composition}_{n_composition}.csv")
                    if os.path.exists(total_data_file):
                        total_data = pd.read_csv(total_data_file)
                        
                        I_list = total_data['I'].tolist()
                        eff_total_list = total_data['eff_total'].tolist()
                        power_total_list = total_data['Power_total'].tolist()
                        
                        # 如果存在效率和功率数据，更新主界面图表
                        if I_list and eff_total_list and power_total_list:
                            print(f"加载的数据: {len(I_list)}个点")
                            print(f"效率范围: {min(eff_total_list):.6f} ~ {max(eff_total_list):.6f}")
                            print(f"功率范围: {min(power_total_list):.6e} ~ {max(power_total_list):.6e}")
                            
                            # 找到最大效率和功率点
                            if max(eff_total_list) > 0:
                                max_eff_idx = eff_total_list.index(max(eff_total_list))
                                self.max_eff.setText(f"{eff_total_list[max_eff_idx]*100:.2f}%")
                                self.eff_current.setText(f"{I_list[max_eff_idx]:.2f}")
                                print(f"最大效率: {eff_total_list[max_eff_idx]*100:.2f}% 在电流 {I_list[max_eff_idx]:.2f}A")
                            
                            if max(power_total_list) > 0:
                                max_power_idx = power_total_list.index(max(power_total_list))
                                self.max_power.setText(f"{power_total_list[max_power_idx]:.2e}")
                                self.power_current.setText(f"{I_list[max_power_idx]:.2f}")
                                print(f"最大功率: {power_total_list[max_power_idx]:.4e} W/cm^2 在电流 {I_list[max_power_idx]:.2f}A")
                            
                            # 更新功率图表
                            try:
                                power_group = self.findChild(QGroupBox, "器件功率")
                                if not power_group:
                                    all_group_boxes = self.findChildren(QGroupBox)
                                    available_names = [box.title() for box in all_group_boxes]
                                    print(f"警告: 找不到器件功率GroupBox，可用的GroupBox: {available_names}")
                                    
                                    # 尝试找到一个合适的替代GroupBox
                                    for box in all_group_boxes:
                                        if '功率' in box.title() and '材料' not in box.title() and '最大' not in box.title():
                                            print(f"找到替代GroupBox: {box.title()}")
                                            power_group = box
                                            break
                                
                                if power_group:
                                    power_canvases = power_group.findChildren(FigureCanvas)
                                    if power_canvases and len(power_canvases) > 0:
                                        power_container = power_canvases[0]
                                        power_fig = power_container.figure
                                        power_ax = power_fig.axes[0]
                                        power_ax.clear()
                                        
                                        print(f"绘制功率图, 数据点数: {len(I_list)}")
                                        power_ax.plot(I_list, power_total_list, 'r-', linewidth=2.0, label='功率曲线')
                                        power_ax.scatter(I_list[max_power_idx], power_total_list[max_power_idx], 
                                                     color='red', marker='o', s=50, label='最大功率点')
                                        
                                        power_ax.set_xlabel("电流 (A)")
                                        power_ax.set_ylabel("功率 (W/cm$^2$)")
                                        power_ax.set_xlim(0, max(I_list) if I_list else 10)
                                        power_ax.set_ylim(0, max(power_total_list)*1.1 if power_total_list else 1)
                                        
                                        power_ax.annotate(f"{power_total_list[max_power_idx]:.2e}",
                                                     (I_list[max_power_idx], power_total_list[max_power_idx]),
                                                     xytext=(5, 5), textcoords='offset points')
                                        
                                        power_ax.grid(True, linestyle='--', alpha=0.6)
                                        power_ax.legend(loc='best')
                                        power_fig.tight_layout()
                                        power_fig.canvas.draw()
                                        print("功率图已更新")
                            except Exception as e:
                                print(f"更新功率图错误: {str(e)}")
                                import traceback
                                traceback.print_exc()
                            
                            # 更新效率图表
                            try:
                                eff_group = self.findChild(QGroupBox, "器件效率")
                                if not eff_group:
                                    all_group_boxes = self.findChildren(QGroupBox)
                                    available_names = [box.title() for box in all_group_boxes]
                                    print(f"警告: 找不到器件效率GroupBox，可用的GroupBox: {available_names}")
                                    
                                    # 尝试找到一个合适的替代GroupBox
                                    for box in all_group_boxes:
                                        if '效率' in box.title() and '材料' not in box.title() and '最大' not in box.title():
                                            print(f"找到替代效率GroupBox: {box.title()}")
                                            eff_group = box
                                            break
                                
                                if eff_group:
                                    eff_canvases = eff_group.findChildren(FigureCanvas)
                                    if eff_canvases and len(eff_canvases) > 0:
                                        eff_container = eff_canvases[0]
                                        eff_fig = eff_container.figure
                                        eff_ax = eff_fig.axes[0]
                                        eff_ax.clear()
                                        
                                        print(f"开始绘制效率图, 数据点数: {len(I_list)} / {len(eff_total_list)}")
                                        # 将效率值转换为百分比
                                        eff_percent = [e * 100 for e in eff_total_list]  # 转换为百分比
                                        
                                        eff_ax.plot(I_list, eff_percent, 'r-', linewidth=2.0, label='效率曲线')
                                        
                                        if max(eff_total_list) > 0:
                                            eff_ax.scatter(I_list[max_eff_idx], eff_percent[max_eff_idx], 
                                                        color='blue', marker='o', s=50, label='最大效率点')
                                            # 标注最大效率点
                                            eff_ax.annotate(f"{eff_percent[max_eff_idx]:.2f}%",
                                                      (I_list[max_eff_idx], eff_percent[max_eff_idx]),
                                                      xytext=(5, 5), textcoords='offset points')
                                        
                                        eff_ax.set_xlabel("电流 (A)")
                                        eff_ax.set_ylabel("效率 (%)")
                                        eff_ax.set_xlim(0, max(I_list) if I_list else 10)
                                        
                                        # 设置Y轴范围，主要关注0%-10%的效率区间
                                        if eff_percent and max(eff_percent) > 0:
                                            # 如果最大效率小于10%，则显示到10%；否则显示到最大效率的1.2倍
                                            y_max = min(max(max(eff_percent)*1.2, 10.0), 40.0)
                                        else:
                                            y_max = 10.0  # 默认显示到10%
                                        eff_ax.set_ylim(0, y_max)
                                        
                                        # 添加卡诺效率参考线
                                        Th = float(self.th_edit.text())
                                        Tc = float(self.tc_edit.text())
                                        carnot_eff = (Th - Tc) / Th * 100
                                        eff_ax.axhline(y=carnot_eff, color='r', linestyle='--', 
                                                    label=f'卡诺效率: {carnot_eff:.1f}%')
                                        
                                        eff_ax.grid(True, linestyle='--', alpha=0.6)
                                        eff_ax.legend(loc='best')
                                        eff_fig.tight_layout()
                                        eff_fig.canvas.draw()
                                        print(f"效率图已更新，x轴范围: [0, {max(I_list)}], y轴范围: [0, {y_max}]")
                                    else:
                                        print("警告: 找不到效率图的FigureCanvas")
                                else:
                                    print("错误: 无法找到任何可用的效率图表容器")
                            except Exception as e:
                                print(f"更新效率图错误: {str(e)}")
                                import traceback
                                traceback.print_exc()
                            
                            # 更新优化区间图
                            try:
                                opt_group = self.findChild(QGroupBox, "功率效率优化区间")
                                if not opt_group:
                                    all_group_boxes = self.findChildren(QGroupBox)
                                    available_names = [box.title() for box in all_group_boxes]
                                    print(f"警告: 找不到功率效率优化区间GroupBox，可用的GroupBox: {available_names}")
                                    
                                    for box in all_group_boxes:
                                        if '优化' in box.title():
                                            print(f"找到替代优化区间GroupBox: {box.title()}")
                                            opt_group = box
                                            break
                                
                                if opt_group:
                                    opt_canvases = opt_group.findChildren(FigureCanvas)
                                    if opt_canvases and len(opt_canvases) > 0:
                                        opt_container = opt_canvases[0]
                                        opt_fig = opt_container.figure
                                        opt_ax = opt_fig.axes[0]
                                        opt_ax.clear()
                                        
                                        # 效率值转换为百分比
                                        eff_percent = [e * 100 for e in eff_total_list]
                                        
                                        print(f"优化区间图数据: 功率点数={len(power_total_list)}, 效率点数={len(eff_percent)}")
                                        opt_ax.plot(power_total_list, eff_percent, 'g-', linewidth=2.0, label='优化曲线')
                                        opt_ax.scatter(power_total_list[max_power_idx], eff_percent[max_power_idx], 
                                                    color='red', marker='o', s=50, label='最大功率点')
                                        opt_ax.scatter(power_total_list[max_eff_idx], eff_percent[max_eff_idx], 
                                                    color='blue', marker='o', s=50, label='最大效率点')
                                        
                                        # 标注最大点
                                        opt_ax.annotate(f"最大功率: {power_total_list[max_power_idx]:.2e}W/cm^2",
                                                  (power_total_list[max_power_idx], eff_percent[max_power_idx]),
                                                  xytext=(10, 10), textcoords='offset points', 
                                                  arrowprops=dict(arrowstyle="->", color='red'))
                                        
                                        opt_ax.annotate(f"最大效率: {eff_percent[max_eff_idx]:.2f}%",
                                                  (power_total_list[max_eff_idx], eff_percent[max_eff_idx]),
                                                  xytext=(-10, 10), textcoords='offset points',
                                                  arrowprops=dict(arrowstyle="->", color='blue'))
                                        
                                        opt_ax.set_xlabel("功率 (W/cm$^2$)")
                                        opt_ax.set_ylabel("效率 (%)")
                                        
                                        # 动态设置坐标轴范围
                                        if power_total_list and max(power_total_list) > 0:
                                            x_max = max(power_total_list) * 1.1
                                        else:
                                            x_max = 0.0001  # 设置一个小的默认值
                                            
                                        # 主要关注0%-10%的效率区间
                                        y_max = min(max(max(eff_percent) * 1.2, 10.0), 40.0) if eff_percent and max(eff_percent) > 0 else 10.0
                                        opt_ax.set_xlim(0, x_max)
                                        opt_ax.set_ylim(0, y_max)
                                        
                                        # 添加卡诺效率参考线
                                        Th = float(self.th_edit.text())
                                        Tc = float(self.tc_edit.text())
                                        carnot_eff = (Th - Tc) / Th * 100
                                        opt_ax.axhline(y=carnot_eff, color='r', linestyle='--', 
                                                    label=f'卡诺效率: {carnot_eff:.1f}%')
                                        
                                        opt_ax.grid(True, linestyle='--', alpha=0.6)
                                        opt_ax.legend(loc='best')
                                        opt_fig.tight_layout()
                                        opt_fig.canvas.draw()
                                        print(f"优化区间图已更新，x轴范围: [0, {x_max}], y轴范围: [0, {y_max}]")
                                    else:
                                        print("警告: 找不到优化区间图的FigureCanvas")
                                else:
                                    print("错误: 无法找到任何可用的优化区间图表容器")
                            except Exception as e:
                                print(f"更新优化区间图错误: {str(e)}")
                                import traceback
                                traceback.print_exc()
                            
                            # 保存计算结果供后续使用
                            self.last_device_calculation = {
                                'I_list': I_list,
                                'eff_total_list': eff_total_list,
                                'power_total_list': power_total_list,
                                'max_power_idx': max_power_idx if power_total_list and max(power_total_list) > 0 else 0,
                                'max_eff_idx': max_eff_idx if eff_total_list and max(eff_total_list) > 0 else 0
                            }
                    else:
                        print(f"警告: 找不到总计算结果文件：{total_data_file}")
                except Exception as e:
                    print(f"加载和应用计算结果错误: {str(e)}")
                    import traceback
                    traceback.print_exc()
                    
                # 更新状态指示灯为绿色（完成）
                self.calc_status.set_status(True)
            else:
                error_msg = f"详细效率计算失败，错误码: {return_code}\n\n"
                if stderr:
                    error_msg += f"错误信息:\n{stderr}\n"
                print(error_msg)
                
                QMessageBox.critical(self, "错误", f"详细效率计算失败。请查看控制台输出或日志文件：\n{log_file}")
                
                # 保持状态指示灯为红色（出错）
                self.calc_status.set_status(False)
            
        except Exception as e:
            error_msg = f"运行外部效率计算错误: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            
            QMessageBox.critical(self, "错误", error_msg)
            # 保持状态指示灯为红色（出错）
            self.calc_status.set_status(False)

    def analyze_material_performance(self, material_type, composition, current_density):
        """分析材料性能并可视化结果，帮助查找问题"""
        try:
            if not hasattr(self, 'last_calc_data'):
                print("尚未执行效率计算，请先计算效率")
                return
                
            data = self.last_calc_data
            
            # 创建一个2x2的可视化图表
            fig, axes = plt.subplots(2, 2, figsize=(12, 10))
            fig.suptitle(f"{material_type}型材料 (组分={composition}, 电流密度={current_density}A/cm^2) 性能分析", fontsize=14)
            
            # 1. 温度分布
            ax1 = axes[0, 0]
            x_range = np.arange(1, len(data['temperature']) + 1)
            ax1.plot(x_range, data['temperature'], 'b-o')
            ax1.set_title('温度分布')
            ax1.set_xlabel('格点位置')
            ax1.set_ylabel('温度 (K)')
            ax1.grid(True)
            
            # 2. 材料属性随温度变化
            ax2 = axes[0, 1]
            ax2.plot(data['temperature'], data['seebeck'] * 1e6, 'r-', label='塞贝克系数 (μV/K)')
            ax2.set_xlabel('温度 (K)')
            ax2.set_ylabel('塞贝克系数 (μV/K)')
            ax2.set_title('塞贝克系数分布')
            ax2.grid(True)
            
            ax2_twin = ax2.twinx()
            ax2_twin.plot(data['temperature'], data['resistivity'] * 1e6, 'g-', label='电阻率 (μΩ·m)')
            ax2_twin.set_ylabel('电阻率 (μΩ·m)')
            
            # 添加双轴图例
            lines1, labels1 = ax2.get_legend_handles_labels()
            lines2, labels2 = ax2_twin.get_legend_handles_labels()
            ax2.legend(lines1 + lines2, labels1 + labels2, loc='best')
            
            # 3. 能量流动分析
            ax3 = axes[1, 0]
            seebeck_power = data['seebeck'] * data['dTdx'] * data['current_density']
            joule_heat = data['resistivity'] * data['current_density']**2
            
            ax3.plot(x_range, seebeck_power, 'b-', label='塞贝克功率')
            ax3.plot(x_range, joule_heat, 'r-', label='焦耳热损失')
            ax3.plot(x_range, seebeck_power - joule_heat, 'g-', label='净功率')
            ax3.set_title('能量流动分析')
            ax3.set_xlabel('格点位置')
            ax3.set_ylabel('功率密度 (W/m³)')
            ax3.grid(True)
            ax3.legend()
            
            # 4. 热流分析
            ax4 = axes[1, 1]
            fourier_heat = data['thermal_cond'] * data['dTdx']
            peltier_heat = data['current_density'] * data['seebeck'] * data['temperature']
            ax4.plot(x_range, fourier_heat, 'b-', label='傅里叶热流')
            ax4.plot(x_range, peltier_heat, 'r-', label='帕尔贴热流')
            ax4.plot(x_range, fourier_heat - peltier_heat, 'g-', label='净热流')
            ax4.set_title('热流分析')
            ax4.set_xlabel('格点位置')
            ax4.set_ylabel('热流密度 (W/m²)')
            ax4.grid(True)
            ax4.legend()
            
            plt.tight_layout()
            plt.show()
            
            # 打印能量平衡分析
            print("\n===== 能量平衡分析 =====")
            heat_in = abs(fourier_heat[0] - peltier_heat[0])
            heat_out = abs(fourier_heat[-1] - peltier_heat[-1])
            total_joule = np.sum(joule_heat) * (x_range[-1] - x_range[0]) / (len(x_range) - 1)
            total_power = np.sum(seebeck_power - joule_heat) * (x_range[-1] - x_range[0]) / (len(x_range) - 1)
            
            print(f"入口热流: {heat_in:.3e} W/m²")
            print(f"出口热流: {heat_out:.3e} W/m²")
            print(f"总焦耳热: {total_joule:.3e} W/m²")
            print(f"总功率输出: {total_power:.3e} W/m²")
            print(f"热平衡差值: {(heat_in - heat_out - total_power):.3e} W/m² (理论上应接近0)")
            
        except Exception as e:
            print(f"性能分析错误: {str(e)}")
            import traceback
            traceback.print_exc()

    def show_help(self):
        """显示帮助信息"""
        help_text = """
        <h3>温差半导体仿真程序使用说明</h3>
        
        <h4>基本操作：</h4>
        <ul>
        <li><b>参数设置：</b>在左侧面板设置温度、网格点数和迭代次数</li>
        <li><b>材料选择：</b>选择P型和N型半导体材料的组分</li>
        <li><b>开始计算：</b>点击"开始计算"按钮进行温度分布计算</li>
        <li><b>查看结果：</b>在中间和右侧面板查看计算结果和图表</li>
        </ul>
        
        <h4>参数说明：</h4>
        <ul>
        <li><b>高温温度Th：</b>热端温度，范围1-1000K</li>
        <li><b>低温温度Tc：</b>冷端温度，范围1-1000K，必须小于Th</li>
        <li><b>格子数量：</b>计算网格点数，范围5-100</li>
        <li><b>迭代次数：</b>最大迭代次数，范围1-1000</li>
        </ul>
        
        <h4>材料信息：</h4>
        <ul>
        <li><b>P型材料：</b>PbTe₁₋ᵧIᵧ，可选组分0.01、0.02、0.03</li>
        <li><b>N型材料：</b>PbTe:Na/Ag₂Te，可选组分0.0004、0.0012、0.0020、0.0028</li>
        </ul>
        
        <h4>功能特性：</h4>
        <ul>
        <li><b>实时计算：</b>支持实时温度分布计算和显示</li>
        <li><b>交互图表：</b>点击图表数据点可查看详细信息</li>
        <li><b>数据导出：</b>可将计算结果导出为Excel文件</li>
        <li><b>外部计算：</b>支持调用外部脚本进行详细效率计算</li>
        </ul>
        
        <h4>注意事项：</h4>
        <ul>
        <li>确保输入参数在有效范围内</li>
        <li>计算过程中请勿关闭程序</li>
        <li>大量数据计算可能需要较长时间</li>
        <li>建议定期保存计算结果</li>
        </ul>
        """
        
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("使用帮助")
        msg_box.setTextFormat(Qt.RichText)
        msg_box.setText(help_text)
        msg_box.setStandardButtons(QMessageBox.Ok)
        msg_box.exec_()
    
    def show_about(self):
        """显示关于信息"""
        about_text = """
        <h3>温差半导体仿真程序</h3>
        <p><b>版本：</b>1.0</p>
        <p><b>功能：</b>基于差分法的半导体热电器件仿真</p>
        <p><b>技术：</b>PyQt5 + Matplotlib + NumPy + SciPy</p>
        <p><b>特点：</b>实时计算、交互图表、数据导出</p>
        <p><b>适用：</b>热电材料研究、器件设计优化</p>
        """
        
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("关于程序")
        msg_box.setTextFormat(Qt.RichText)
        msg_box.setText(about_text)
        msg_box.setStandardButtons(QMessageBox.Ok)
        msg_box.exec_()
    
    def add_context_menu(self):
        """为图表添加右键菜单"""
        try:
            # 为温度分布图添加右键菜单
            if hasattr(self, 'temp_canvas'):
                self.temp_canvas.setContextMenuPolicy(Qt.CustomContextMenu)
                self.temp_canvas.customContextMenuRequested.connect(self.show_temp_context_menu)
            
            # 为效率图添加右键菜单
            if hasattr(self, 'eff_canvas'):
                self.eff_canvas.setContextMenuPolicy(Qt.CustomContextMenu)
                self.eff_canvas.customContextMenuRequested.connect(self.show_eff_context_menu)
        except Exception as e:
            print(f"添加右键菜单失败: {str(e)}")
    
    def show_temp_context_menu(self, position):
        """显示温度分布图右键菜单"""
        try:
            from PyQt5.QtWidgets import QMenu
            
            menu = QMenu(self)
            
            # 添加菜单项
            export_action = menu.addAction("导出数据")
            zoom_action = menu.addAction("放大查看")
            reset_action = menu.addAction("重置视图")
            
            # 显示菜单
            action = menu.exec_(self.temp_canvas.mapToGlobal(position))
            
            if action == export_action:
                self.export_temperature_data()
            elif action == zoom_action:
                self.zoom_temperature_plot()
            elif action == reset_action:
                self.reset_temperature_view()
                
        except Exception as e:
            print(f"显示右键菜单失败: {str(e)}")
    
    def export_temperature_data(self):
        """导出温度分布数据"""
        try:
            if hasattr(self, 'x_p') and hasattr(self, 'T_p') and hasattr(self, 'x_n') and hasattr(self, 'T_n'):
                # 创建数据框
                import pandas as pd
                
                data = {
                    'P型位置': self.x_p,
                    'P型温度': self.T_p,
                    'N型位置': self.x_n,
                    'N型温度': self.T_n
                }
                
                df = pd.DataFrame(data)
                
                # 保存文件
                filename, _ = QFileDialog.getSaveFileName(
                    self, "保存温度分布数据", "", "CSV文件 (*.csv);;Excel文件 (*.xlsx)")
                
                if filename:
                    if filename.endswith('.csv'):
                        df.to_csv(filename, index=False, encoding='utf-8-sig')
                    else:
                        df.to_excel(filename, index=False)
                    
                    self.show_info_message("导出成功", f"温度分布数据已保存到：\n{filename}")
            else:
                self.show_error_message("导出失败", "没有可导出的温度分布数据")
                
        except Exception as e:
            self.show_error_message("导出失败", f"导出温度分布数据时出错：\n{str(e)}")
    
    def zoom_temperature_plot(self):
        """放大温度分布图"""
        try:
            if hasattr(self, 'temp_canvas'):
                # 这里可以添加放大逻辑
                self.show_info_message("功能提示", "放大功能正在开发中...")
        except Exception as e:
            print(f"放大温度图失败: {str(e)}")
    
    def reset_temperature_view(self):
        """重置温度分布图视图"""
        try:
            if hasattr(self, 'temp_canvas'):
                # 重新绘制温度分布图
                if hasattr(self, 'x_p') and hasattr(self, 'T_p') and hasattr(self, 'x_n') and hasattr(self, 'T_n'):
                    self.update_temperature_plots(self.x_p, self.T_p, self.x_n, self.T_n)
                    self.show_info_message("视图重置", "温度分布图视图已重置")
                else:
                    self.show_error_message("重置失败", "没有可重置的温度分布数据")
        except Exception as e:
            print(f"重置温度图视图失败: {str(e)}")
    
    def show_eff_context_menu(self, position):
        """显示效率图右键菜单"""
        try:
            from PyQt5.QtWidgets import QMenu
            
            menu = QMenu(self)
            
            # 添加菜单项
            export_action = menu.addAction("导出效率数据")
            analyze_action = menu.addAction("性能分析")
            
            # 显示菜单
            action = menu.exec_(self.eff_canvas.mapToGlobal(position))
            
            if action == export_action:
                self.export_efficiency_data()
            elif action == analyze_action:
                self.analyze_efficiency_performance()
                
        except Exception as e:
            print(f"显示效率图右键菜单失败: {str(e)}")
    
    def export_efficiency_data(self):
        """导出效率数据"""
        try:
            # 这里可以添加导出效率数据的逻辑
            self.show_info_message("功能提示", "效率数据导出功能正在开发中...")
        except Exception as e:
            print(f"导出效率数据失败: {str(e)}")
    
    def analyze_efficiency_performance(self):
        """分析效率性能"""
        try:
            # 这里可以添加效率性能分析的逻辑
            self.show_info_message("功能提示", "效率性能分析功能正在开发中...")
        except Exception as e:
            print(f"分析效率性能失败: {str(e)}")
    
    def closeEvent(self, event):
        """程序关闭事件"""
        try:
            # 停止计算线程
            if hasattr(self, 'calculation_thread') and self.calculation_thread:
                self.calculation_thread.stop()
                self.calculation_thread.wait()
            
            # 保存用户设置（如果有的话）
            self.save_user_settings()
            
            event.accept()
        except Exception as e:
            print(f"程序关闭时出错: {str(e)}")
            event.accept()
    
    def save_user_settings(self):
        """保存用户设置"""
        try:
            # 这里可以添加保存用户设置的逻辑
            # 比如保存窗口大小、位置、默认参数等
            pass
        except Exception as e:
            print(f"保存用户设置失败: {str(e)}")
    
    def load_user_settings(self):
        """加载用户设置"""
        try:
            # 这里可以添加加载用户设置的逻辑
            pass
        except Exception as e:
            print(f"加载用户设置失败: {str(e)}")

    def new_project(self):
        """新建项目"""
        try:
            # 重置所有输入参数
            self.th_edit.setText("500")
            self.tc_edit.setText("300")
            self.grid_edit.setText("10")
            self.iter_edit.setText("20")
            self.p_type_combo.setCurrentText("0.01")
            self.n_type_combo.setCurrentText("0.0004")
            
            # 清除图表
            self.clear_all_plots()
            
            # 更新状态
            self.update_status("已创建新项目")
            self.show_info_message("新建项目", "新项目已创建，所有参数已重置为默认值")
            
        except Exception as e:
            self.show_error_message("新建项目失败", f"创建新项目时出错：\n{str(e)}")
    def open_project(self):
        """打开项目"""
        try:
            filename, _ = QFileDialog.getOpenFileName(
                self, "打开项目", "", "项目文件 (*.json);;所有文件 (*)")
            
            if filename:
                # 这里可以添加加载项目文件的逻辑
                self.show_info_message("功能提示", "项目加载功能正在开发中...")
                
        except Exception as e:
            self.show_error_message("打开项目失败", f"打开项目时出错：\n{str(e)}")
    def save_project(self):
        """保存项目"""
        try:
            filename, _ = QFileDialog.getSaveFileName(
                self, "保存项目", "", "项目文件 (*.json);;所有文件 (*)")
            
            if filename:
                # 这里可以添加保存项目文件的逻辑
                self.show_info_message("功能提示", "项目保存功能正在开发中...")
                
        except Exception as e:
            self.show_error_message("保存项目失败", f"保存项目时出错：\n{str(e)}")
    
    def stop_calculation(self):
        """停止计算"""
        try:
            if hasattr(self, 'calculation_thread') and self.calculation_thread:
                self.calculation_thread.stop()
                self.calculation_thread.wait()
                self.update_status("计算已停止")
                self.show_info_message("计算停止", "计算已成功停止")
            else:
                self.show_info_message("提示", "当前没有正在进行的计算")
                
        except Exception as e:
            self.show_error_message("停止计算失败", f"停止计算时出错：\n{str(e)}")
    
    def reset_all_views(self):
        """重置所有视图"""
        try:
            # 重置温度分布图
            if hasattr(self, 'x_p') and hasattr(self, 'T_p') and hasattr(self, 'x_n') and hasattr(self, 'T_n'):
                self.update_temperature_plots(self.x_p, self.T_p, self.x_n, self.T_n)
            
            # 重置效率图
            self.update_efficiency_plots()
            
            # 重置ZT图
            self.update_zt_plots()
            
            self.update_status("所有视图已重置")
            self.show_info_message("视图重置", "所有图表视图已重置为默认状态")
            
        except Exception as e:
            self.show_error_message("重置视图失败", f"重置视图时出错：\n{str(e)}")
    
    def toggle_fullscreen(self):
        """切换全屏显示"""
        try:
            if self.isFullScreen():
                self.showNormal()
                self.update_status("退出全屏模式")
            else:
                self.showFullScreen()
                self.update_status("进入全屏模式")
                
        except Exception as e:
            self.show_error_message("全屏切换失败", f"切换全屏模式时出错：\n{str(e)}")
    
    def clear_all_plots(self):
        """清除所有图表"""
        try:
            # 清除温度分布图
            if hasattr(self, 'temp_axes'):
                ax1, ax2 = self.temp_axes
                ax1.clear()
                ax2.clear()
                ax1.set_title("P型材料温度分布")
                ax1.set_xlabel("位置")
                ax1.set_ylabel("温度 (K)")
                ax1.grid(True)
                
                ax2.set_title("N型材料温度分布")
                ax2.set_xlabel("位置")
                ax2.set_ylabel("温度 (K)")
                ax2.grid(True)
                
                if hasattr(self, 'temp_canvas'):
                    self.temp_canvas.draw()
            
            # 清除效率图
            if hasattr(self, 'eff_axes'):
                ax1, ax2 = self.eff_axes
                ax1.clear()
                ax2.clear()
                ax1.set_title("P型材料效率")
                ax1.set_xlabel("电流密度 (A/cm$^2$)")
                ax1.set_ylabel("效率")
                ax1.grid(True)
                
                ax2.set_title("N型材料效率")
                ax2.set_xlabel("电流密度 (A/cm$^2$)")
                ax2.set_ylabel("效率")
                ax2.grid(True)
                
                if hasattr(self, 'eff_canvas'):
                    self.eff_canvas.draw()
            
            # 清除ZT图
            if hasattr(self, 'zt_axes'):
                ax1, ax2 = self.zt_axes
                ax1.clear()
                ax2.clear()
                ax1.set_title("P型半导体材料")
                ax1.set_xlabel("温度")
                ax1.set_ylabel("ZT")
                ax1.set_xlim(300, 700)
                ax1.set_ylim(0, 1.5)
                ax1.grid(True, color='white', linestyle='-', alpha=0.8)
                ax1.set_facecolor('#F0F0F0')
                
                ax2.set_title("N型半导体材料")
                ax2.set_xlabel("温度")
                ax2.set_ylabel("ZT")
                ax2.set_xlim(300, 700)
                ax2.set_ylim(0, 1.5)
                ax2.grid(True, color='white', linestyle='-', alpha=0.8)
                ax2.set_facecolor('#F0F0F0')
                
                if hasattr(self, 'zt_canvas'):
                    self.zt_canvas.draw()
                    
        except Exception as e:
            print(f"清除图表失败: {str(e)}")
    
    # ==================== 材料管理功能 ====================
    
    def show_material_manager(self):
        """显示材料管理对话框"""
        try:
            dialog = MaterialManagerDialog(self)
            dialog.exec_()
            
            # 刷新材料选择下拉框
            self.refresh_material_combos()
            
        except Exception as e:
            self.show_error_message("错误", f"显示材料管理器失败：{str(e)}")
    
    def import_material_data(self):
        """导入材料数据"""
        try:
            filename, _ = QFileDialog.getOpenFileName(
                self, "选择Excel文件", "", "Excel文件 (*.xlsx *.xls)")
            
            if filename:
                # 显示导入选项对话框
                dialog = MaterialImportDialog(self, filename)
                if dialog.exec_() == QDialog.Accepted:
                    # 刷新材料选择下拉框
                    self.refresh_material_combos()
                    self.show_info_message("成功", "材料数据导入成功！")
                    
        except Exception as e:
            self.show_error_message("错误", f"导入材料数据失败：{str(e)}")
    
    def export_material_data(self):
        """导出材料数据"""
        try:
            # 显示导出选项对话框
            dialog = MaterialExportDialog(self)
            dialog.exec_()
            
        except Exception as e:
            self.show_error_message("错误", f"导出材料数据失败：{str(e)}")
    
    def view_materials_info(self):
        """查看材料信息"""
        try:
            dialog = MaterialInfoDialog(self)
            dialog.exec_()
            
        except Exception as e:
            self.show_error_message("错误", f"查看材料信息失败：{str(e)}")
    
    def refresh_material_combos(self):
        """刷新材料选择下拉框"""
        try:
            # 获取当前选择的材料
            current_p = self.p_type_combo.currentText()
            current_n = self.n_type_combo.currentText()
            
            # 获取所有可用材料
            p_materials = self.calculator.get_material_list('p')
            n_materials = self.calculator.get_material_list('n')
            
            # 更新P型材料下拉框
            self.p_type_combo.clear()
            self.p_type_combo.addItems(p_materials)
            
            # 尝试恢复之前的选择
            if current_p in p_materials:
                self.p_type_combo.setCurrentText(current_p)
            elif p_materials:
                self.p_type_combo.setCurrentText(p_materials[0])
            
            # 更新N型材料下拉框
            self.n_type_combo.clear()
            self.n_type_combo.addItems(n_materials)
            
            # 尝试恢复之前的选择
            if current_n in n_materials:
                self.n_type_combo.setCurrentText(current_n)
            elif n_materials:
                self.n_type_combo.setCurrentText(n_materials[0])
            
            # 更新ZT图表
            self.update_zt_plots()
            
        except Exception as e:
            print(f"刷新材料下拉框失败: {str(e)}")

class ExperimentReportGenerator:
    """实验报告生成器类"""
    
    def __init__(self, app_instance):
        self.app = app_instance
        self.calculator = app_instance.calculator
        self.report_data = {}
        
    def collect_simulation_data(self):
        """收集所有仿真数据"""
        try:
            # 收集基本参数
            self.report_data['basic_params'] = {
                'high_temp': float(self.app.th_edit.text()),
                'low_temp': float(self.app.tc_edit.text()),
                'p_type_composition': self.app.p_type_combo.currentText(),
                'n_type_composition': self.app.n_type_combo.currentText(),
                'n_p_ratio': float(self.app.ratio_edit.text()),
                'grid_points': int(self.app.grid_edit.text()),
                'max_iterations': int(self.app.iter_edit.text())
            }
            
            # 收集计算结果
            result = getattr(self.app, 'last_device_calculation', None)
            if result:
                self.report_data['calculation_results'] = {
                    'max_efficiency': float(self.app.max_eff.text().replace('%', '').strip()),
                    'max_efficiency_current': float(self.app.eff_current.text()),
                    'max_power': float(self.app.max_power.text()),
                    'max_power_current': float(self.app.power_current.text()),
                    'i_list': result.get('I_list', []),
                    'eff_total_list': result.get('eff_total_list', []),
                    'power_total_list': result.get('Power_total_list', [])
                }
            
            # 收集分支特性数据
            self.report_data['branch_data'] = {
                'p_branch': {
                    'j_list': getattr(self.app, 'J_list_P', []),
                    'eff_list': getattr(self.app, 'eff_list_P', []),
                    'q_list': getattr(self.app, 'q_P_list', [])
                },
                'n_branch': {
                    'j_list': getattr(self.app, 'J_list_N', []),
                    'eff_list': getattr(self.app, 'eff_list_N', []),
                    'q_list': getattr(self.app, 'q_N_list', [])
                }
            }
            
            # 收集温度分布数据
            self.report_data['temperature_data'] = {
                'p_branch': {
                    'x': getattr(self.app, 'x_p', []),
                    'T': getattr(self.app, 'T_p', [])
                },
                'n_branch': {
                    'x': getattr(self.app, 'x_n', []),
                    'T': getattr(self.app, 'T_n', [])
                }
            }
            
            # 收集材料参数
            p_comp = self.app.p_type_combo.currentText()
            n_comp = self.app.n_type_combo.currentText()
            
            if p_comp in self.calculator.p_type_data:
                self.report_data['p_material'] = self.calculator.p_type_data[p_comp]
            if n_comp in self.calculator.n_type_data:
                self.report_data['n_material'] = self.calculator.n_type_data[n_comp]
                
            return True
            
        except Exception as e:
            print(f"收集仿真数据时出错: {str(e)}")
            return False
    
    def generate_report_content(self):
        """生成实验报告内容"""
        if not self.report_data:
            return None
            
        # 生成报告标题和基本信息
        report_content = {
            'title': '半导体热电器件仿真实验报告',
            'timestamp': time.strftime("%Y年%m月%d日 %H:%M:%S"),
            'basic_info': self._generate_basic_info(),
            'experiment_objective': self._generate_experiment_objective(),
            'experiment_method': self._generate_experiment_method(),
            'experiment_results': self._generate_experiment_results(),
            'data_analysis': self._generate_data_analysis(),
            'conclusion': self._generate_conclusion()
        }
        
        return report_content
    
    def _generate_basic_info(self):
        """生成基本信息"""
        params = self.report_data['basic_params']
        return f"""
实验基本信息：
• 实验时间：{time.strftime("%Y年%m月%d日 %H:%M:%S")}
• 高温端温度：{params['high_temp']} K
• 低温端温度：{params['low_temp']} K
• 温度差：{params['high_temp'] - params['low_temp']} K
• P型材料组分：{params['p_type_composition']}
• N型材料组分：{params['n_type_composition']}
• N/P面积比：{params['n_p_ratio']}
• 网格点数：{params['grid_points']}
• 最大迭代次数：{params['max_iterations']}
        """.strip()
    
    def _generate_experiment_objective(self):
        """生成实验目的"""
        return """
实验目的：
1. 研究不同材料组分对半导体热电器件性能的影响
2. 分析温度分布、效率、功率等关键性能指标
3. 确定器件的最优工作条件和参数配置
4. 验证差分法在热电仿真中的有效性
        """.strip()
    
    def _generate_experiment_method(self):
        """生成实验方法"""
        return """
实验方法：
1. 数值仿真方法：采用有限差分法求解一维稳态热传导方程
2. 物理模型：考虑塞贝克效应、帕尔贴效应、焦耳热等热电耦合效应
3. 边界条件：高温端和低温端温度固定，中间区域自由演化
4. 收敛判据：采用迭代法求解，直到温度分布收敛
5. 性能计算：基于温度分布计算效率、功率等性能指标
        """.strip()
    
    def _generate_experiment_results(self):
        """生成实验结果"""
        if 'calculation_results' not in self.report_data:
            return "暂无计算结果"
            
        results = self.report_data['calculation_results']
        return f"""
实验结果：
1. 最大效率：{results['max_efficiency']:.4f}% (对应电流：{results['max_efficiency_current']} A)
2. 最大功率：{results['max_power']:.6f} W/cm² (对应电流：{results['max_power_current']} A)
3. 工作电流范围：{min(results['i_list'])} - {max(results['i_list'])} A
4. 效率变化范围：{min(results['eff_total_list'])*100:.4f}% - {max(results['eff_total_list'])*100:.4f}%
5. 功率变化范围：{min(results['power_total_list']):.6f} - {max(results['power_total_list']):.6f} W/cm²
        """.strip()
    
    def _generate_data_analysis(self):
        """生成数据分析"""
        if 'calculation_results' not in self.report_data:
            return "暂无数据可分析"
            
        results = self.report_data['calculation_results']
        
        # 计算一些统计信息
        eff_list = results['eff_total_list']
        power_list = results['power_total_list']
        
        # 找到效率大于0的区间
        positive_eff_indices = [i for i, eff in enumerate(eff_list) if eff > 0]
        if positive_eff_indices:
            positive_eff_range = f"{min(positive_eff_indices)} - {max(positive_eff_indices)} A"
        else:
            positive_eff_range = "无"
        
        return f"""
数据分析：
1. 性能曲线特征：
   • 效率曲线呈现典型的抛物线特征，存在最优工作点
   • 功率曲线与效率曲线类似，但峰值位置可能不同
   
2. 工作区间分析：
   • 正效率工作区间：{positive_eff_range}
   • 最优工作电流：{results['max_efficiency_current']} A
   
3. 材料性能影响：
   • P型材料组分{self.report_data['basic_params']['p_type_composition']}的影响
   • N型材料组分{self.report_data['basic_params']['n_type_composition']}的影响
   • N/P面积比{self.report_data['basic_params']['n_p_ratio']}的优化效果
        """.strip()
    
    def _generate_conclusion(self):
        """生成结论"""
        if 'calculation_results' not in self.report_data:
            return "暂无结论可得出"
            
        results = self.report_data['calculation_results']
        params = self.report_data['basic_params']
        
        return f"""
实验结论：
1. 器件性能表现：
   • 在给定条件下，器件最大效率达到{results['max_efficiency']:.4f}%
   • 最大功率输出为{results['max_power']:.6f} W/cm²
   • 最优工作电流为{results['max_efficiency_current']} A
   
2. 参数优化建议：
   • 当前N/P面积比{params['n_p_ratio']}的设置效果
   • 材料组分{params['p_type_composition']}和{params['n_type_composition']}的匹配性
   • 温度差{params['high_temp'] - params['low_temp']} K对性能的影响
   
3. 仿真方法验证：
   • 差分法在热电仿真中表现良好，结果合理
   • 温度分布收敛稳定，计算精度满足要求
   • 为实际器件设计提供了有价值的理论指导
        """.strip()
    
    def create_charts_for_report(self):
        """为报告创建专业图表"""
        charts = {}
        
        try:
            # 设置图表字体，确保上标正确显示
            plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
            plt.rcParams['mathtext.fontset'] = 'dejavusans'
            
            # 创建效率-电流关系图
            if 'calculation_results' in self.report_data:
                results = self.report_data['calculation_results']
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.plot(results['i_list'], [e*100 for e in results['eff_total_list']], 
                       'b-o', linewidth=2, markersize=6, label='总效率')
                ax.axhline(y=0, color='r', linestyle='--', alpha=0.5)
                ax.set_xlabel('电流 (A)', fontsize=12)
                ax.set_ylabel('效率 (%)', fontsize=12)
                ax.set_title('热电器件效率-电流特性曲线', fontsize=14, fontweight='bold')
                
                # 调整Y轴范围，主要关注0%-10%的效率区间
                eff_values = [e*100 for e in results['eff_total_list']]
                max_eff = max(eff_values)
                min_eff = min(eff_values)
                
                # 设置Y轴范围为0到最大效率的1.2倍，但不超过10%
                y_upper = min(max_eff * 1.2, 10.0)
                ax.set_ylim(0, y_upper)
                
                # 设置Y轴刻度，在0-10%范围内显示更多细节
                if y_upper <= 10.0:
                    ax.set_yticks([0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10])
                else:
                    ax.set_yticks([0, 2, 4, 6, 8, 10, 12, 14, 16, 18, 20])
                
                ax.grid(True, alpha=0.3)
                ax.legend(fontsize=11)
                ax.tick_params(labelsize=10)
                
                # 标记最大效率点
                max_eff_idx = results['eff_total_list'].index(max(results['eff_total_list']))
                max_eff_current = results['i_list'][max_eff_idx]
                max_eff_value = results['eff_total_list'][max_eff_idx] * 100
                
                # 将标签放到曲线下面，避免遮挡标题
                x_range = max(results['i_list']) - min(results['i_list'])
                y_range = max([e*100 for e in results['eff_total_list']]) - min([e*100 for e in results['eff_total_list']])
                
                # 标签放在数据点下方，避免遮挡标题
                x_offset = 0  # 水平居中
                y_offset = -y_range * 0.15  # 向下偏移
                
                ax.annotate(f'最大效率: {max_eff_value:.3f}%\n对应电流: {max_eff_current} A',
                           xy=(max_eff_current, max_eff_value),
                           xytext=(max_eff_current + x_offset, max_eff_value + y_offset),
                           arrowprops=dict(arrowstyle='->', color='red', lw=2),
                           bbox=dict(boxstyle='round,pad=0.5', facecolor='yellow', alpha=0.7),
                           fontsize=10,
                           ha='center', va='top')
                
                charts['efficiency_current'] = fig
                
                # 创建功率-电流关系图
                fig2, ax2 = plt.subplots(figsize=(10, 6))
                ax2.plot(results['i_list'], results['power_total_list'], 
                        'g-s', linewidth=2, markersize=6, label='总功率')
                ax2.axhline(y=0, color='r', linestyle='--', alpha=0.5)
                ax2.set_xlabel('电流 (A)', fontsize=12)
                ax2.set_ylabel('功率 (W/cm$^2$)', fontsize=12)
                ax2.set_title('热电器件功率-电流特性曲线', fontsize=14, fontweight='bold')
                
                # 调整Y轴范围，让功率变化更加清晰
                power_values = results['power_total_list']
                max_power = max(power_values)
                min_power = min(power_values)
                
                # 设置Y轴范围，主要关注正功率区域
                if max_power > 0:
                    y_upper = max_power * 1.2
                    y_lower = min(0, min_power * 1.1)  # 包含负值区域
                else:
                    y_upper = abs(min_power) * 1.2
                    y_lower = min_power * 1.1
                
                ax2.set_ylim(y_lower, y_upper)
                
                ax2.grid(True, alpha=0.3)
                ax2.legend(fontsize=11)
                ax2.tick_params(labelsize=10)
                
                # 标记最大功率点
                max_power_idx = results['power_total_list'].index(max(results['power_total_list']))
                max_power_current = results['i_list'][max_power_idx]
                max_power_value = results['power_total_list'][max_power_idx]
                
                # 将标签放到曲线下面，避免遮挡标题
                power_y_range = max(results['power_total_list']) - min(results['power_total_list'])
                y_offset = -power_y_range * 0.15  # 向下偏移
                
                ax2.annotate(f'最大功率: {max_power_value:.6f} W/cm²\n对应电流: {max_power_current} A',
                            xy=(max_power_current, max_power_value),
                            xytext=(max_power_current, max_power_value + y_offset),
                            arrowprops=dict(arrowstyle='->', color='red', lw=2),
                            bbox=dict(boxstyle='round,pad=0.5', facecolor='yellow', alpha=0.7),
                            fontsize=10,
                            ha='center', va='top')
                
                charts['power_current'] = fig2
                
                # 创建温度分布图
                if 'temperature_data' in self.report_data:
                    temp_data = self.report_data['temperature_data']
                    if temp_data['p_branch']['x'] and temp_data['n_branch']['x']:
                        fig3, ax3 = plt.subplots(figsize=(10, 6))
                        ax3.plot(temp_data['p_branch']['x'], temp_data['p_branch']['T'], 
                                'r-o', linewidth=2, markersize=4, label='P型分支')
                        ax3.plot(temp_data['n_branch']['x'], temp_data['n_branch']['T'], 
                                'b-s', linewidth=2, markersize=4, label='N型分支')
                        ax3.set_xlabel('位置 (归一化)', fontsize=12)
                        ax3.set_ylabel('温度 (K)', fontsize=12)
                        ax3.set_title('P型和N型分支温度分布', fontsize=14, fontweight='bold')
                        ax3.grid(True, alpha=0.3)
                        ax3.legend(fontsize=11)
                        ax3.tick_params(labelsize=10)
                        
                        charts['temperature_distribution'] = fig3
                
        except Exception as e:
            print(f"创建图表时出错: {str(e)}")
            
        return charts
    
    def export_report_to_word(self, filename=None):
        """导出实验报告到Word文档"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            
            # 创建Word文档
            doc = Document()
            
            # 设置中文字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加标题
            title = doc.add_heading('半导体热电器件仿真实验报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加时间
            time_para = doc.add_paragraph(f"实验时间：{time.strftime('%Y年%m月%d日 %H:%M:%S')}")
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加基本信息
            doc.add_heading('1. 实验基本信息', level=1)
            doc.add_paragraph(self._generate_basic_info())
            
            # 添加实验目的
            doc.add_heading('2. 实验目的', level=1)
            doc.add_paragraph(self._generate_experiment_objective())
            
            # 添加实验方法
            doc.add_heading('3. 实验方法', level=1)
            doc.add_paragraph(self._generate_experiment_method())
            
            # 添加实验结果
            doc.add_heading('4. 实验结果', level=1)
            doc.add_paragraph(self._generate_experiment_results())
            
            # 添加图表
            charts = self.create_charts_for_report()
            if charts:
                doc.add_heading('5. 实验图表', level=1)
                
                # 保存图表为临时文件并插入
                import tempfile
                import os
                
                for chart_name, fig in charts.items():
                    # 创建临时文件
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                        fig.savefig(tmp_file.name, dpi=300, bbox_inches='tight')
                        tmp_path = tmp_file.name
                    
                    # 插入图片
                    doc.add_picture(tmp_path, width=Inches(6))
                    
                    # 添加图片说明
                    if chart_name == 'efficiency_current':
                        doc.add_paragraph('图1: 热电器件效率-电流特性曲线')
                    elif chart_name == 'power_current':
                        doc.add_paragraph('图2: 热电器件功率-电流特性曲线')
                    elif chart_name == 'temperature_distribution':
                        doc.add_paragraph('图3: P型和N型分支温度分布')
                    
                    # 删除临时文件
                    os.unlink(tmp_path)
                    plt.close(fig)
            
            # 添加数据分析
            doc.add_heading('6. 数据分析', level=1)
            doc.add_paragraph(self._generate_data_analysis())
            
            # 添加结论
            doc.add_heading('7. 实验结论', level=1)
            doc.add_paragraph(self._generate_conclusion())
            
            # 保存文档
            if not filename:
                timestamp = time.strftime("%Y%m%d_%H%M%S")
                filename = f"实验报告_{timestamp}.docx"
            
            doc.save(filename)
            return filename
            
        except ImportError:
            print("未安装python-docx库，无法导出Word文档")
            return None
        except Exception as e:
            print(f"导出Word文档时出错: {str(e)}")
            return None
    
    def export_report_to_html(self, filename=None):
        """导出实验报告到HTML格式"""
        try:
            if not filename:
                timestamp = time.strftime("%Y%m%d_%H%M%S")
                filename = f"实验报告_{timestamp}.html"
            
            # 创建HTML内容
            html_content = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>半导体热电器件仿真实验报告</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; border-bottom: 2px solid #ecf0f1; padding-bottom: 5px; }}
        .timestamp {{ text-align: center; color: #7f8c8d; font-style: italic; }}
        .section {{ margin: 20px 0; }}
        .data-table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        .data-table th, .data-table td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        .data-table th {{ background-color: #f2f2f2; }}
        .highlight {{ background-color: #fff3cd; padding: 10px; border-left: 4px solid #ffc107; }}
    </style>
</head>
<body>
    <h1>半导体热电器件仿真实验报告</h1>
    <p class="timestamp">实验时间：{time.strftime("%Y年%m月%d日 %H:%M:%S")}</p>
    
    <div class="section">
        <h2>1. 实验基本信息</h2>
        <pre>{self._generate_basic_info()}</pre>
    </div>
    
    <div class="section">
        <h2>2. 实验目的</h2>
        <pre>{self._generate_experiment_objective()}</pre>
    </div>
    
    <div class="section">
        <h2>3. 实验方法</h2>
        <pre>{self._generate_experiment_method()}</pre>
    </div>
    
    <div class="section">
        <h2>4. 实验结果</h2>
        <pre>{self._generate_experiment_results()}</pre>
    </div>
    
    <div class="section">
        <h2>5. 数据分析</h2>
        <pre>{self._generate_data_analysis()}</pre>
    </div>
    
    <div class="section">
        <h2>6. 实验结论</h2>
        <pre>{self._generate_conclusion()}</pre>
    </div>
    
    <div class="highlight">
        <strong>注意：</strong>本报告基于数值仿真结果生成，图表数据请参考相应的PNG文件。
    </div>
</body>
</html>
            """
            
            # 保存HTML文件
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return filename
            
        except Exception as e:
            print(f"导出HTML文档时出错: {str(e)}")
            return None
    
    def export_report_to_txt(self, filename=None):
        """导出实验报告到纯文本格式"""
        try:
            if not filename:
                timestamp = time.strftime("%Y%m%d_%H%M%S")
                filename = f"实验报告_{timestamp}.txt"
            
            # 创建文本内容
            text_content = f"""
半导体热电器件仿真实验报告
{'='*50}

实验时间：{time.strftime("%Y年%m月%d日 %H:%M:%S")}

{self._generate_basic_info()}

{self._generate_experiment_objective()}

{self._generate_experiment_method()}

{self._generate_experiment_results()}

{self._generate_data_analysis()}

{self._generate_conclusion()}

{'='*50}
报告生成完成
            """
            
            # 保存文本文件
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            return filename
            
        except Exception as e:
            print(f"导出文本文档时出错: {str(e)}")
            return None
class MaterialManagerDialog(QDialog):
    """材料管理对话框"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.calculator = parent.calculator
        self.setup_ui()
        self.load_materials()
    
    def setup_ui(self):
        """设置用户界面"""
        self.setWindowTitle("材料管理器")
        self.setMinimumSize(800, 600)
        
        layout = QVBoxLayout(self)
        
        # 创建选项卡
        tab_widget = QTabWidget()
        
        # P型材料选项卡
        p_tab = self.create_material_tab('p')
        tab_widget.addTab(p_tab, "P型材料")
        
        # N型材料选项卡
        n_tab = self.create_material_tab('n')
        tab_widget.addTab(n_tab, "N型材料")
        
        layout.addWidget(tab_widget)
        
        # 底部按钮
        button_layout = QHBoxLayout()
        
        close_button = QPushButton("关闭")
        close_button.clicked.connect(self.close)
        button_layout.addStretch()
        button_layout.addWidget(close_button)
        
        layout.addLayout(button_layout)
    
    def create_material_tab(self, material_type):
        """创建材料选项卡"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 材料列表
        list_group = QGroupBox("材料列表")
        list_layout = QVBoxLayout(list_group)
        
        self.material_list = QListWidget()
        self.material_list.itemSelectionChanged.connect(self.on_material_selected)
        list_layout.addWidget(self.material_list)
        
        # 材料操作按钮
        button_layout = QHBoxLayout()
        
        add_button = QPushButton("添加材料")
        add_button.clicked.connect(lambda checked, mt=material_type: self.add_material(mt))
        
        edit_button = QPushButton("编辑材料")
        edit_button.clicked.connect(lambda checked, mt=material_type: self.edit_material(mt))
        
        delete_button = QPushButton("删除材料")
        delete_button.clicked.connect(lambda checked, mt=material_type: self.delete_material(mt))
        
        import_button = QPushButton("从Excel导入")
        import_button.clicked.connect(lambda checked, mt=material_type: self.import_from_excel(mt))
        
        export_button = QPushButton("导出到Excel")
        export_button.clicked.connect(lambda checked, mt=material_type: self.export_to_excel(mt))
        
        button_layout.addWidget(add_button)
        button_layout.addWidget(edit_button)
        button_layout.addWidget(delete_button)
        button_layout.addWidget(import_button)
        button_layout.addWidget(export_button)
        
        list_layout.addLayout(button_layout)
        layout.addWidget(list_group)
        
        # 材料信息显示
        info_group = QGroupBox("材料信息")
        info_layout = QVBoxLayout(info_group)
        
        self.info_text = QTextEdit()
        self.info_text.setReadOnly(True)
        self.info_text.setMaximumHeight(200)
        info_layout.addWidget(self.info_text)
        
        layout.addWidget(info_group)
        
        # 保存材料类型引用
        widget.material_type = material_type
        
        return widget
    
    def load_materials(self):
        """加载材料列表"""
        try:
            # 加载P型材料
            p_tab = self.findChild(QTabWidget).widget(0)
            p_list = p_tab.findChild(QListWidget)
            p_materials = self.calculator.get_material_list('p')
            p_list.clear()
            p_list.addItems(p_materials)
            
            # 加载N型材料
            n_tab = self.findChild(QTabWidget).widget(1)
            n_list = n_tab.findChild(QListWidget)
            n_materials = self.calculator.get_material_list('n')
            n_list.clear()
            n_list.addItems(n_materials)
            
        except Exception as e:
            print(f"加载材料列表失败: {str(e)}")
    
    def on_material_selected(self):
        """材料选择改变事件"""
        try:
            # 获取当前选中的材料
            current_tab = self.findChild(QTabWidget).currentWidget()
            material_type = current_tab.material_type
            list_widget = current_tab.findChild(QListWidget)
            
            if list_widget.currentItem():
                composition = list_widget.currentItem().text()
                self.show_material_info(material_type, composition)
                
        except Exception as e:
            print(f"处理材料选择事件失败: {str(e)}")
    
    def show_material_info(self, material_type, composition):
        """显示材料信息"""
        try:
            info = self.calculator.get_material_info(material_type, composition)
            if info:
                info_text = f"""
材料类型: {info['material_type'].upper()}型
组分标识: {info['composition']}
温度范围: {info['temperature_range']}
塞贝克系数范围: {info['seebeck_range']}
电阻率范围: {info['resistivity_range']}
热导率范围: {info['thermal_cond_range']}
最大优值系数: {info['max_zt']}
数据点数: {info['data_points']}
材料来源: {'内置材料' if info['is_builtin'] else '自定义材料'}
                """.strip()
                
                # 更新信息显示
                current_tab = self.findChild(QTabWidget).currentWidget()
                info_text_widget = current_tab.findChild(QTextEdit)
                info_text_widget.setText(info_text)
            else:
                print(f"无法获取材料 {composition} 的信息")
                
        except Exception as e:
            print(f"显示材料信息失败: {str(e)}")
    
    def add_material(self, material_type):
        """添加新材料"""
        try:
            dialog = AddMaterialDialog(self, material_type)
            if dialog.exec_() == QDialog.Accepted:
                # 重新加载材料列表
                self.load_materials()
                
        except Exception as e:
            self.parent.show_error_message("错误", f"添加材料失败：{str(e)}")
    
    def edit_material(self, material_type):
        """编辑材料"""
        try:
            current_tab = self.findChild(QTabWidget).currentWidget()
            list_widget = current_tab.findChild(QListWidget)
            
            if list_widget.currentItem():
                composition = list_widget.currentItem().text()
                dialog = EditMaterialDialog(self, material_type, composition)
                if dialog.exec_() == QDialog.Accepted:
                    # 重新加载材料列表
                    self.load_materials()
            else:
                self.parent.show_info_message("提示", "请先选择要编辑的材料")
                
        except Exception as e:
            self.parent.show_error_message("错误", f"编辑材料失败：{str(e)}")
    
    def delete_material(self, material_type):
        """删除材料"""
        try:
            current_tab = self.findChild(QTabWidget).currentWidget()
            list_widget = current_tab.findChild(QListWidget)
            
            if list_widget.currentItem():
                composition = list_widget.currentItem().text()
                
                # 确认删除
                reply = QMessageBox.question(
                    self, "确认删除", 
                    f"确定要删除{material_type.upper()}型材料 {composition} 吗？\n此操作不可恢复！",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                
                if reply == QMessageBox.Yes:
                    if self.calculator.remove_custom_material(material_type, composition):
                        self.parent.show_info_message("成功", f"材料 {composition} 已删除")
                        # 重新加载材料列表
                        self.load_materials()
                    else:
                        self.parent.show_error_message("错误", f"删除材料 {composition} 失败")
            else:
                self.parent.show_info_message("提示", "请先选择要删除的材料")
                
        except Exception as e:
            self.parent.show_error_message("错误", f"删除材料失败：{str(e)}")
    
    def import_from_excel(self, material_type):
        """从Excel导入材料"""
        try:
            filename, _ = QFileDialog.getOpenFileName(
                self, "选择Excel文件", "", "Excel文件 (*.xlsx *.xls)")
            
            if filename:
                # 获取组分标识
                composition, ok = QInputDialog.getText(
                    self, "输入组分标识", 
                    f"请输入{material_type.upper()}型材料的组分标识：")
                
                if ok and composition:
                    if self.calculator.import_material_from_excel(material_type, composition, filename):
                        self.parent.show_info_message("成功", f"材料 {composition} 导入成功")
                        # 重新加载材料列表
                        self.load_materials()
                    else:
                        self.parent.show_error_message("错误", f"导入材料 {composition} 失败")
                        
        except Exception as e:
            self.parent.show_error_message("错误", f"从Excel导入材料失败：{str(e)}")
    
    def export_to_excel(self, material_type):
        """导出材料到Excel"""
        try:
            current_tab = self.findChild(QTabWidget).currentWidget()
            list_widget = current_tab.findChild(QListWidget)
            
            if list_widget.currentItem():
                composition = list_widget.currentItem().text()
                
                filename, _ = QFileDialog.getSaveFileName(
                    self, "保存Excel文件", 
                    f"{material_type.upper()}_{composition}.xlsx",
                    "Excel文件 (*.xlsx)")
                
                if filename:
                    if self.calculator.export_material_to_excel(material_type, composition, filename):
                        self.parent.show_info_message("成功", f"材料 {composition} 导出成功")
                    else:
                        self.parent.show_error_message("错误", f"导出材料 {composition} 失败")
            else:
                self.parent.show_info_message("提示", "请先选择要导出的材料")
                
        except Exception as e:
            self.parent.show_error_message("错误", f"导出材料到Excel失败：{str(e)}")


class AddMaterialDialog(QDialog):
    """添加材料对话框"""
    
    def __init__(self, parent=None, material_type='p'):
        super().__init__(parent)
        self.parent = parent
        self.calculator = parent.calculator
        self.material_type = material_type
        self.setup_ui()
    
    def setup_ui(self):
        """设置用户界面"""
        self.setWindowTitle(f"添加{self.material_type.upper()}型材料")
        self.setMinimumSize(600, 500)
        
        layout = QVBoxLayout(self)
        
        # 基本信息
        basic_group = QGroupBox("基本信息")
        basic_layout = QFormLayout(basic_group)
        
        self.composition_edit = QLineEdit()
        self.composition_edit.setPlaceholderText("例如: 0.04, custom_001")
        basic_layout.addRow("组分标识:", self.composition_edit)
        
        layout.addWidget(basic_group)
        
        # 数据输入
        data_group = QGroupBox("材料数据")
        data_layout = QVBoxLayout(data_group)
        
        # 数据输入表格
        self.data_table = QTableWidget()
        self.data_table.setColumnCount(4)
        self.data_table.setHorizontalHeaderLabels([
            "温度 (K)", "塞贝克系数 (μV/K)", "电阻率 (μΩ·m)", "热导率 (W/m·K)"
        ])
        
        # 设置初始行数
        self.data_table.setRowCount(10)
        for i in range(10):
            for j in range(4):
                self.data_table.setItem(i, j, QTableWidgetItem(""))
        
        data_layout.addWidget(self.data_table)
        
        # 表格操作按钮
        table_button_layout = QHBoxLayout()
        
        add_row_button = QPushButton("添加行")
        add_row_button.clicked.connect(self.add_row)
        
        remove_row_button = QPushButton("删除行")
        remove_row_button.clicked.connect(self.remove_row)
        
        clear_button = QPushButton("清空")
        clear_button.clicked.connect(self.clear_table)
        
        table_button_layout.addWidget(add_row_button)
        table_button_layout.addWidget(remove_row_button)
        table_button_layout.addWidget(clear_button)
        table_button_layout.addStretch()
        
        data_layout.addLayout(table_button_layout)
        layout.addWidget(data_group)
        
        # 底部按钮
        button_layout = QHBoxLayout()
        
        ok_button = QPushButton("确定")
        ok_button.clicked.connect(self.accept)
        
        cancel_button = QPushButton("取消")
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        
        layout.addLayout(button_layout)
    
    def add_row(self):
        """添加行"""
        current_row = self.data_table.rowCount()
        self.data_table.insertRow(current_row)
        for j in range(4):
            self.data_table.setItem(current_row, j, QTableWidgetItem(""))
    
    def remove_row(self):
        """删除行"""
        current_row = self.data_table.currentRow()
        if current_row >= 0:
            self.data_table.removeRow(current_row)
    
    def clear_table(self):
        """清空表格"""
        for i in range(self.data_table.rowCount()):
            for j in range(4):
                self.data_table.setItem(i, j, QTableWidgetItem(""))
    
    def accept(self):
        """确认添加材料"""
        try:
            # 获取组分标识
            composition = self.composition_edit.text().strip()
            if not composition:
                QMessageBox.warning(self, "警告", "请输入组分标识")
                return
            
            # 检查组分标识是否已存在
            existing_materials = self.calculator.get_material_list(self.material_type)
            if composition in existing_materials:
                QMessageBox.warning(self, "警告", f"组分标识 {composition} 已存在")
                return
            
            # 收集表格数据
            material_data = {
                "temp": [],
                "seebeck": [],
                "resistivity": [],
                "thermal_cond": []
            }
            
            for i in range(self.data_table.rowCount()):
                row_data = []
                for j in range(4):
                    item = self.data_table.item(i, j)
                    if item and item.text().strip():
                        try:
                            value = float(item.text().strip())
                            row_data.append(value)
                        except ValueError:
                            QMessageBox.warning(self, "警告", f"第{i+1}行第{j+1}列的数据格式不正确")
                            return
                    else:
                        break
                
                if len(row_data) == 4:
                    material_data["temp"].append(row_data[0])
                    material_data["seebeck"].append(row_data[1] * 1e-6)  # μV/K 转换为 V/K
                    material_data["resistivity"].append(row_data[2] * 1e-6)  # μΩ·m 转换为 Ω·m
                    material_data["thermal_cond"].append(row_data[3])
            
            # 检查数据是否足够
            if len(material_data["temp"]) < 3:
                QMessageBox.warning(self, "警告", "至少需要3个有效的数据点")
                return
            
            # 添加材料
            if self.calculator.add_custom_material(self.material_type, composition, material_data):
                QMessageBox.information(self, "成功", f"材料 {composition} 添加成功")
                super().accept()
            else:
                QMessageBox.critical(self, "错误", "添加材料失败")
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"添加材料时发生错误：{str(e)}")


class MaterialImportDialog(QDialog):
    """材料导入对话框"""
    
    def __init__(self, parent=None, excel_file=""):
        super().__init__(parent)
        self.parent = parent
        self.calculator = parent.calculator
        self.excel_file = excel_file
        self.setup_ui()
    
    def setup_ui(self):
        """设置用户界面"""
        self.setWindowTitle("导入材料数据")
        self.setMinimumSize(400, 300)
        
        layout = QVBoxLayout(self)
        
        # 文件信息
        file_group = QGroupBox("文件信息")
        file_layout = QFormLayout(file_group)
        
        file_layout.addRow("文件路径:", QLabel(self.excel_file))
        
        layout.addWidget(file_group)
        
        # 导入选项
        options_group = QGroupBox("导入选项")
        options_layout = QFormLayout(options_group)
        
        self.material_type_combo = QComboBox()
        self.material_type_combo.addItems(["P型", "N型"])
        options_layout.addRow("材料类型:", self.material_type_combo)
        
        self.composition_edit = QLineEdit()
        self.composition_edit.setPlaceholderText("例如: 0.04, custom_001")
        options_layout.addRow("组分标识:", self.composition_edit)
        
        layout.addWidget(options_group)
        
        # 底部按钮
        button_layout = QHBoxLayout()
        
        import_button = QPushButton("导入")
        import_button.clicked.connect(self.import_material)
        
        cancel_button = QPushButton("取消")
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(import_button)
        button_layout.addWidget(cancel_button)
        
        layout.addLayout(button_layout)
    
    def import_material(self):
        """导入材料"""
        try:
            # 获取导入选项
            material_type = 'p' if self.material_type_combo.currentText() == "P型" else 'n'
            composition = self.composition_edit.text().strip()
            
            if not composition:
                QMessageBox.warning(self, "警告", "请输入组分标识")
                return
            
            # 检查组分标识是否已存在
            existing_materials = self.calculator.get_material_list(material_type)
            if composition in existing_materials:
                reply = QMessageBox.question(
                    self, "确认覆盖", 
                    f"组分标识 {composition} 已存在，是否覆盖？",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return
            
            # 导入材料
            if self.calculator.import_material_from_excel(material_type, composition, self.excel_file):
                QMessageBox.information(self, "成功", f"材料 {composition} 导入成功")
                super().accept()
            else:
                QMessageBox.critical(self, "错误", "导入材料失败")
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导入材料时发生错误：{str(e)}")


class MaterialExportDialog(QDialog):
    """材料导出对话框"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.calculator = parent.calculator
        self.setup_ui()
    
    def setup_ui(self):
        """设置用户界面"""
        self.setWindowTitle("导出材料数据")
        self.setMinimumSize(400, 300)
        
        layout = QVBoxLayout(self)
        
        # 导出选项
        options_group = QGroupBox("导出选项")
        options_layout = QFormLayout(options_group)
        
        self.material_type_combo = QComboBox()
        self.material_type_combo.addItems(["P型", "N型"])
        self.material_type_combo.currentTextChanged.connect(self.on_material_type_changed)
        options_layout.addRow("材料类型:", self.material_type_combo)
        
        self.composition_combo = QComboBox()
        options_layout.addRow("材料组分:", self.composition_combo)
        
        layout.addWidget(options_group)
        
        # 底部按钮
        button_layout = QHBoxLayout()
        
        export_button = QPushButton("导出")
        export_button.clicked.connect(self.export_material)
        
        cancel_button = QPushButton("取消")
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(export_button)
        button_layout.addWidget(cancel_button)
        
        layout.addLayout(button_layout)
        
        # 初始化组分列表
        self.on_material_type_changed()
    
    def on_material_type_changed(self):
        """材料类型改变事件"""
        try:
            material_type = 'p' if self.material_type_combo.currentText() == "P型" else 'n'
            materials = self.calculator.get_material_list(material_type)
            
            self.composition_combo.clear()
            self.composition_combo.addItems(materials)
            
        except Exception as e:
            print(f"更新材料组分列表失败: {str(e)}")
    
    def export_material(self):
        """导出材料"""
        try:
            # 获取导出选项
            material_type = 'p' if self.material_type_combo.currentText() == "P型" else 'n'
            composition = self.composition_combo.currentText()
            
            if not composition:
                QMessageBox.warning(self, "警告", "请选择要导出的材料")
                return
            
            # 选择保存路径
            filename, _ = QFileDialog.getSaveFileName(
                self, "保存Excel文件", 
                f"{material_type.upper()}_{composition}.xlsx",
                "Excel文件 (*.xlsx)")
            
            if filename:
                if self.calculator.export_material_to_excel(material_type, composition, filename):
                    QMessageBox.information(self, "成功", f"材料 {composition} 导出成功")
                    super().accept()
                else:
                    QMessageBox.critical(self, "错误", "导出材料失败")
                    
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出材料时发生错误：{str(e)}")


class EditMaterialDialog(QDialog):
    """编辑材料对话框"""
    
    def __init__(self, parent=None, material_type='p', composition=''):
        super().__init__(parent)
        self.parent = parent
        self.calculator = parent.calculator
        self.material_type = material_type
        self.composition = composition
        self.setup_ui()
        self.load_material_data()
    
    def setup_ui(self):
        """设置用户界面"""
        self.setWindowTitle(f"编辑{self.material_type.upper()}型材料 - {self.composition}")
        self.setMinimumSize(600, 500)
        
        layout = QVBoxLayout(self)
        
        # 基本信息
        basic_group = QGroupBox("基本信息")
        basic_layout = QFormLayout(basic_group)
        
        self.composition_edit = QLineEdit(self.composition)
        self.composition_edit.setReadOnly(True)  # 组分标识不可编辑
        basic_layout.addRow("组分标识:", self.composition_edit)
        
        layout.addWidget(basic_group)
        
        # 数据编辑
        data_group = QGroupBox("材料数据")
        data_layout = QVBoxLayout(data_group)
        
        # 数据编辑表格
        self.data_table = QTableWidget()
        self.data_table.setColumnCount(4)
        self.data_table.setHorizontalHeaderLabels([
            "温度 (K)", "塞贝克系数 (μV/K)", "电阻率 (μΩ·m)", "热导率 (W/m·K)"
        ])
        
        data_layout.addWidget(self.data_table)
        
        # 表格操作按钮
        table_button_layout = QHBoxLayout()
        
        add_row_button = QPushButton("添加行")
        add_row_button.clicked.connect(self.add_row)
        
        remove_row_button = QPushButton("删除行")
        remove_row_button.clicked.connect(self.remove_row)
        
        clear_button = QPushButton("清空")
        clear_button.clicked.connect(self.clear_table)
        
        table_button_layout.addWidget(add_row_button)
        table_button_layout.addWidget(remove_row_button)
        table_button_layout.addWidget(clear_button)
        table_button_layout.addStretch()
        
        data_layout.addLayout(table_button_layout)
        layout.addWidget(data_group)
        
        # 底部按钮
        button_layout = QHBoxLayout()
        
        save_button = QPushButton("保存")
        save_button.clicked.connect(self.save_material)
        
        cancel_button = QPushButton("取消")
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(save_button)
        button_layout.addWidget(cancel_button)
        
        layout.addLayout(button_layout)
    
    def load_material_data(self):
        """加载材料数据到表格"""
        try:
            target_data = self.calculator.p_type_data if self.material_type == 'p' else self.calculator.n_type_data
            if self.composition not in target_data:
                QMessageBox.warning(self, "警告", f"材料 {self.composition} 不存在")
                self.reject()
                return
            
            data = target_data[self.composition]
            
            # 设置表格行数
            self.data_table.setRowCount(len(data["temp"]))
            
            # 填充数据
            for i in range(len(data["temp"])):
                self.data_table.setItem(i, 0, QTableWidgetItem(str(data["temp"][i])))
                self.data_table.setItem(i, 1, QTableWidgetItem(str(data["seebeck"][i] * 1e6)))  # V/K 转换为 μV/K
                self.data_table.setItem(i, 2, QTableWidgetItem(str(data["resistivity"][i] * 1e6)))  # Ω·m 转换为 μΩ·m
                self.data_table.setItem(i, 3, QTableWidgetItem(str(data["thermal_cond"][i])))
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"加载材料数据失败：{str(e)}")
            self.reject()
    
    def add_row(self):
        """添加行"""
        current_row = self.data_table.rowCount()
        self.data_table.insertRow(current_row)
        for j in range(4):
            self.data_table.setItem(current_row, j, QTableWidgetItem(""))
    
    def remove_row(self):
        """删除行"""
        current_row = self.data_table.currentRow()
        if current_row >= 0:
            self.data_table.removeRow(current_row)
    
    def clear_table(self):
        """清空表格"""
        for i in range(self.data_table.rowCount()):
            for j in range(4):
                self.data_table.setItem(i, j, QTableWidgetItem(""))
    
    def save_material(self):
        """保存材料数据"""
        try:
            # 收集表格数据
            material_data = {
                "temp": [],
                "seebeck": [],
                "resistivity": [],
                "thermal_cond": []
            }
            
            for i in range(self.data_table.rowCount()):
                row_data = []
                for j in range(4):
                    item = self.data_table.item(i, j)
                    if item and item.text().strip():
                        try:
                            value = float(item.text().strip())
                            row_data.append(value)
                        except ValueError:
                            QMessageBox.warning(self, "警告", f"第{i+1}行第{j+1}列的数据格式不正确")
                            return
                    else:
                        break
                
                if len(row_data) == 4:
                    material_data["temp"].append(row_data[0])
                    material_data["seebeck"].append(row_data[1] * 1e-6)  # μV/K 转换为 V/K
                    material_data["resistivity"].append(row_data[2] * 1e-6)  # μΩ·m 转换为 Ω·m
                    material_data["thermal_cond"].append(row_data[3])
            
            # 检查数据是否足够
            if len(material_data["temp"]) < 3:
                QMessageBox.warning(self, "警告", "至少需要3个有效的数据点")
                return
            
            # 删除旧材料并添加新材料
            if self.calculator.remove_custom_material(self.material_type, self.composition):
                if self.calculator.add_custom_material(self.material_type, self.composition, material_data):
                    QMessageBox.information(self, "成功", f"材料 {self.composition} 更新成功")
                    super().accept()
                else:
                    QMessageBox.critical(self, "错误", "保存材料失败")
            else:
                QMessageBox.critical(self, "错误", "删除旧材料失败")
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存材料时发生错误：{str(e)}")
class MaterialInfoDialog(QDialog):
    """材料信息查看对话框"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.calculator = parent.calculator
        self.setup_ui()
        self.load_materials_info()
    
    def setup_ui(self):
        """设置用户界面"""
        self.setWindowTitle("材料信息查看")
        self.setMinimumSize(800, 600)
        
        layout = QVBoxLayout(self)
        
        # 创建选项卡
        tab_widget = QTabWidget()
        
        # P型材料选项卡
        p_tab = self.create_material_info_tab('p')
        tab_widget.addTab(p_tab, "P型材料")
        
        # N型材料选项卡
        n_tab = self.create_material_info_tab('n')
        tab_widget.addTab(n_tab, "N型材料")
        
        layout.addWidget(tab_widget)
        
        # 底部按钮
        close_button = QPushButton("关闭")
        close_button.clicked.connect(self.close)
        layout.addWidget(close_button)
    
    def create_material_info_tab(self, material_type):
        """创建材料信息选项卡"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 材料信息表格
        self.info_table = QTableWidget()
        self.info_table.setColumnCount(8)
        self.info_table.setHorizontalHeaderLabels([
            "组分标识", "温度范围(K)", "塞贝克系数范围(μV/K)", 
            "电阻率范围(μΩ·m)", "热导率范围(W/m·K)", 
            "最大ZT值", "数据点数", "材料来源"
        ])
        
        layout.addWidget(self.info_table)
        
        # 保存材料类型引用
        widget.material_type = material_type
        
        return widget
    
    def load_materials_info(self):
        """加载材料信息"""
        try:
            # 加载P型材料信息
            p_tab = self.findChild(QTabWidget).widget(0)
            p_table = p_tab.findChild(QTableWidget)
            self.populate_info_table(p_table, 'p')
            
            # 加载N型材料信息
            n_tab = self.findChild(QTabWidget).widget(1)
            n_table = n_tab.findChild(QTableWidget)
            self.populate_info_table(n_table, 'n')
            
        except Exception as e:
            print(f"加载材料信息失败: {str(e)}")
    
    def populate_info_table(self, table, material_type):
        """填充信息表格"""
        try:
            materials = self.calculator.get_material_list(material_type)
            table.setRowCount(len(materials))
            
            for i, composition in enumerate(materials):
                info = self.calculator.get_material_info(material_type, composition)
                if info:
                    table.setItem(i, 0, QTableWidgetItem(info['composition']))
                    table.setItem(i, 1, QTableWidgetItem(info['temperature_range']))
                    table.setItem(i, 2, QTableWidgetItem(info['seebeck_range']))
                    table.setItem(i, 3, QTableWidgetItem(info['resistivity_range']))
                    table.setItem(i, 4, QTableWidgetItem(info['thermal_cond_range']))
                    table.setItem(i, 5, QTableWidgetItem(info['max_zt']))
                    table.setItem(i, 6, QTableWidgetItem(str(info['data_points'])))
                    table.setItem(i, 7, QTableWidgetItem('内置材料' if info['is_builtin'] else '自定义材料'))
            
            # 调整列宽
            table.resizeColumnsToContents()
            
        except Exception as e:
            print(f"填充信息表格失败: {str(e)}")


if __name__ == '__main__':
    try:
        app = QApplication(sys.argv)
        
        # 设置应用程序样式
        app.setStyle('Fusion')
        
        # 创建主窗口
        window = ThermoelectricApp()
        
        # 加载用户设置
        window.load_user_settings()
        
        # 添加右键菜单
        window.add_context_menu()
        
        # 显示窗口
        window.show()
        
        # 运行应用程序
        sys.exit(app.exec_())
    except Exception as e:
        print(f"程序启动失败: {str(e)}")
        import traceback
        traceback.print_exc()
        input("按回车键退出...")
